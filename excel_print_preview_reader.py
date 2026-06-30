import re
import html
import base64
from pathlib import Path
from datetime import datetime
from io import BytesIO

import pandas as pd
import streamlit as st
import streamlit.components.v1 as components

from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_CELL_VERTICAL_ALIGNMENT
from docx.oxml import OxmlElement
from docx.oxml.ns import qn

APP_TITLE = "PLAN STRATÉGIQUE USJ 2032"
SHEET_NAME = "All answers 29-6-2026"

USJ_BLUE = "#001F5B"
USJ_BLUE_2 = "#1F3C88"
USJ_RED = "#8B1538"
USJ_LIGHT_BLUE = "#EAF2F8"
USJ_TEXT = "#1B2A41"
LOGO_PATH = "LOGO NEW UAQ.png"


def clean(value):
    if pd.isna(value):
        return ""
    return str(value).strip()


def normalize(value):
    value = clean(value).lower()
    replacements = {
        "é": "e", "è": "e", "ê": "e", "ë": "e",
        "à": "a", "â": "a",
        "ù": "u", "û": "u",
        "ç": "c",
        "î": "i", "ï": "i",
        "ô": "o",
    }
    for old, new in replacements.items():
        value = value.replace(old, new)
    return value


def esc(value):
    return html.escape(clean(value)).replace("\n", "<br>")


def safe_filename(value):
    value = clean(value)
    value = re.sub(r"[^A-Za-z0-9À-ÿ._-]+", "_", value)
    return value.strip("_") or "rapport"


def html_block(content):
    st.markdown(content, unsafe_allow_html=True)


def image_to_base64(image_path):
    image_path = Path(image_path)
    if not image_path.exists():
        return ""
    suffix = image_path.suffix.lower().replace(".", "")
    mime_type = "png" if suffix == "png" else suffix
    with open(image_path, "rb") as img_file:
        encoded = base64.b64encode(img_file.read()).decode()
    return f"data:image/{mime_type};base64,{encoded}"


PRINT_CSS = f"""
body {{
    font-family: Candara, Calibri, Arial, sans-serif;
    background:white;
    color:{USJ_TEXT};
    margin:0;
    padding:0;
}}

.print-report {{
    background:white;
    border:1px solid #D0D6E0;
    border-radius:12px;
    padding:28px 34px;
}}

.report-page {{
    background:white;
    border:none;
    border-radius:0;
    padding:0;
    margin:0;
    width:auto;
    min-height:auto;
    box-sizing:border-box;
}}

.report-page + .report-page {{
    margin-top:28px;
}}

.usj-main-header {{
    display:flex;
    justify-content:space-between;
    align-items:flex-start;
    border-bottom:2px solid #D0D6E0;
    padding-bottom:18px;
    margin-bottom:28px;
}}

.usj-main-header h1 {{
    font-size:42px;
    margin:0 0 6px 0;
    color:{USJ_BLUE};
    font-weight:800;
}}

.usj-main-header p {{
    font-size:18px;
    font-weight:700;
    color:{USJ_BLUE_2};
    margin:0;
}}

.usj-logo-box {{
    display:flex;
    justify-content:flex-end;
    align-items:flex-start;
}}

.usj-logo {{
    width:170px;
    height:auto;
    object-fit:contain;
}}

.group-title {{
    text-align:center;
    font-size:28px;
    font-weight:800;
    color:{USJ_RED};
    margin:24px 0;
}}

.names {{
    text-align:center;
    font-size:21px;
    font-weight:700;
    color:#111;
    margin-bottom:32px;
}}

.section-header {{
    background-color:{USJ_LIGHT_BLUE};
    padding:12px 18px;
    border-radius:10px;
    border-left:7px solid {USJ_BLUE};
    margin-top:24px;
    margin-bottom:16px;
    page-break-inside:avoid;
    box-shadow:0 2px 10px rgba(0,0,0,0.06);
}}

.section-header h2 {{
    font-size:28px;
    margin:0;
    color:{USJ_BLUE};
}}

.two-cols {{
    display:grid;
    grid-template-columns:1fr 1fr;
    gap:20px;
}}

.col-title {{
    background:white;
    color:{USJ_BLUE};
    padding:10px 14px;
    text-align:center;
    font-size:21px;
    font-weight:900;
    border-radius:6px;
    margin-bottom:10px;
}}

.answer-box {{
    background:#E3DED9;
    border:1.5px solid #595959;
    padding:12px 16px;
    margin-bottom:10px;
    font-size:18px;
    line-height:1.35;
    color:#000;
    page-break-inside:avoid;
    break-inside:avoid;
}}

.answer-number {{
    color:{USJ_RED};
    font-weight:800;
    margin-right:6px;
}}

.conclusion-label {{
    color:{USJ_BLUE};
    font-size:19px;
    font-weight:800;
    margin:28px 0 10px 4px;
}}

.conclusion-box {{
    background:#E3DED9;
    border:1.5px solid #595959;
    padding:14px 18px;
    margin-bottom:18px;
    min-height:72px;
    font-size:20px;
    line-height:1.35;
    color:#000;
    page-break-inside:avoid;
    break-inside:avoid;
}}

.swot-section-wrapper {{
    break-inside:avoid;
    page-break-inside:avoid;
}}

.swot-matrix-title {{
    background-color:{USJ_LIGHT_BLUE};
    padding:12px 18px;
    border-radius:10px;
    border-left:7px solid {USJ_BLUE};
    margin-top:24px;
    margin-bottom:16px;
    color:{USJ_BLUE};
    font-size:28px;
    font-weight:800;
    page-break-after:avoid;
    break-after:avoid;
}}

.swot-grid {{
    display:grid;
    grid-template-columns:1fr 1fr;
    gap:18px;
    margin-top:14px;
    page-break-before:avoid;
    break-before:avoid;
}}

.swot-card {{
    border-radius:14px;
    padding:16px;
    border:2px solid var(--accent);
    background:var(--bg);
    page-break-inside:avoid;
    break-inside:avoid;
}}

.swot-card h3 {{
    margin:0 0 10px 0;
    color:var(--accent);
    font-size:24px;
    font-weight:900;
    border-bottom:2px solid var(--accent);
    padding-bottom:8px;
}}

.swot-card ul {{
    margin:0;
    padding-left:22px;
}}

.swot-card li {{
    margin-bottom:8px;
    font-size:16px;
    line-height:1.3;
}}

.print-page-break {{
    page-break-before:always;
}}

.print-button {{
    display:inline-block;
    background:#8B1538;
    color:white;
    border:1px solid #8B1538;
    border-radius:8px;
    padding:10px 22px;
    font-family:Candara,Calibri,Arial,sans-serif;
    font-size:18px;
    font-weight:800;
    cursor:pointer;
    margin:0 12px 20px 0;
    text-decoration:none;
}}

.print-actions {{
    margin-bottom:20px;
}}

@media print {{
    @page {{
        size:A4 portrait;
        margin:8mm 9mm;
    }}

    body {{
        margin:0;
    }}

    .print-button {{
        display:none;
    }}

    .print-report {{
        border:none;
        padding:0;
        background:white;
    }}

    .report-page {{
        border:none;
        border-radius:0;
        padding:0;
        margin:0;
        width:auto;
        min-height:auto;
        page-break-after:always;
        break-after:page;
    }}

    .report-page:last-child {{
        page-break-after:auto;
        break-after:auto;
    }}

    .usj-main-header h1 {{
        font-size:22px;
        line-height:1.05;
        margin-bottom:1mm;
    }}

    .usj-main-header p {{
        font-size:10px;
    }}

    .usj-logo {{
        width:36mm;
        max-width:36mm;
        height:auto;
    }}

    .group-title {{
        font-size:18px;
        margin:8mm 0 4mm 0;
    }}

    .names {{
        font-size:12px;
        margin-bottom:8mm;
    }}

    .section-header {{
        margin-top:5mm;
        margin-bottom:3mm;
        padding:7px 11px;
        box-shadow:none;
    }}

    .section-header h2 {{
        font-size:16px;
    }}

    .swot-section-wrapper {{
        break-inside:avoid;
        page-break-inside:avoid;
    }}

    .swot-matrix-title {{
        font-size:16px;
        margin-top:5mm;
        margin-bottom:3mm;
        padding:7px 11px;
        page-break-after:avoid;
        break-after:avoid;
    }}

    .col-title {{
        font-size:13px;
        padding:6px;
        color:{USJ_BLUE};
        background:white;
        font-weight:900;
    }}

    .answer-box,
    .conclusion-box {{
        font-size:10.5px;
        line-height:1.2;
        padding:6px;
        margin-bottom:2mm;
        page-break-inside:avoid;
        break-inside:avoid;
    }}

    .conclusion-label {{
        font-size:11px;
        margin:4mm 0 2mm 0;
    }}

    .swot-card h3 {{
        font-size:13px;
    }}

    .swot-card li {{
        font-size:9.5px;
        line-height:1.15;
    }}
}}
"""


APP_CSS = f"""
<style>
html, body, .stApp {{
    font-family: Candara, Calibri, Arial, sans-serif !important;
    color:{USJ_TEXT};
}}

header, footer, #MainMenu {{
    display:none !important;
}}

.block-container {{
    max-width:100% !important;
    padding-left:3rem !important;
    padding-right:3rem !important;
}}

h1, h2, h3 {{
    color:{USJ_BLUE} !important;
    font-weight:800 !important;
}}

.stButton button,
.stDownloadButton button {{
    background-color:#0070C0 !important;
    color:white !important;
    border-radius:8px !important;
    border:1px solid #0070C0 !important;
    font-weight:800 !important;
    font-size:18px !important;
}}

.stButton button p,
.stDownloadButton button p {{
    color:white !important;
}}
</style>
"""


def get_col(df, possible_names):
    cols = {normalize(c): c for c in df.columns}
    for name in possible_names:
        n = normalize(name)
        if n in cols:
            return cols[n]
    return None


def get_answers(df_group, section_contains=None, category_contains=None):
    temp = df_group.copy()

    if section_contains:
        temp = temp[temp["section_norm"].str.contains(section_contains, na=False)]

    if category_contains:
        temp = temp[temp["category_norm"].str.contains(category_contains, na=False)]

    return [clean(x) for x in temp["Final_Answer"].tolist() if clean(x)]


def answer_boxes(answers):
    if not answers:
        return '<div class="answer-box">Aucune réponse disponible.</div>'

    return "".join(
        f"""
        <div class="answer-box">
            <span class="answer-number">{i}.</span>{esc(answer)}
        </div>
        """
        for i, answer in enumerate(answers, start=1)
    )


def swot_card(title, answers, accent, bg):
    if not answers:
        li = "<li>Aucune réponse disponible.</li>"
    else:
        li = "".join(f"<li>{esc(x)}</li>" for x in answers)

    return f"""
    <div class="swot-card" style="--accent:{accent}; --bg:{bg};">
        <h3>{title}</h3>
        <ul>{li}</ul>
    </div>
    """


def swot_matrix_html(forces, faiblesses, opportunites, menaces):
    return f"""
    <div class="swot-section-wrapper">
        <div class="swot-matrix-title">
            Matrice SWOT - Niveau USJ
        </div>

        <div class="swot-grid">
            {swot_card("FORCES", forces, USJ_BLUE, "#DDEFF7")}
            {swot_card("FAIBLESSES", faiblesses, USJ_RED, "#FBE3C3")}
            {swot_card("OPPORTUNITÉS", opportunites, "#2F6B2F", "#E2F2D3")}
            {swot_card("MENACES", menaces, USJ_RED, "#F4C6C4")}
        </div>
    </div>
    """


def conclusion_html(df_group):
    conclusion_rows = df_group[
        df_group["section_norm"].str.contains("conclusion", na=False)
    ].copy()

    if conclusion_rows.empty:
        return '<div class="conclusion-box">Aucune réponse disponible.</div>'

    blocks = ""
    question_order = []

    for question in conclusion_rows["question"].tolist():
        question = clean(question)
        if question and question not in question_order:
            question_order.append(question)

    for question in question_order:
        temp = conclusion_rows[conclusion_rows["question"] == question]
        answers = [clean(x) for x in temp["Final_Answer"].tolist() if clean(x)]

        blocks += f"""
        <div class="conclusion-label">• {esc(question)}</div>
        """

        if not answers:
            blocks += '<div class="conclusion-box"></div>'
        else:
            for answer in answers:
                blocks += f"""
                <div class="conclusion-box">{esc(answer)}</div>
                """

    return blocks


def build_one_report_html(df_group, participant_type, title_label, hide_names):
    names = sorted(set(clean(x) for x in df_group["participants"].dropna() if clean(x)))

    names_html = ""
    if not hide_names and names:
        names_html = f"""
        <div class="names">
            Nom des participants : {esc("; ".join(names))}
        </div>
        """

    logo_src = image_to_base64(LOGO_PATH)
    logo_html = f'<img src="{logo_src}" class="usj-logo">' if logo_src else ""

    forces = get_answers(df_group, section_contains="forces", category_contains="force")
    faiblesses = get_answers(df_group, section_contains="forces", category_contains="faiblesse")
    opportunites = get_answers(df_group, section_contains="opportunites", category_contains="opportun")
    menaces = get_answers(df_group, section_contains="opportunites", category_contains="menace")
    priorites = get_answers(df_group, section_contains="prior")

    return f"""
<div class="print-report">

    <div class="report-page">
        <div class="usj-main-header">
            <div>
                <h1>PLAN STRATÉGIQUE USJ 2032</h1>
                <p>Focus groupe</p>
            </div>

            <div class="usj-logo-box">
                {logo_html}
            </div>
        </div>

        <div class="group-title">{esc(title_label)}</div>

        {names_html}

        <div class="section-header">
            <h2>I - Forces et faiblesses</h2>
        </div>

        <div class="two-cols">
            <div>
                <div class="col-title">Forces</div>
                {answer_boxes(forces)}
            </div>
            <div>
                <div class="col-title">Faiblesses</div>
                {answer_boxes(faiblesses)}
            </div>
        </div>
    </div>

    <div class="report-page">
        <div class="section-header">
            <h2>II - Opportunités et menaces</h2>
        </div>

        <div class="two-cols">
            <div>
                <div class="col-title">Opportunités</div>
                {answer_boxes(opportunites)}
            </div>
            <div>
                <div class="col-title">Menaces</div>
                {answer_boxes(menaces)}
            </div>
        </div>

        {swot_matrix_html(forces, faiblesses, opportunites, menaces)}
    </div>

    <div class="report-page">
        <div class="section-header">
            <h2>III - Priorités</h2>
        </div>

        <div>
            {answer_boxes(priorites)}
        </div>
    </div>

    <div class="report-page">
        <div class="section-header">
            <h2>IV - Conclusion</h2>
        </div>

        <div>
            {conclusion_html(df_group)}
        </div>
    </div>

</div>
"""


def build_full_html(html_report, selected_type, selected_label, docx_base64="", docx_filename="rapport.docx"):
    word_button = ""
    if docx_base64:
        word_button = f"""
<a class="print-button"
   href="data:application/vnd.openxmlformats-officedocument.wordprocessingml.document;base64,{docx_base64}"
   download="{esc(docx_filename)}">
   Télécharger Word
</a>
"""

    return f"""
<!DOCTYPE html>
<html>
<head>
<meta charset="UTF-8">
<title>{esc(selected_type)} - {esc(selected_label)}</title>
<style>
{PRINT_CSS}
</style>
</head>
<body>
<div class="print-actions">
    <button class="print-button" onclick="window.print()">Imprimer / Enregistrer en PDF</button>
    {word_button}
</div>
{html_report}
</body>
</html>
"""



def set_cell_shading(cell, fill):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:fill"), fill.replace("#", ""))


def set_cell_border(cell, color="595959", size="8"):
    tc_pr = cell._tc.get_or_add_tcPr()
    borders = tc_pr.first_child_found_in("w:tcBorders")
    if borders is None:
        borders = OxmlElement("w:tcBorders")
        tc_pr.append(borders)
    for edge in ("top", "left", "bottom", "right"):
        tag = "w:{}".format(edge)
        element = borders.find(qn(tag))
        if element is None:
            element = OxmlElement(tag)
            borders.append(element)
        element.set(qn("w:val"), "single")
        element.set(qn("w:sz"), size)
        element.set(qn("w:space"), "0")
        element.set(qn("w:color"), color.replace("#", ""))


def clear_cell(cell):
    cell.text = ""
    if not cell.paragraphs:
        cell.add_paragraph()
    return cell.paragraphs[0]


def set_paragraph_font(paragraph, size=10.5, bold=False, color="000000"):
    for run in paragraph.runs:
        run.font.name = "Candara"
        run._element.rPr.rFonts.set(qn("w:eastAsia"), "Candara")
        run.font.size = Pt(size)
        run.font.bold = bold
        run.font.color.rgb = RGBColor.from_string(color.replace("#", ""))


def add_word_header(document, participant_type):
    table = document.add_table(rows=1, cols=2)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.autofit = False
    table.columns[0].width = Inches(5.1)
    table.columns[1].width = Inches(1.5)

    for row in table.rows:
        for cell in row.cells:
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.TOP
            set_cell_border(cell, "FFFFFF", "0")

    left = table.cell(0, 0)
    right = table.cell(0, 1)

    p = clear_cell(left)

    r = p.add_run("PLAN STRATÉGIQUE USJ 2032")
    r.bold = True
    r.font.name = "Candara"
    r._element.rPr.rFonts.set(qn("w:eastAsia"), "Candara")
    r.font.size = Pt(20)
    r.font.color.rgb = RGBColor(0, 31, 91)
    
    p2 = left.add_paragraph()
    r2 = p2.add_run(f"Focus groupe - {participant_type}")
    r2.bold = True
    r2.font.name = "Candara"
    r2._element.rPr.rFonts.set(qn("w:eastAsia"), "Candara")
    r2.font.size = Pt(10)
    r2.font.color.rgb = RGBColor(31, 60, 136)

    p_logo = clear_cell(right)
    p_logo.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    if Path(LOGO_PATH).exists():
        p_logo.add_run().add_picture(LOGO_PATH, width=Inches(1.42))

    line = document.add_paragraph()
    line.paragraph_format.space_after = Pt(10)
    line_run = line.add_run("_")
    line_run.font.color.rgb = RGBColor(208, 214, 224)
    line_run.font.size = Pt(1)


def add_section_header(document, text):
    table = document.add_table(rows=1, cols=1)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    cell = table.cell(0, 0)
    set_cell_shading(cell, USJ_LIGHT_BLUE)
    set_cell_border(cell, USJ_LIGHT_BLUE, "6")
    cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
    p = clear_cell(cell)
    p.paragraph_format.left_indent = Inches(0.05)
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.space_after = Pt(0)
    r = p.add_run(text)
    r.bold = True
    r.font.name = "Candara"
    r._element.rPr.rFonts.set(qn("w:eastAsia"), "Candara")
    r.font.size = Pt(16)
    r.font.color.rgb = RGBColor(0, 31, 91)
    document.add_paragraph().paragraph_format.space_after = Pt(2)


def add_answer_to_cell(cell, number, answer, font_size=10.5):
    set_cell_shading(cell, "#E3DED9")
    set_cell_border(cell, "595959", "8")
    cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.TOP
    p = clear_cell(cell)
    p.paragraph_format.space_after = Pt(0)
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.left_indent = Inches(0.02)
    r1 = p.add_run(f"{number}. ")
    r1.bold = True
    r1.font.name = "Candara"
    r1._element.rPr.rFonts.set(qn("w:eastAsia"), "Candara")
    r1.font.size = Pt(font_size)
    r1.font.color.rgb = RGBColor(139, 21, 56)
    r2 = p.add_run(answer)
    r2.font.name = "Candara"
    r2._element.rPr.rFonts.set(qn("w:eastAsia"), "Candara")
    r2.font.size = Pt(font_size)
    r2.font.color.rgb = RGBColor(0, 0, 0)


def add_single_answer_box(document, number, answer, font_size=10.5, min_height=None):
    table = document.add_table(rows=1, cols=1)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    cell = table.cell(0, 0)
    if min_height:
        table.rows[0].height = min_height
    add_answer_to_cell(cell, number, answer, font_size=font_size)
    document.add_paragraph().paragraph_format.space_after = Pt(1)


def add_two_column_answers(document, left_title, left_answers, right_title, right_answers):
    max_len = max(len(left_answers), len(right_answers), 1)

    table = document.add_table(rows=max_len + 1, cols=2)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.autofit = False
    table.style = "Table Grid"
    table.columns[0].width = Inches(3.15)
    table.columns[1].width = Inches(3.15)

    # Repeat the header row if the table continues on another page.
    header_tr_pr = table.rows[0]._tr.get_or_add_trPr()
    tbl_header = OxmlElement("w:tblHeader")
    tbl_header.set(qn("w:val"), "true")
    header_tr_pr.append(tbl_header)

    for i, title in enumerate([left_title, right_title]):
        cell = table.cell(0, i)
        set_cell_border(cell, "595959", "8")
        cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
        p = clear_cell(cell)
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p.paragraph_format.space_before = Pt(0)
        p.paragraph_format.space_after = Pt(0)
        r = p.add_run(title)
        r.bold = True
        r.font.name = "Candara"
        r._element.rPr.rFonts.set(qn("w:eastAsia"), "Candara")
        r.font.size = Pt(13)
        r.font.color.rgb = RGBColor(0, 31, 91)

    for i in range(max_len):
        left_text = left_answers[i] if i < len(left_answers) else ""
        right_text = right_answers[i] if i < len(right_answers) else ""

        for col_idx, text in enumerate([left_text, right_text]):
            cell = table.cell(i + 1, col_idx)
            set_cell_border(cell, "595959", "8")
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.TOP
            p = clear_cell(cell)
            p.paragraph_format.space_after = Pt(0)
            p.paragraph_format.space_before = Pt(0)
            p.paragraph_format.left_indent = Inches(0.02)

            if text:
                r = p.add_run(text)
                r.font.name = "Candara"
                r._element.rPr.rFonts.set(qn("w:eastAsia"), "Candara")
                r.font.size = Pt(10.2)
                r.font.color.rgb = RGBColor(0, 0, 0)

    document.add_paragraph().paragraph_format.space_after = Pt(2)


def add_swot_cell(cell, title, answers, accent_hex, bg_hex):
    set_cell_shading(cell, bg_hex)
    set_cell_border(cell, accent_hex, "12")
    p = clear_cell(cell)
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    r = p.add_run(title)
    r.bold = True
    r.font.name = "Candara"
    r._element.rPr.rFonts.set(qn("w:eastAsia"), "Candara")
    r.font.size = Pt(12)
    r.font.color.rgb = RGBColor.from_string(accent_hex.replace("#", ""))
    
    for answer in answers or ["Aucune réponse disponible."]:
        p_item = cell.add_paragraph(style=None)
        p_item.paragraph_format.left_indent = Inches(0.12)
        p_item.paragraph_format.space_after = Pt(2)
    
        answer_clean = answer.replace("\r\n", "\n").replace("\r", "\n").strip()

        if "\n" in answer_clean:
            parts = [x.strip() for x in answer_clean.split("\n") if x.strip()]
        else:
            parts = re.split(r"(?<=[.!?])\s+(?=[A-ZÉÈÊÀÂÎÏÔÙÛÇ])", answer_clean)
            parts = [x.strip() for x in parts if x.strip()]
    
        if parts:
            run = p_item.add_run("• " + parts[0])
            run.font.name = "Candara"
            run._element.rPr.rFonts.set(qn("w:eastAsia"), "Candara")
            run.font.size = Pt(9.2)
            run.font.color.rgb = RGBColor(0, 0, 0)
    
            for extra_part in parts[1:]:
                run_break = p_item.add_run()
                run_break.add_break()
    
                run_extra = p_item.add_run(extra_part)
                run_extra.font.name = "Candara"
                run_extra._element.rPr.rFonts.set(qn("w:eastAsia"), "Candara")
                run_extra.font.size = Pt(9.2)
                run_extra.font.color.rgb = RGBColor(0, 0, 0)
    


def add_swot_matrix_word(document, forces, faiblesses, opportunites, menaces):
    add_section_header(document, "Matrice SWOT - Niveau USJ")
    table = document.add_table(rows=2, cols=2)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.autofit = False
    table.columns[0].width = Inches(3.15)
    table.columns[1].width = Inches(3.15)
    add_swot_cell(table.cell(0, 0), "FORCES", forces, "#001F5B", "#DDEFF7")
    add_swot_cell(table.cell(0, 1), "FAIBLESSES", faiblesses, "#8B1538", "#FBE3C3")
    add_swot_cell(table.cell(1, 0), "OPPORTUNITÉS", opportunites, "#2F6B2F", "#E2F2D3")
    add_swot_cell(table.cell(1, 1), "MENACES", menaces, "#8B1538", "#F4C6C4")
    document.add_paragraph().paragraph_format.space_after = Pt(2)


def add_conclusion_word(document, df_group):
    conclusion_rows = df_group[df_group["section_norm"].str.contains("conclusion", na=False)].copy()
    if conclusion_rows.empty:
        add_single_answer_box(document, 1, "Aucune réponse disponible.", font_size=10.5, min_height=Inches(0.65))
        return

    question_order = []
    for question in conclusion_rows["question"].tolist():
        question = clean(question)
        if question and question not in question_order:
            question_order.append(question)

    for question in question_order:
        p_q = document.add_paragraph()
        p_q.paragraph_format.space_before = Pt(4)
        p_q.paragraph_format.space_after = Pt(3)
        r_q = p_q.add_run("• " + question)
        r_q.bold = True
        r_q.font.name = "Candara"
        r_q._element.rPr.rFonts.set(qn("w:eastAsia"), "Candara")
        r_q.font.size = Pt(11)
        r_q.font.color.rgb = RGBColor(0, 31, 91)

        answers = [
            clean(x)
            for x in conclusion_rows[conclusion_rows["question"] == question]["Final_Answer"].tolist()
            if clean(x)
        ]

        if not answers:
            answers = [""]
        for answer in answers:
            table = document.add_table(rows=1, cols=1)
            table.alignment = WD_TABLE_ALIGNMENT.CENTER
            table.rows[0].height = Inches(0.72)
            cell = table.cell(0, 0)
            set_cell_shading(cell, "#E3DED9")
            set_cell_border(cell, "595959", "8")
            p_a = clear_cell(cell)
            p_a.paragraph_format.space_after = Pt(0)
            r_a = p_a.add_run(answer)
            r_a.font.name = "Candara"
            r_a._element.rPr.rFonts.set(qn("w:eastAsia"), "Candara")
            r_a.font.size = Pt(10.5)
            r_a.font.color.rgb = RGBColor(0, 0, 0)
            document.add_paragraph().paragraph_format.space_after = Pt(1)


def build_word_docx(df_group, participant_type, title_label, hide_names):
    document = Document()

    section = document.sections[0]
    section.page_width = Inches(8.27)
    section.page_height = Inches(11.69)
    section.top_margin = Inches(0.32)
    section.bottom_margin = Inches(0.32)
    section.left_margin = Inches(0.35)
    section.right_margin = Inches(0.35)

    footer = section.footer

    footer_table = footer.add_table(rows=1, cols=3, width=Inches(7.57))
    footer_table.alignment = WD_TABLE_ALIGNMENT.CENTER
    footer_table.autofit = False
    
    footer_table.columns[0].width = Inches(2.4)
    footer_table.columns[1].width = Inches(2.7)
    footer_table.columns[2].width = Inches(2.4)
    
    date_text = get_focus_group_date(participant_type)
    
    footer_table.cell(0, 0).text = date_text
    footer_table.cell(0, 1).text = f"Focus groupe - {participant_type}"
    
    p_page = footer_table.cell(0, 2).paragraphs[0]
    p_page.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    p_page.add_run("PAGE/NUMPAGES")
    
    for row in footer_table.rows:
        for cell in row.cells:
            set_cell_shading(cell, "66E6E6")
            set_cell_border(cell, "66E6E6", "0")
            for p in cell.paragraphs:
                for r in p.runs:
                    r.font.name = "Candara"
                    r._element.rPr.rFonts.set(qn("w:eastAsia"), "Candara")
                    r.font.size = Pt(9)
                    r.font.color.rgb = RGBColor(0, 31, 91)

    styles = document.styles
    styles["Normal"].font.name = "Candara"
    styles["Normal"]._element.rPr.rFonts.set(qn("w:eastAsia"), "Candara")
    styles["Normal"].font.size = Pt(10.5)

    names = sorted(set(clean(x) for x in df_group["participants"].dropna() if clean(x)))

    forces = get_answers(df_group, section_contains="forces", category_contains="force")
    faiblesses = get_answers(df_group, section_contains="forces", category_contains="faiblesse")
    opportunites = get_answers(df_group, section_contains="opportunites", category_contains="opportun")
    menaces = get_answers(df_group, section_contains="opportunites", category_contains="menace")
    priorites = get_answers(df_group, section_contains="prior")

    add_word_header(document, participant_type)

    title = document.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    title.paragraph_format.space_before = Pt(4)
    title.paragraph_format.space_after = Pt(8)
    run = title.add_run(title_label)
    run.bold = True
    run.font.name = "Candara"
    run._element.rPr.rFonts.set(qn("w:eastAsia"), "Candara")
    run.font.size = Pt(18)
    run.font.color.rgb = RGBColor(139, 21, 56)

    if names and not hide_names:
        p_names = document.add_paragraph()
        p_names.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p_names.paragraph_format.space_after = Pt(10)
        r_names = p_names.add_run("Nom des participants : " + "; ".join(names))
        r_names.bold = True
        r_names.font.name = "Candara"
        r_names._element.rPr.rFonts.set(qn("w:eastAsia"), "Candara")
        r_names.font.size = Pt(12)
        r_names.font.color.rgb = RGBColor(0, 0, 0)

    add_section_header(document, "I - Forces et faiblesses")
    add_two_column_answers(document, "Forces", forces, "Faiblesses", faiblesses)

    document.add_page_break()
    add_section_header(document, "II - Opportunités et menaces")
    add_two_column_answers(document, "Opportunités", opportunites, "Menaces", menaces)
    
    document.add_page_break()
    add_swot_matrix_word(document, forces, faiblesses, opportunites, menaces)

    document.add_page_break()
    add_section_header(document, "III - Priorités")
    if priorites:
        for i, answer in enumerate(priorites, start=1):
            add_single_answer_box(document, i, answer, font_size=10.5)
    else:
        add_single_answer_box(document, 1, "Aucune réponse disponible.", font_size=10.5)

    document.add_page_break()
    add_section_header(document, "IV - Conclusion")
    add_conclusion_word(document, df_group)

    buffer = BytesIO()
    document.save(buffer)
    buffer.seek(0)
    return buffer.getvalue()

def main():
    st.set_page_config(
        page_title=APP_TITLE,
        layout="wide",
        initial_sidebar_state="collapsed"
    )

    html_block(APP_CSS)

    html_block(f"""
    <div class="screen-only">
        <h1 style="margin-bottom:0;">{APP_TITLE}</h1>
        <h3 style="color:{USJ_RED} !important; margin-top:4px;">
            Lecture Excel - Print Preview des réponses corrigées
        </h3>
    </div>
    """)

    uploaded_file = st.file_uploader(
        "Uploader le fichier Excel corrigé",
        type=["xlsx", "xls"]
    )

    if uploaded_file is None:
        st.info(f"Veuillez uploader le fichier Excel contenant la feuille : {SHEET_NAME}")
        st.stop()

    try:
        df = pd.read_excel(uploaded_file, sheet_name=SHEET_NAME, engine="openpyxl")
    except Exception as e:
        st.error(f"Impossible de lire la feuille : {SHEET_NAME}")
        st.exception(e)
        st.stop()

    df = df.dropna(how="all")

    type_col = get_col(df, ["Respondent_Type", "Type participant", "Type de participants"])
    group_col = get_col(df, ["groupe", "Sous groupe", "Subgroup"])
    names_col = get_col(df, ["participants", "Participants", "Nom participants"])
    section_col = get_col(df, ["section"])
    category_col = get_col(df, ["category", "catégorie", "categorie"])
    answer_col = get_col(df, ["Final_Answer", "Final Answer", "Réponse finale", "Reponse finale"])

    missing = []
    if not type_col:
        missing.append("Respondent_Type")
    if not group_col:
        missing.append("groupe")
    if not names_col:
        missing.append("participants")
    if not section_col:
        missing.append("section")
    if not category_col:
        missing.append("category")
    if not answer_col:
        missing.append("Final_Answer")

    if missing:
        st.error("Colonnes manquantes : " + ", ".join(missing))
        st.stop()

    df = df.rename(columns={
        type_col: "Respondent_Type",
        group_col: "groupe",
        names_col: "participants",
        section_col: "section",
        category_col: "category",
        answer_col: "Final_Answer"
    })

    df["question"] = df["category"]

    for col in ["Respondent_Type", "groupe", "participants", "section", "category", "question", "Final_Answer"]:
        df[col] = df[col].apply(clean)

    df = df[(df["Respondent_Type"] != "") & (df["groupe"] != "")]
    df["section_norm"] = df["section"].apply(normalize)
    df["question_norm"] = df["question"].apply(normalize)
    df["category_norm"] = df["category"].apply(normalize)

    participant_types = sorted(df["Respondent_Type"].dropna().unique().tolist())

    c1, c2, c3, c4 = st.columns([1.2, 1.3, 1.2, 1])

    with c1:
        selected_type = st.selectbox("Type de participants", participant_types)

    df_type = df[df["Respondent_Type"] == selected_type].copy()

    with c2:
        report_mode = st.selectbox(
            "Mode d'affichage",
            [
                "Tous les sous-groupes du type sélectionné",
                "Un sous-groupe spécifique"
            ]
        )

    subgroups = sorted(df_type["groupe"].dropna().unique().tolist())

    with c3:
        if report_mode == "Un sous-groupe spécifique":
            selected_subgroup = st.selectbox("Sous-groupe", subgroups)
        else:
            selected_subgroup = "Tous les sous-groupes"

    with c4:
        hide_names = st.checkbox("Hide participants names", value=False)

    if report_mode == "Un sous-groupe spécifique":
        df_report = df_type[df_type["groupe"] == selected_subgroup].copy()
        title_label = selected_subgroup
    else:
        df_report = df_type.copy()
        title_label = "Tous les sous-groupes"

    html_report = build_one_report_html(
        df_group=df_report,
        participant_type=selected_type,
        title_label=title_label,
        hide_names=hide_names
    )

    file_base = (
        f"USJ2032_{safe_filename(selected_type)}_"
        f"{safe_filename(title_label)}_"
        f"{datetime.now().strftime('%Y%m%d_%H%M')}"
    )

    docx_bytes = build_word_docx(
        df_group=df_report,
        participant_type=selected_type,
        title_label=title_label,
        hide_names=hide_names
    )
    docx_base64 = base64.b64encode(docx_bytes).decode("utf-8")

    full_html = build_full_html(
        html_report=html_report,
        selected_type=selected_type,
        selected_label=title_label,
        docx_base64=docx_base64,
        docx_filename=f"{file_base}.docx"
    )

    components.html(
        full_html,
        height=1300,
        scrolling=True
    )


if __name__ == "__main__":
    main()
