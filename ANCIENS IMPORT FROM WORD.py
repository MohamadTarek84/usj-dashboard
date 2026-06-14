#!/usr/bin/env python
# coding: utf-8

import sqlite3
import json
import base64
import html as html_lib
import re
import textwrap
import unicodedata
from datetime import datetime
from pathlib import Path
from io import BytesIO
from docx import Document

import pandas as pd
import streamlit as st
import streamlit.components.v1 as components
from PIL import Image, ImageDraw, ImageFont

import matplotlib
matplotlib.use("Agg")
import matplotlib.pyplot as plt
from matplotlib.patches import FancyBboxPatch, Circle


APP_TITLE = "PLAN STRATÉGIQUE USJ 2032"

if Path("/home/site").exists():
    DB_DIR = Path("/home/data")          # Azure persistent storage
else:
    DB_DIR = Path(".")                   # Streamlit Cloud / local

DB_DIR.mkdir(parents=True, exist_ok=True)

DB_PATH = DB_DIR / "focus_group_anciens.db"

LOGO_PATH = Path("LogoUAQ.png")
INTRO_IMAGE_PATH = Path("Intro_schema.png")
ANNEXE_A_PATH = Path("Annexe_A.png")
ANNEXE_B_PATH = Path("Annexe_B.png")
ANNEXE_C_PATH = Path("Annexe_C.png")
PRINT_ICON_PATH = Path("Print.png")

USJ_BLUE = "#001F5B"
USJ_BLUE_2 = "#1F3C88"
USJ_RED = "#8B1538"
USJ_GOLD = "#C9A227"
USJ_LIGHT_BLUE = "#EAF2F8"
USJ_TEXT = "#1B2A41"

AUTHORIZED_TEST_CODES = {
    # Old test codes
    "USJ-HS-2032": {"responsable": "Hadi Sawaya", "institution": "ESIB"},
    "USJ-IM-2032": {"responsable": "Irma Majdalani", "institution": "FSE"},
    "USJ-NRH-2032": {"responsable": "Nadine Riachi Haddad", "institution": "FDLT"},
    "USJ-UEH-2032": {"responsable": "Ursula El Hage", "institution": "FGM"},
    "USJ-LKG-2032": {"responsable": "Lina Koleilat Ghalayini", "institution": "FSE"},
    "USJ-TH-2032": {"responsable": "Tarek Halabi", "institution": ""},

    # Focus groups
    "USJ-PSG-FG1-2032": {"responsable": "", "institution": "Sous groupe 1"},
    "USJ-PSG-FG2-2032": {"responsable": "", "institution": "Sous groupe 2"},
    "USJ-PSG-FG3-2032": {"responsable": "", "institution": "Sous groupe 3"},
    "USJ-PSG-FG4-2032": {"responsable": "", "institution": "Sous groupe 4"},
}

def html_block(content):
    if hasattr(st, "html"):
        st.html(content)
    else:
        st.markdown(content, unsafe_allow_html=True)


def image_to_base64(image_path):
    if not image_path.exists():
        return None
    suffix = image_path.suffix.lower().replace(".", "")
    mime_type = "png" if suffix == "png" else suffix
    with open(image_path, "rb") as img_file:
        encoded = base64.b64encode(img_file.read()).decode()
    return f"data:image/{mime_type};base64,{encoded}"


def init_db():
    conn = sqlite3.connect(DB_PATH)
    cur = conn.cursor()

    

    cur.execute("""
        CREATE TABLE IF NOT EXISTS responses (
            id INTEGER PRIMARY KEY AUTOINCREMENT,
            submitted_at TEXT NOT NULL,
            respondent_name TEXT,
            respondent_unit TEXT,
            respondent_email TEXT,
            statut TEXT DEFAULT 'Brouillon',
            draft_code TEXT,
            data_json TEXT NOT NULL,
            admin_data_json TEXT
        )
    """)

    existing_cols = [row[1] for row in cur.execute("PRAGMA table_info(responses)").fetchall()]

    if "statut" not in existing_cols:
        cur.execute("ALTER TABLE responses ADD COLUMN statut TEXT DEFAULT 'Brouillon'")

    if "draft_code" not in existing_cols:
        cur.execute("ALTER TABLE responses ADD COLUMN draft_code TEXT")

    if "admin_data_json" not in existing_cols:
        cur.execute("ALTER TABLE responses ADD COLUMN admin_data_json TEXT")

    conn.commit()
    conn.close()

def save_response(metadata, data):
    institution = metadata.get("institution", "").strip()
    responsable = metadata.get("responsable", "").strip()
    email = metadata.get("email", "").strip()
    statut = metadata.get("statut", "Brouillon")
    draft_code = metadata.get("draft_code", "").strip().upper()

    if not draft_code:
        raise ValueError("Un code de reprise est obligatoire pour enregistrer la réponse.")

    data["draft_code"] = draft_code
    data["metadata"]["draft_code"] = draft_code

    conn = sqlite3.connect(DB_PATH)
    cur = conn.cursor()

    existing = cur.execute("""
        SELECT id FROM responses
        WHERE draft_code = ?
        ORDER BY id DESC
        LIMIT 1
    """, (draft_code,)).fetchone()

    if existing:
        cur.execute("""
            UPDATE responses
            SET submitted_at = ?,
                respondent_name = ?,
                respondent_unit = ?,
                respondent_email = ?,
                statut = ?,
                data_json = ?
            WHERE id = ?
        """, (
            "2026-06-12",
            responsable,
            institution,
            email,
            statut,
            json.dumps(data, ensure_ascii=False),
            existing[0],
        ))
    else:
        cur.execute("""
            INSERT INTO responses (
                submitted_at,
                respondent_name,
                respondent_unit,
                respondent_email,
                statut,
                draft_code,
                data_json
            )
            VALUES (?, ?, ?, ?, ?, ?, ?)
        """, (
            datetime.now().strftime("%Y-%m-%d %H:%M:%S"),
            responsable,
            institution,
            email,
            statut,
            draft_code,
            json.dumps(data, ensure_ascii=False),
        ))

    conn.commit()
    conn.close()

    return draft_code

def keep_session_alive():
    html_block("""
<script>
setInterval(function() {
    fetch(window.location.href, {cache: "no-store"}).catch(function(error) {});
}, 240000);
</script>
""")

def load_responses():
    conn = sqlite3.connect(DB_PATH)
    df = pd.read_sql_query("SELECT * FROM responses ORDER BY id DESC", conn)
    conn.close()
    return df


def load_existing_draft_by_code(draft_code):
    draft_code = draft_code.strip().upper()

    if not draft_code:
        return None

    conn = sqlite3.connect(DB_PATH)
    cur = conn.cursor()

    row = cur.execute("""
        SELECT data_json, statut
        FROM responses
        WHERE draft_code = ?
        ORDER BY id DESC
        LIMIT 1
    """, (draft_code,)).fetchone()

    conn.close()

    if row is None:
        return None

    data = json.loads(row[0])
    data["loaded_statut"] = row[1]
    data["draft_code"] = draft_code
    return data


def preload_draft_into_session(data):
    if not data:
        return

    st.session_state["current_draft_code"] = data.get("draft_code", "")

    metadata = data.get("metadata", {})
    st.session_state["institution"] = metadata.get("institution", "")
    st.session_state["responsable"] = metadata.get("responsable", "")

    stakeholder_rows = data.get("stakeholders", {}).get("rows", [])

    stakeholder_options = [
        "Responsables institution",
        "Enseignants cadrés",
        "Enseignants non-cadrés",
        "PSG",
        "Étudiants",
        "Anciens",
        "Employeurs",
        "Conseil d’orientation stratégique",
    ]

    row_types = []
    number_of_rows_to_load = max(8, len(stakeholder_rows))

    for i in range(1, number_of_rows_to_load + 1):
        row = stakeholder_rows[i - 1] if i <= len(stakeholder_rows) else {}

        categorie = row.get("categorie", "")

        if categorie and categorie not in stakeholder_options:
            row_types.append("autres")
            st.session_state[f"stakeholder_category_autre_{i}"] = categorie
        else:
            row_types.append("standard")
            st.session_state[f"stakeholder_category_{i}"] = categorie if categorie in stakeholder_options else None

        st.session_state[f"stakeholder_nom_{i}"] = row.get("nom", "")
        st.session_state[f"stakeholder_poste_{i}"] = row.get("poste", "")
        st.session_state[f"stakeholder_organisme_{i}"] = row.get("organisme_affiliation", "")

    st.session_state["stakeholder_row_types"] = row_types

    for theme, value in data.get("internal_analysis", {}).items():
        st.session_state[f"internal_{theme}"] = value

    for theme, value in data.get("external_analysis", {}).items():
        st.session_state[f"external_{theme}"] = value

    for section_key, rows in data.get("swot_analysis", {}).items():
        prefix = "swot_internal" if section_key == "facteurs_internes" else "swot_external"
    
        if section_key == "facteurs_internes":
            forces_count = max(
                5,
                max(
                    [i for i, row in enumerate(rows, start=1) if str(row.get("Forces", "")).strip()],
                    default=5
                )
            )
    
            faiblesses_count = max(
                5,
                max(
                    [i for i, row in enumerate(rows, start=1) if str(row.get("Faiblesses", "")).strip()],
                    default=5
                )
            )
    
            st.session_state[f"{prefix}_Forces_rows"] = forces_count
            st.session_state[f"{prefix}_Faiblesses_rows"] = faiblesses_count
    
        elif section_key == "facteurs_externes":
            opportunites_count = max(
                5,
                max(
                    [i for i, row in enumerate(rows, start=1) if str(row.get("Opportunités", "")).strip()],
                    default=5
                )
            )
    
            menaces_count = max(
                5,
                max(
                    [i for i, row in enumerate(rows, start=1) if str(row.get("Menaces", "")).strip()],
                    default=5
                )
            )
    
            st.session_state[f"{prefix}_Opportunités_rows"] = opportunites_count
            st.session_state[f"{prefix}_Menaces_rows"] = menaces_count
    
        for i, row in enumerate(rows, start=1):
            if section_key == "facteurs_internes":
                st.session_state[f"{prefix}_Forces_{i}"] = row.get("Forces", "")
                st.session_state[f"{prefix}_Faiblesses_{i}"] = row.get("Faiblesses", "")
    
            elif section_key == "facteurs_externes":
                st.session_state[f"{prefix}_Opportunités_{i}"] = row.get("Opportunités", "")
                st.session_state[f"{prefix}_Menaces_{i}"] = row.get("Menaces", "")


    priorities_data = data.get("priorities_initiatives", {})
    priority_indices = []

    if isinstance(priorities_data, dict):
        for key, value in priorities_data.items():
            m = re.match(r"^priorite_(\d+)$", str(key))
            if m:
                idx = int(m.group(1))
                priority_indices.append(idx)
                st.session_state[f"priority_only_{idx}"] = value

    st.session_state["priority_only_rows"] = max(5, max(priority_indices, default=5))

    pour_finir = data.get("pour_finir", {})

    for i in range(1, 4):
        indices = []
        if isinstance(pour_finir, dict):
            for key, value in pour_finir.items():
                m = re.match(rf"^pour_finir_{i}_(\d+)$", str(key))
                if m:
                    idx = int(m.group(1))
                    indices.append(idx)
                    st.session_state[key] = value
        st.session_state[f"pour_finir_rows_{i}"] = max(2, max(indices, default=1) + 1)

def flatten_response(row):
    base = {
        "id": row["id"],
        "submitted_at": row["submitted_at"],
        "responsable": row["respondent_name"],
        "institution": row["respondent_unit"],
        "email": row["respondent_email"],
    }

    if "statut" in row.index:
        base["statut"] = row["statut"]

    if "draft_code" in row.index:
        base["draft_code"] = row["draft_code"]

    data = json.loads(row["data_json"])

    for section, values in data.items():
        if isinstance(values, dict):
            for key, value in values.items():
                if isinstance(value, list):
                    base[f"{section}_{key}"] = json.dumps(value, ensure_ascii=False)
                elif isinstance(value, dict):
                    base[f"{section}_{key}"] = json.dumps(value, ensure_ascii=False)
                else:
                    base[f"{section}_{key}"] = value
        else:
            base[section] = values

    return base

def apply_usj_style():
    html_block(f"""
<style>

.admin-narrow-box {{
    width: 75% !important;
    margin-left: auto !important;
    margin-right: auto !important;
}}

div[class*="st-key-admin_edit_"] textarea {{
    font-family: Candara, Calibri, Arial, sans-serif !important;
    font-size: 20px !important;
    font-weight: 400 !important;
    line-height: 1.35 !important;
}}

.st-key-download_export_csv button,
.st-key-download_export_excel button {{
    background-color: #6A1B9A !important;
    border: 1px solid #6A1B9A !important;
    color: white !important;
}}

/* SWOT add buttons */
div[class*="st-key-add_swot_"] button {{
    background-color: #8B1538 !important;
    border: 1px solid #8B1538 !important;
    color: white !important;
}}

div[class*="st-key-add_swot_"] button:hover {{
    background-color: #761130 !important;
    border: 1px solid #761130 !important;
    color: white !important;
}}

div[class*="st-key-add_swot_"] button p {{
    color: white !important;
}}

.st-key-download_export_csv button,
.st-key-download_export_excel button {{
    background-color: #6A1B9A !important;
    border: 1px solid #6A1B9A !important;
    color: white !important;
}}

.st-key-download_export_csv button:hover,
.st-key-download_export_excel button:hover {{
    background-color: #4A0F73 !important;
    border: 1px solid #4A0F73 !important;
}}

#MainMenu {{
    display: none !important;
    visibility: hidden !important;
}}

.stDownloadButton button {{
    background-color: #6A1B9A !important;
    border: 1px solid #6A1B9A !important;
    color: white !important;
    height: 42px !important;
    min-height: 42px !important;
    max-height: 42px !important;
}}

.stDownloadButton button:hover {{
    background-color: #A9851F !important;
    border: 1px solid #A9851F !important;
}}

header {{
    display: none !important;
    visibility: hidden !important;
}}

div[data-testid="stSelectbox"] div[data-baseweb="select"] > div {{
    background-color: #E3DED9 !important;
    border: none !important;
    border-radius: 6px !important;
    color: #000000 !important;
}}

div[data-testid="stSelectbox"] span {{
    color: #000000 !important;
}}

[data-testid="stToolbar"] {{
    display: none !important;
    visibility: hidden !important;
}}

[data-testid="stDecoration"] {{
    display: none !important;
    visibility: hidden !important;
}}

[data-testid="stStatusWidget"] {{
    display: none !important;
    visibility: hidden !important;
}}

button[data-testid="stBaseButton-header"] {{
    display: none !important;
    visibility: hidden !important;
}}

[data-testid="stBaseButton-header"] {{
    display: none !important;
    visibility: hidden !important;
}}

html, body, [class*="css"], [class*="st-"], .stApp {{
    font-family: Candara, Calibri, Arial, sans-serif !important;
    color: {USJ_TEXT};
}}

/* Clean Streamlit file uploader: keep one readable Upload label */
div[data-testid="stFileUploader"] [data-testid="stIconMaterial"],
div[data-testid="stFileUploader"] svg {{
    display: none !important;
    visibility: hidden !important;
    width: 0 !important;
    min-width: 0 !important;
    max-width: 0 !important;
    margin: 0 !important;
    padding: 0 !important;
}}

div[data-testid="stFileUploader"] button {{
    width: 220px !important;
    min-width: 220px !important;
    max-width: 220px !important;
    height: 58px !important;
    min-height: 58px !important;
    max-height: 58px !important;
    display: flex !important;
    align-items: center !important;
    justify-content: center !important;
    padding: 0 18px !important;
    white-space: nowrap !important;
}}

div[data-testid="stFileUploader"] button::before {{
    content: "Upload" !important;
    display: block !important;
    font-family: Candara, Calibri, Arial, sans-serif !important;
    font-size: 18px !important;
    font-weight: 700 !important;
    color: inherit !important;
}}

div[data-testid="stFileUploader"] button * {{
    display: none !important;
    visibility: hidden !important;
}}


/* =====================================================
   AZURE / STREAMLIT FIX
   Hide Streamlit multipage/sidebar elements and the
   broken keyboard_double_arrow_right text in Azure.
===================================================== */
section[data-testid="stSidebar"] {{
    display: none !important;
    visibility: hidden !important;
    width: 0px !important;
    min-width: 0px !important;
    max-width: 0px !important;
}}

[data-testid="stSidebarNav"] {{
    display: none !important;
    visibility: hidden !important;
}}

[data-testid="collapsedControl"],
[data-testid="stSidebarCollapsedControl"],
button[kind="header"] {{
    display: none !important;
    visibility: hidden !important;
}}

div[data-testid="stAppViewContainer"] {{
    margin-left: 0px !important;
    padding-left: 0px !important;
}}

.main .block-container,
div[data-testid="stAppViewContainer"] .block-container {{
    padding-left: 3rem !important;
    padding-right: 3rem !important;
    max-width: 100% !important;
}}

div[data-testid="stTextInput"] label,
div[data-testid="stTextInput"] label *,
div[data-testid="stDateInput"] label,
div[data-testid="stDateInput"] label * {{
    font-weight: 700 !important;
    color: #000000 !important;
}}

div[data-testid="stTextInput"] input,
div[data-testid="stDateInput"] input {{
    color: #000000 !important;
    -webkit-text-fill-color: #000000 !important;
    opacity: 1 !important;
    background-color: #E3DED9 !important;
    border: none !important;
    border-radius: 6px !important;
    box-shadow: none !important;
    outline: none !important;
}}

div[data-testid="stTextInput"] input:focus,
div[data-testid="stDateInput"] input:focus {{
    background-color: #E3DED9 !important;
    border: none !important;
    box-shadow: none !important;
    outline: none !important;
}}

div[data-testid="stTextInput"] input:disabled,
div[data-testid="stDateInput"] input:disabled {{
    background-color: #E3DED9 !important;
    border: none !important;
    border-radius: 6px !important;
    color: #000000 !important;
    -webkit-text-fill-color: #000000 !important;
    opacity: 1 !important;
}}

div[data-testid="stTextInput"] input::placeholder,
div[data-testid="stDateInput"] input::placeholder {{
    color: #000000 !important;
    -webkit-text-fill-color: #000000 !important;
    opacity: 1 !important;
}}

.annexe-a-hover {{
    position: relative;
    color: #0000FF;
    text-decoration: underline;
    font-weight: 700;
    cursor: help;
}}

.annexe-a-popup {{
    display: none;
    position: fixed;
    z-index: 99999;
    left: 50%;
    top: 90px;
    transform: translateX(-50%);
    background: white;
    padding: 8px;
    border: 1px solid #595959;
    box-shadow: 0 4px 16px rgba(0,0,0,0.20);
    max-width: 95vw;
}}

.annexe-a-popup img {{
    width: 900px !important;
    max-width: 92vw !important;
    height: auto !important;
}}

.annexe-a-hover:hover .annexe-a-popup {{
    display: block;
}}

div[data-testid="stTextArea"] {{
    border: 0.75px solid #595959 !important;
    border-radius: 0px !important;
    background-color: #E3DED9 !important;
}}

div[data-testid="stTextArea"] > div {{
    border: none !important;
    box-shadow: none !important;
}}

div[data-testid="stTextArea"] textarea {{
    background-color: #E3DED9 !important;
    border: none !important;
    border-radius: 6px !important;
    color: #000000 !important;
    -webkit-text-fill-color: #000000 !important;
    opacity: 1 !important;
    caret-color: #000000 !important;
    box-shadow: none !important;
}}

div[data-testid="stTextArea"] textarea:focus {{
    border: 2px solid #595959 !important;
    box-shadow: none !important;
    outline: none !important;
}}

div[data-testid="stTextArea"] textarea::placeholder {{
    color: #000000 !important;
    -webkit-text-fill-color: #000000 !important;
    opacity: 1 !important;
}}

h1, h2, h3, h4, h5, h6 {{
    font-family: Candara, Calibri, Arial, sans-serif !important;
    color: {USJ_BLUE} !important;
    font-weight: 700 !important;
}}

p, div, span, label, button, input, textarea, select {{
    font-family: Candara, Calibri, Arial, sans-serif !important;
}}

.stTextArea textarea {{
    resize: vertical !important;
    overflow-y: auto !important;
}}

div[data-testid="stForm"] {{
    border: none !important;
    border-radius: 0px !important;
    padding: 0px !important;
    background-color: #FFFFFF;
    box-shadow: none !important;
}}




.admin-action-row iframe {{
    width: 100% !important;
}}

.admin-action-row button,
.admin-action-row div[data-testid="stButton"] button {{
    height: 58px !important;
    min-height: 58px !important;
    width: 100% !important;
    min-width: 100% !important;
    max-width: 100% !important;
    font-family: Candara, Calibri, Arial, sans-serif !important;
    font-size: 18px !important;
    font-weight: 800 !important;
    border-radius: 8px !important;
    padding: 10px 22px !important;
}}


/* FINAL ADMIN ACTION BUTTON ALIGNMENT */
.admin-action-row {{
    width: 100% !important;
}}

.admin-action-row iframe {{
    display: block !important;
    width: 100% !important;
    min-width: 100% !important;
    height: 58px !important;
    min-height: 58px !important;
    margin: 0 !important;
    padding: 0 !important;
}}

div[data-testid="stButton"] button[kind="secondary"] {{
    height: 42px !important;
    min-height: 42px !important;
    font-family: Candara, Calibri, Arial, sans-serif !important;
    font-size: 18px !important;
    font-weight: 800 !important;
    border-radius: 8px !important;
    padding: 10px 22px !important;
    display: flex !important;
    align-items: center !important;
    justify-content: center !important;
}}

.st-key-access_report_button button {{
    width: 190px !important;
    min-width: 190px !important;
    max-width: 190px !important;
    height: 42px !important;
    min-height: 42px !important;
    max-height: 42px !important;
    padding: 4px 12px !important;
    font-size: 18px !important;
}}

div[data-testid="stButton"] button[kind="secondary"] p {{
    font-family: Candara, Calibri, Arial, sans-serif !important;
    font-size: 18px !important;
    font-weight: 800 !important;
    color: white !important;
    margin: 0 !important;
}}


.st-key-save_admin_all_button button p {{
    color: white !important;
    font-family: Candara, Calibri, Arial, sans-serif !important;
    font-size: 18px !important;
    font-weight: 800 !important;
    line-height: 1 !important;
    margin: 0 !important;
    padding: 0 !important;
}}

/* Normal buttons */
.stButton button,
.stDownloadButton button,
div[data-testid="stFormSubmitButton"] button {{
    background-color: #0070C0 !important;
    color: white !important;
    border-radius: 8px !important;
    border: 1px solid #0070C0 !important;
    font-weight: 800 !important;
    font-size: 18px !important;
    padding: 10px 22px !important;
    white-space: nowrap !important;
}}

.stButton button p,
.stDownloadButton button p,
div[data-testid="stFormSubmitButton"] button p {{
    color: white !important;
    white-space: nowrap !important;
}}

.st-key-access_report_button button {{
    width: 200px !important;
    min-width: 200px !important;
    max-width: 200px !important;
    height: 30px !important;
    min-height: 30px !important;
    padding: 8px 14px !important;
    font-size: 18px !important;
    white-space: nowrap !important;
}}

.st-key-access_report_button button p {{
    font-size: 18px !important;
    white-space: nowrap !important;
    margin: 0 !important;
}}

/* ONLY these two stakeholder add-row buttons */
.st-key-add_stakeholder_standard button,
.st-key-add_stakeholder_autres button {{
    background-color: #001F5B !important;
    border: 1px solid #001F5B !important;
    color: white !important;
}}

.st-key-add_stakeholder_standard button:hover,
.st-key-add_stakeholder_autres button:hover {{
    background-color: #001352 !important;
    border: 1px solid #001352 !important;
    color: white !important;
}}

.st-key-add_stakeholder_standard button p,
.st-key-add_stakeholder_autres button p {{
    color: white !important;
    white-space: nowrap !important;
}}


/* Action buttons: same size for save and final buttons */
.st-key-quick_save_after_stakeholders,
.st-key-quick_save_after_internal,
.st-key-quick_save_after_external,
.st-key-quick_save_after_swot,
.st-key-quick_save_after_priorities,
.st-key-save_draft_button,
.st-key-submit_final_button {{
    display: flex !important;
    justify-content: flex-start !important;
}}

.st-key-quick_save_after_stakeholders button,
.st-key-quick_save_after_internal button,
.st-key-quick_save_after_external button,
.st-key-quick_save_after_swot button,
.st-key-quick_save_after_priorities button,
.st-key-save_draft_button button,
.st-key-submit_final_button button {{
    width: 360px !important;
    min-width: 360px !important;
    max-width: 360px !important;
    height: 58px !important;
    min-height: 58px !important;
    padding: 10px 22px !important;
    display: flex !important;
    align-items: center !important;
    justify-content: center !important;
}}

/* Final submit button only */

button[kind="primary"] {{
    background-color: #8B1538 !important;
    border: 1px solid #8B1538 !important;
    color: white !important;
    min-height: 42px !important;
    height: 42px !important;
    border-radius: 8px !important;
    font-size: 16px !important;
    font-weight: 700 !important;
    padding: 0.5rem 1rem !important;
}}


button[kind="primary"] p {{
    color: white !important;
    white-space: nowrap !important;
    font-size: 18px !important;
    font-weight: 700 !important;
}}

div[data-testid="InputInstructions"] {{
    display: none !important;
}}

[data-testid="stTextInput"] [data-testid="InputInstructions"],
[data-testid="stSelectbox"] [data-testid="InputInstructions"],
[data-testid="stTextArea"] [data-testid="InputInstructions"] {{
    display: none !important;
    visibility: hidden !important;
    height: 0px !important;
}}



hr {{
    border: none !important;
    height: 3px !important;
    background-color: #D0D6E0 !important;
    margin-top: 14px !important;
    margin-bottom: 14px !important;
}}

.final-action-line {{
    border: none !important;
    height: 3px !important;
    background-color: #D0D6E0 !important;
    margin-top: 14px !important;
    margin-bottom: 14px !important;
}}

/* Move the print icon iframe upward while keeping enough height so the image is not cut */
div[data-testid="stIFrame"] {{
    margin-top: 0px !important;
}}

/* =========================
   CLEAN PRINT / PDF MODE
   Cleaned: one print block only.
========================= */

.print-answer-text {{
    display: none;
}}

.pour-finir-print-row {{
    display: none;
}}

.admin-print-title,
.swot-print-only {{
    display: none !important;
}}

/* FINAL ADMIN BUTTON FIX - same width, same height, same font */
.admin-print-button {{
    width: 100% !important;
    height: 58px !important;
    min-height: 58px !important;
    background-color: #8B1538 !important;
    color: #ffffff !important;
    border: 1px solid #8B1538 !important;
    border-radius: 8px !important;
    padding: 10px 22px !important;
    font-family: Candara, Calibri, Arial, sans-serif !important;
    font-size: 18px !important;
    font-weight: 800 !important;
    display: flex !important;
    align-items: center !important;
    justify-content: center !important;
    white-space: nowrap !important;
    cursor: pointer !important;
}}



.st-key-save_admin_all_button button {{
    height: 58px !important;
    min-height: 58px !important;
    max-height: 58px !important;
    width: 100% !important;
    min-width: 100% !important;
    max-width: 100% !important;
    background-color: #0070C0 !important;
    border: 1px solid #0070C0 !important;
    border-radius: 8px !important;
    padding: 10px 22px !important;
    margin: 0 !important;
    display: flex !important;
    align-items: center !important;
    justify-content: center !important;
    font-family: Candara, Calibri, Arial, sans-serif !important;
    font-size: 18px !important;
    font-weight: 800 !important;
    line-height: 1 !important;
}}

.st-key-save_admin_all_button button p {{
    color: white !important;
    font-family: Candara, Calibri, Arial, sans-serif !important;
    font-size: 18px !important;
    font-weight: 800 !important;
    line-height: 1 !important;
    margin: 0 !important;
    padding: 0 !important;
}}

/* Admin print cover must be hidden on screen. It is shown only during print. */
.admin-print-cover-header {{
    display: none !important;
    visibility: hidden !important;
    height: 0 !important;
    min-height: 0 !important;
    max-height: 0 !important;
    margin: 0 !important;
    padding: 0 !important;
    overflow: hidden !important;
}}

.admin-print-cover-header {{
    display: none !important;
    visibility: hidden !important;
    height: 0 !important;
    min-height: 0 !important;
    max-height: 0 !important;
    margin: 0 !important;
    padding: 0 !important;
    overflow: hidden !important;
}}

div.st-key-download_export_csv button,
div.st-key-download_export_excel button,
div.st-key-download_export_csv div[data-testid="stDownloadButton"] button,
div.st-key-download_export_excel div[data-testid="stDownloadButton"] button {{
    background-color: #6A1B9A !important;
    border-color: #6A1B9A !important;
    color: #ffffff !important;
    font-weight: 800 !important;
    font-size: 18px !important;
}}

div.st-key-download_export_csv button p,
div.st-key-download_export_excel button p {{
    font-weight: 800 !important;
    font-size: 18px !important;
    color: #ffffff !important;
}}

div.st-key-download_export_csv button:hover,
div.st-key-download_export_excel button:hover,
div.st-key-download_export_csv div[data-testid="stDownloadButton"] button:hover,
div.st-key-download_export_excel div[data-testid="stDownloadButton"] button:hover {{
    background-color: #4A0F73 !important;
    border-color: #4A0F73 !important;
}}

@media print {{


    @page {{
        size: A4 portrait;
        margin: 8mm 9mm 8mm 9mm;
    }}

    html, body, .stApp {{
        background: white !important;
        overflow: visible !important;
        width: 100% !important;
    }}

    .block-container {{
        max-width: 190mm !important;
        width: 190mm !important;
        padding: 0 !important;
        margin: 0 auto !important;
    }}

    header,
    footer,
    #MainMenu,
    .stDeployButton,
    div[data-testid="stToolbar"],
    div[data-testid="stDecoration"],
    div[data-testid="stStatusWidget"],
    .print-button-wrapper,
    div[data-testid="stButton"],
    div[data-testid="stDownloadButton"],
    div[data-testid="stAlert"],
    div[data-testid="stSelectbox"],
    iframe,
    .admin-screen-only,
    .admin-debug-info,
    .admin-action-row,
    .admin-action-row-fix,
    .admin-action-row-fix *,
    .admin-print-button,
    .admin-save-button-wrapper,
    .swot-screen-only,
    .admin-word-import-block,
    .admin-word-import-block *,
    .st-key-admin_word_import_block,
    .st-key-admin_word_import_block *,
    div[data-testid="stFileUploader"],
    div[data-testid="stFileUploader"] *,
    div[data-testid="stCheckbox"],
    div[data-testid="stCheckbox"] *,
    .st-key-admin_word_import_submitted,
    .st-key-admin_word_import_submitted *,
    .st-key-admin_word_import_button,
    .st-key-admin_word_import_button * {{
        display: none !important;
        visibility: hidden !important;
        height: 0 !important;
        min-height: 0 !important;
        max-height: 0 !important;
        margin: 0 !important;
        padding: 0 !important;
        overflow: hidden !important;
    }}

    .usj-main-header {{
        display: flex !important;
        align-items: flex-start !important;
        justify-content: space-between !important;
        gap: 12mm !important;
        width: 100% !important;
        margin: 0 !important;
        padding: 0 0 1mm 0 !important;
        border-bottom: 1px solid #D0D6E0 !important;
        min-height: 0 !important;
        height: auto !important;
        break-after: avoid !important;
        page-break-after: avoid !important;
    }}

    .usj-main-header h1 {{
        font-size: 22px !important;
        line-height: 1.05 !important;
        margin: 0 0 1mm 0 !important;
        color: #001F5B !important;
        font-weight: 800 !important;
    }}

    .usj-main-header p,
    .usj-main-header p span {{
        font-size: 10px !important;
        line-height: 1.05 !important;
        color: #1F3C88 !important;
        font-weight: 700 !important;
        margin: 0 !important;
        padding: 0 !important;
    }}

    .usj-main-header-logo {{
        width: 36mm !important;
        min-width: 36mm !important;
        max-width: 36mm !important;
        display: flex !important;
        justify-content: flex-end !important;
        align-items: flex-start !important;
        padding: 0 !important;
        margin: 0 !important;
    }}

    .usj-main-header-logo img {{
        display: block !important;
        visibility: visible !important;
        width: 36mm !important;
        max-width: 36mm !important;
        height: auto !important;
        max-height: 15mm !important;
        object-fit: contain !important;
        opacity: 1 !important;
    }}

    .admin-print-cover-header {{
        display: block !important;
        visibility: visible !important;
        height: auto !important;
        min-height: 0 !important;
        max-height: none !important;
        margin: 0 !important;
        padding: 0 !important;
        overflow: visible !important;
        break-before: auto !important;
        page-break-before: auto !important;
        break-after: avoid !important;
        page-break-after: avoid !important;
    }}

    .admin-print-cover-header img {{
        display: none !important;
    }}

    .admin-print-cover-title {{
        display: block !important;
        visibility: visible !important;
        position: static !important;
        transform: none !important;
        width: 100% !important;
        text-align: center !important;
        color: #001F5B !important;
        font-size: 22px !important;
        line-height: 1.05 !important;
        font-weight: 800 !important;
        margin: -18mm 0 -14mm 0 !important;
        padding: 0 !important;
    }}

    .admin-print-title {{
        display: none !important;
    }}

    hr {{
        display: none !important;
        height: 0 !important;
        margin: 0 !important;
        padding: 0 !important;
    }}

    h1 {{
        font-size: 22px !important;
        line-height: 1.05 !important;
        margin-top: 0 !important;
        margin-bottom: 2mm !important;
    }}

    h2 {{
        font-size: 18px !important;
        line-height: 1.15 !important;
    }}

    h3, h4, p, div, span, label {{
        font-size: 11px !important;
        line-height: 1.25 !important;
    }}

    div[style*="border-left:7px"] {{
        break-after: avoid !important;
        page-break-after: avoid !important;
        break-inside: avoid !important;
        page-break-inside: avoid !important;
        margin-top: -6mm !important;
        margin-bottom: 2mm !important;
        padding: 7px 11px !important;
        box-shadow: none !important;
    }}

    .print-page-break,
    .admin-print-page-break,
    .admin-print-field-page-break {{
        clear: both !important;
        break-before: page !important;
        page-break-before: always !important;
        break-after: auto !important;
        page-break-after: auto !important;
        height: 0 !important;
        max-height: 0 !important;
        line-height: 0 !important;
        margin: 0 !important;
        padding: 0 !important;
        overflow: hidden !important;
    }}

    .admin-conclusion-no-page-break {{
        display: none !important;
    }}

    .print-answer-block,
    .print-answer-text,
    .print-answer-content,
    .admin-answer-row-wrapper,
    .admin-answer-row,
    div[data-testid="stHorizontalBlock"] {{
        break-inside: avoid !important;
        page-break-inside: avoid !important;
    }}

    div[data-testid="stTextArea"],
    div[data-testid="stTextArea"] > div,
    div[data-testid="stTextArea"] textarea,
    .admin-original-answer-box {{
        display: block !important;
        min-height: 20mm !important;
        height: 20mm !important;
        max-height: 20mm !important;
        margin-top: 0 !important;
        margin-bottom: 0.3mm !important;
        box-sizing: border-box !important;
        overflow: hidden !important;
        break-inside: avoid !important;
        page-break-inside: avoid !important;
    }}

    div[data-testid="stTextArea"] textarea {{
        background-color: #E3DED9 !important;
        border: 1.5px solid #595959 !important;
        color: #000000 !important;
        -webkit-text-fill-color: #000000 !important;
        font-size: 10.5px !important;
        line-height: 1.18 !important;
        padding: 6px !important;
        resize: none !important;
    }}

    .admin-answer-row-wrapper {{
        margin-top: 0 !important;
        margin-bottom: 1.5mm !important;
    }}

    .print-answer-text {{
        display: block !important;
        width: 100% !important;
        margin-top: 2px !important;
        margin-bottom: 3px !important;
    }}

    .print-answer-content {{
        display: block !important;
        width: 100% !important;
        box-sizing: border-box !important;
        border: 1px solid #595959 !important;
        background-color: #E3DED9 !important;
        color: #000000 !important;
        font-size: 10.5px !important;
        line-height: 1.2 !important;
        padding: 6px !important;
        white-space: pre-wrap !important;
        overflow: hidden !important;
    }}

    .print-answer-content-empty {{
        min-height: 34px !important;
        height: 34px !important;
    }}

    .print-answer-content-filled {{
        min-height: 40px !important;
        height: auto !important;
    }}

    .word-counter-status,
    div[data-testid="InputInstructions"],
    div[style*="min-height:24px"] {{
        display: none !important;
        height: 0 !important;
        min-height: 0 !important;
        margin: 0 !important;
        padding: 0 !important;
    }}

    .st-key-pour_finir_1_1,
    .st-key-pour_finir_1_2,
    .st-key-pour_finir_1_3,
    .st-key-pour_finir_2_1,
    .st-key-pour_finir_2_2,
    .st-key-pour_finir_2_3,
    .st-key-pour_finir_3_1,
    .st-key-pour_finir_3_2,
    .st-key-pour_finir_3_3,
    .pour-finir-screen-label {{
        display: none !important;
        height: 0 !important;
        min-height: 0 !important;
        margin: 0 !important;
        padding: 0 !important;
    }}

    .pour-finir-print-row {{
        display: block !important;
        width: 100% !important;
        break-inside: avoid !important;
        page-break-inside: avoid !important;
        margin: 8px 0 12px 0 !important;
        clear: both !important;
    }}

    .pour-finir-print-label {{
        display: block !important;
        width: 100% !important;
        color: #001F5B !important;
        font-weight: 700 !important;
        font-size: 11px !important;
        line-height: 1.25 !important;
        margin-bottom: 5px !important;
        white-space: normal !important;
    }}

    .pour-finir-print-box {{
        display: block !important;
        width: 100% !important;
        box-sizing: border-box !important;
        min-height: 28px !important;
        border: 1px solid #595959 !important;
        background-color: #E3DED9 !important;
        color: #000000 !important;
        font-size: 10.5px !important;
        line-height: 1.2 !important;
        padding: 6px !important;
        margin: 0 0 5px 0 !important;
        white-space: pre-wrap !important;
        break-inside: avoid !important;
        page-break-inside: avoid !important;
    }}

    .swot-print-only {{
        display: block !important;
        break-before: page !important;
        page-break-before: always !important;
        break-after: auto !important;
        page-break-after: auto !important;
        break-inside: avoid !important;
        page-break-inside: avoid !important;
        width: 100% !important;
        margin: 0 !important;
        padding: 0 !important;
    }}

    .swot-print-shell {{
        width: 100% !important;
        min-height: auto !important;
        max-height: none !important;
        height: auto !important;
        padding: 4mm !important;
        margin: 0 !important;
        border: 1px solid #D0D6E0 !important;
        border-radius: 8px !important;
        box-shadow: none !important;
        background: white !important;
        overflow: visible !important;
        break-inside: avoid !important;
        page-break-inside: avoid !important;
    }}

    .swot-print-title {{
        text-align: center !important;
        color: #001F5B !important;
        font-size: 18px !important;
        line-height: 1.1 !important;
        font-weight: 800 !important;
        margin: 0 0 1.5mm 0 !important;
        padding: 0 !important;
    }}

    .swot-print-group {{
        text-align: center !important;
        color: #555555 !important;
        font-size: 10px !important;
        line-height: 1.1 !important;
        font-weight: 700 !important;
        margin: 0 0 3mm 0 !important;
        padding: 0 !important;
    }}

    .swot-print-grid {{
        display: grid !important;
        grid-template-columns: 1fr 1fr !important;
        gap: 3mm !important;
        break-inside: avoid !important;
        page-break-inside: avoid !important;
    }}

    .swot-print-card {{
        min-height: 36mm !important;
        max-height: 52mm !important;
        padding: 3mm !important;
        border: 1.5px solid var(--accent) !important;
        border-radius: 7px !important;
        background: var(--bg) !important;
        overflow: hidden !important;
        break-inside: avoid !important;
        page-break-inside: avoid !important;
    }}

    .swot-print-card h3 {{
        color: var(--accent) !important;
        font-size: 13px !important;
        line-height: 1.1 !important;
        font-weight: 900 !important;
        margin: 0 0 2mm 0 !important;
        padding-bottom: 1.5mm !important;
        border-bottom: 1.5px solid var(--accent) !important;
    }}

    .swot-print-card ul {{
        margin: 0 !important;
        padding-left: 4mm !important;
    }}

    .swot-print-card li {{
        font-size: 9px !important;
        line-height: 1.15 !important;
        margin-bottom: 1mm !important;
        color: #222222 !important;
    }}

    .admin-section-iii-page-break,
    .swot-print-only + .admin-print-page-break,
    .swot-print-only + div .admin-print-page-break {{
        display: none !important;
        height: 0 !important;
        min-height: 0 !important;
        max-height: 0 !important;
        margin: 0 !important;
        padding: 0 !important;
        overflow: hidden !important;
        break-before: auto !important;
        page-break-before: auto !important;
        break-after: auto !important;
        page-break-after: auto !important;
    }}

    div[data-testid="stVerticalBlock"],
    div[data-testid="stElementContainer"],
    div[data-testid="element-container"] {{
        margin-top: 0 !important;
        padding-top: 0 !important;
    }}
}}
</style>
""")
def section_header(title):
    html_block(f"""
<div style="background-color:{USJ_LIGHT_BLUE}; padding:12px 18px; border-radius:10px; border-left:7px solid {USJ_BLUE}; box-shadow:0 2px 10px rgba(0,0,0,0.08); margin-top:0px; margin-bottom:8px;">
    <h2 style="font-size:26px; color:{USJ_BLUE}; margin:0; font-weight:700;">
        {title}
    </h2>
</div>
""")


def text_area(label, key, height=920, placeholder=None):
    return st.text_area(
        label,
        key=key,
        height=height,
        placeholder=placeholder
    )

def render_print_icon_button():
    print_icon_src = image_to_base64(PRINT_ICON_PATH)

    if print_icon_src:
        components.html(
            f"""
            <div class="print-button-wrapper" style="
                height:150px;
                display:flex;
                align-items:center;
                justify-content:center;
                overflow:visible;
                padding:0;
                margin:0;
            ">
                <button onclick="window.parent.print()" title="Imprimer / Enregistrer en PDF" style="
                    background-color:transparent;
                    border:none;
                    cursor:pointer;
                    padding:0;
                    margin:0;
                    width:130px;
                    height:130px;
                    display:flex;
                    align-items:center;
                    justify-content:center;
                ">
                    <img src="{print_icon_src}" alt="Imprimer / Enregistrer en PDF" style="
                        width:130px;
                        height:130px;
                        object-fit:contain;
                        display:block;
                    ">
                </button>
            </div>
            """,
            height=155
        )
    else:
        st.warning("Print.png non trouvé. Placez Print.png dans le même dossier que le script.")

def render_first_page_header():
    logo_src = image_to_base64(LOGO_PATH)

    if logo_src:
        logo_html = f"""
        <img src="{logo_src}" alt="USJ - Unité Assurance Qualité" style="
            width:350px;
            max-width:100%;
            height:auto;
            object-fit:contain;
            display:block;
        " 
        """
    else:
        logo_html = '<div style="color:#8B1538; font-weight:700;">LogoUAQ.png non trouvé</div>'

    html_block(f"""
<div class="usj-main-header" style="
    width:100%;
    display:flex;
    align-items:flex-start;
    justify-content:space-between;
    gap:32px;
    padding-top:0px;
    margin-bottom:26px;
">
    <div style="flex:1; padding-top:0px;">
        <h1 style="font-size:42px; margin:0 0 18px 0; color:{USJ_BLUE}; line-height:1.1;">
            PLAN STRATÉGIQUE USJ 2032
        </h1>

        <p style="font-size:18px; font-weight:700; color:{USJ_BLUE_2}; margin:0; line-height:1.4;">
            <span style="font-size:26px; font-weight:700; color:{USJ_BLUE_2};">Focus groupe - Associations des Anciens Et Comité exécutif de la Fédération des associations des anciens</span>
        </p>
    </div>

    <div class="usj-main-header-logo" style="
        width:350px;
        min-width:350px;
        display:flex;
        justify-content:flex-end;
        align-items:flex-start;
        padding-top:0px;
    ">
        {logo_html}
    </div>
</div>
""")

    st.markdown("---")

def render_fixed_introduction():
    intro_image_src = image_to_base64(INTRO_IMAGE_PATH)

    image_html = ""
    if intro_image_src:
        image_html = f"""
        <div style="text-align:center; margin:18px 0;">
            <img src="{intro_image_src}" style="width:650px; max-width:100%; height:auto; border-radius:8px;">
        </div>
        """

    html_block(f"""
    <div style="background-color:#ffffff; padding:24px 34px; border-radius:12px; border-left:none; border-top:none; border-bottom:none; box-shadow:0 2px 10px rgba(0,0,0,0.08); margin-bottom:25px;">
        <p style="text-align:justify; font-size:19px; line-height:1.55; color:{USJ_BLUE};">
        L’enseignement supérieur est aujourd’hui confronté à des transformations rapides, à des contraintes économiques croissantes et à une intensification de la concurrence, tant nationale qu’internationale. Les évolutions technologiques, les attentes accrues des étudiants et des parties prenantes, ainsi que les exigences renforcées en matière de qualité et de performance, imposent une réflexion stratégique à la fois rigoureuse et collective. Les universités sont ainsi appelées à réinterroger en profondeur leurs modèles académiques, organisationnels et opérationnels.
        </p>

        <p style="text-align:justify; font-size:19px; line-height:1.55; color:{USJ_BLUE};">
        <strong>Le Plan stratégique USJ 2032</strong> s’inscrit dans cette dynamique. Il constitue une feuille de route institutionnelle visant à traduire la mission, la vision et les valeurs de l’USJ en priorités stratégiques claires, en objectifs cohérents et en initiatives concrètes, capables de renforcer durablement son positionnement, sa résilience ainsi que son impact académique et sociétal.
        </p>

        <p style="text-align:justify; font-size:19px; line-height:1.55; color:{USJ_BLUE};">
        L’élaboration de ce plan stratégique se décline en plusieurs étapes (voir le schéma ci-dessous), dont la première est consacrée à <strong>l’analyse de données relatives à l’état actuel de l’Université</strong>. L’ensemble des acteurs de l’Université, ainsi que les parties prenantes, sont invités à y contribuer. Ce Focus groupe a pour objectif d’identifier <strong>les forces à consolider, les fragilités à traiter,
les opportunités de développement et les risques à maîtriser à l’échelle de l’Université</strong>.
        </p>

        {image_html}
    </div>
""")


    
def render_internal_intro():
    html_block(f"""
<div style="background-color:#ffffff; padding:24px 34px 10px 34px; border-radius:12px; border-left:none; border-top:none; border-bottom:none; box-shadow:0 2px 10px rgba(0,0,0,0.08); margin-bottom:6px;">
    <p style="text-align:justify; font-size:17px; line-height:1.55; color:{USJ_BLUE};">
    L’analyse interne vise à apprécier dans quelle mesure l’USJ dispose des ressources nécessaires pour soutenir sa mission et mettre en œuvre ses orientations stratégiques. Elle porte également sur l’évaluation des modes d’organisation et des pratiques de gestion qui influencent directement la performance et l’efficacité de l’Université. Cette analyse permettra d’identifier dans une étape ultérieure les forces et les faiblesses de l’Université. Elle constitue un élément central du diagnostic institutionnel et contribue à éclairer les choix stratégiques, en assurant la cohérence entre les ambitions, les moyens disponibles et les capacités opérationnelles à l’échelle de l’USJ.
    </p>

    <p style="text-align:justify; font-size:17px; line-height:1.55; color:{USJ_BLUE};">
    Dans cette première étape, vous êtes appelés donc à analyser et évaluer, d’après votre expérience, l’état actuel des volets suivants (6 au minimum) au niveau de l’Université.
    </p>

    <p style="text-align:justify; font-size:17px; line-height:1.55; color:#0070C0; font-weight:700; font-style:italic; margin-bottom:2px;">
    Nous vous remercions de bien vouloir compléter les tableaux ci-dessous en vous appuyant sur les données disponibles et sur l’avis de votre institution et de ses parties prenantes, en traitant <span style="text-decoration:underline;">au moins six</span> des thèmes proposés, dans une perspective globale à l’échelle de l’Université.
    </p>
</div>
""")


def count_words(text):
    if text is None:
        return 0
    return len(str(text).split())

def word_limited_text_area(label, key, height=300, max_words=None):
    read_only = st.session_state.get("read_only_submitted", False)
    current_value = st.session_state.get(key, "")

    value = st.text_area(
        label=label,
        value=current_value,
        key=key,
        height=height,
        placeholder="Merci de saisir votre réponse ici",
        label_visibility="collapsed",
        disabled=read_only,
        on_change=lambda: st.session_state.update({"autosave_requested": True})
    )

    printable_value = html_lib.escape(value or "")
    printable_value = printable_value.replace("\n", "<br>")
    print_answer_class = "print-answer-content-filled"

    if not printable_value.strip():
        printable_value = "&nbsp;"
        print_answer_class = "print-answer-content-empty"

    html_block(f"""
<div class="print-answer-text">
    <div class="print-answer-content {print_answer_class}">{printable_value}</div>
</div>
""")

    return value



def render_internal_analysis():
    return {}
    
def render_external_intro():
    html_block(f"""
<div style="background-color:#ffffff; padding:24px 34px 10px 34px; border-radius:12px; border-left:none; border-top:none; border-bottom:none; box-shadow:0 2px 10px rgba(0,0,0,0.08); margin-bottom:6px;">

    <p style="text-align:justify; font-size:17px; line-height:1.55; color:{USJ_BLUE};">
    L’analyse de l’environnement externe constitue une étape essentielle du processus de planification stratégique. Elle vise à situer l’USJ dans son écosystème institutionnel, académique, économique et réglementaire, afin d’identifier les facteurs externes susceptibles d’influencer ses orientations, ses performances et sa soutenabilité à moyen et à long terme.
    </p>

    <p style="text-align:justify; font-size:17px; line-height:1.55; color:{USJ_BLUE};">
    Cette analyse permettra d’identifier dans une étape ultérieure, les Opportunités et les Menaces auxquelles l’Université est confrontée, en examinant les tendances clés, les enjeux critiques, les exigences réglementaires et les normes du secteur, ainsi que les pratiques et références observées auprès d’institutions ou d’organisations comparables (benchmarking).
    </p>

    <p style="text-align:justify; font-size:17px; line-height:1.55; color:{USJ_BLUE};">
    Au-delà des opinions recueillies auprès de certaines parties prenantes, cette analyse repose également sur l’exploitation de données issues de sources externes variées. Chaque institution est ainsi invitée à s’appuyer sur des recherches documentaires, des ressources provenant d’associations professionnelles ou de rapports sectoriels, ainsi que sur des entretiens menés avec des institutions ou organisations pertinentes.
    </p>

    <p style="text-align:justify; font-size:17px; line-height:1.55; color:{USJ_BLUE};">
    Dans cette deuxième étape, vous êtes appelés donc à analyser et évaluer, d’après votre expérience, l’état actuel des dimensions suivantes au niveau de l’Université.
    </p>

    <p style="text-align:justify; font-size:17px; line-height:1.55; color:#0070C0; font-weight:700; font-style:italic; margin-bottom:2px;">
    Nous vous remercions de bien vouloir compléter les tableaux ci-dessous, dans une perspective globale à l’échelle de l’Université.
    </p>

</div>
""")


def render_external_analysis():
    external_themes = [
        "Exigences ministérielles et Environnement réglementaire",
        "Marché du travail et Associations professionnelles",
        "Institutions paires : Concurrence et Benchmarking (si des informations précises sont disponibles prière de les fournir dans un document à part si nécessaire)",
        "L’Intelligence artificielle",
        "Attractivité vis-à-vis des élèves des écoles",
        "Réputation et Image",
        "Autres Menaces ou Opportunités éventuelles",
        "Suggestions de meilleures pratiques ou de programmes innovants",
    ]

    external_analysis = {}

    for theme in external_themes:

        if theme.startswith("Institutions paires"):
            title_html = f"""
            <span style="font-weight:700;">
                &bull; Institutions paires : Concurrence et Benchmarking
            </span>
            <span style="font-weight:400; font-style:italic;">
                (si des informations précises sont disponibles prière de les fournir dans un document à part si nécessaire)
            </span>
            """
        else:
            title_html = f"""
            <span style="font-weight:700;">
                &bull; {theme}
            </span>
            """

        html_block('<div class="print-answer-block">')

        html_block(f"""
<div style="padding:2px 0px; margin-top:2px; margin-bottom:4px;">
    <p style="font-size:17px; line-height:1.25; color:{USJ_RED}; font-weight:700; margin:0;">
        {title_html}
    </p>
</div>
""")

        external_analysis[theme] = word_limited_text_area(
            label=theme,
            key=f"external_{theme}",
            height=300,
            max_words=500
        )
        html_block('</div>')

    return external_analysis

def render_swot_intro():
    html_block(f"""
<div style="background-color:#ffffff; padding:24px 34px 10px 34px; border-radius:12px; border:1px solid #E0E0E0; box-shadow:0 2px 10px rgba(0,0,0,0.08); margin-bottom:6px;">

    <p style="font-size:19px; line-height:1.45; color:{USJ_BLUE}; font-weight:700; margin-bottom:10px;">
    Thématiques à prendre en considération pour répondre aux questions 1 et 2 (au moins 6) :
    </p>

    <ul style="font-size:19px; line-height:1.45; color:{USJ_BLUE}; margin-top:0; margin-bottom:18px;">
        <li>Soutenabilité financière</li>
        <li>Gouvernance et Leadership (gestion, relation, représentation, etc.)</li>
        <li>Stratégie académique et qualité d’enseignement</li>
        <li>Recherche et Innovation</li>
        <li>Ressources documentaires et Environnement digital</li>
        <li>Succès des étudiants (recrutement, accompagnement, services de support, employabilité, etc.)</li>
        <li>Ressources humaines</li>
        <li>Stratégie et mobilité internationales</li>
        <li>Mission sociétale</li>
        <li>Espace et infrastructures</li>
        <li>Environnement de travail</li>
        <li>Diversité et inclusion</li>
        <li>Développement Durable (ODD)</li>
        <li>Autre</li>
    </ul>

    <p style="font-size:19px; line-height:1.5; color:{USJ_BLUE}; font-weight:700; margin-bottom:12px;">
    1. <span style="color:#C00000; font-style:italic;">Forces :</span>
    Quels sont les politiques, processus ou projets universitaires actuels qui constituent les éléments de réussite de l’Université ?
    </p>

    <p style="font-size:19px; line-height:1.5; color:{USJ_BLUE}; font-weight:700; margin-bottom:12px;">
    2. <span style="color:#C00000; font-style:italic;">Faiblesses :</span>
    Quels sont les politiques, processus ou projets universitaires actuels à améliorer ?
    </p>

    <p style="font-size:19px; line-height:1.5; color:#0070C0; font-style:italic; margin-bottom:0;">
    Nous vous remercions de bien vouloir compléter le tableau ci-dessous en indiquant au 
    <strong>maximum cinq forces et cinq faiblesses</strong>, avec un 
    <strong>maximum de 30 mots par ligne.</strong>
    </p>

</div>
""")


def render_swot_table(section_key, left_title, right_title):
    rows = []

    left_count_key = f"{section_key}_{left_title}_rows"
    right_count_key = f"{section_key}_{right_title}_rows"

    if left_count_key not in st.session_state:
        st.session_state[left_count_key] = 5

    if right_count_key not in st.session_state:
        st.session_state[right_count_key] = 5

    col1, col2 = st.columns(2)

    with col1:
        html_block(f"""
<div style="background:{USJ_BLUE}; color:white; padding:10px 12px; min-height:42px; display:flex; align-items:center; justify-content:center; font-weight:700; border-radius:6px;">
    {left_title}
</div>
""")

    with col2:
        html_block(f"""
<div style="background:{USJ_BLUE}; color:white; padding:10px 12px; min-height:42px; display:flex; align-items:center; justify-content:center; font-weight:700; border-radius:6px;">
    {right_title}
</div>
""")

    total_rows = max(
        st.session_state[left_count_key],
        st.session_state[right_count_key]
    )

    for i in range(1, total_rows + 1):
        col1, col2 = st.columns(2)

        left_value = ""
        right_value = ""

        with col1:
            if i <= st.session_state[left_count_key]:
                left_value = word_limited_text_area(
                    label=f"{left_title} {i}",
                    key=f"{section_key}_{left_title}_{i}",
                    height=95,
                    max_words=30
                )

        with col2:
            if i <= st.session_state[right_count_key]:
                right_value = word_limited_text_area(
                    label=f"{right_title} {i}",
                    key=f"{section_key}_{right_title}_{i}",
                    height=95,
                    max_words=30
                )

        rows.append({
            left_title: left_value,
            right_title: right_value,
        })

    return rows

def render_swot_analysis():
    swot_data = {}

    annexe_b_src = image_to_base64(ANNEXE_B_PATH)

    annexe_b_hover_html = "Annexe&nbsp;B"
    if annexe_b_src:
        annexe_b_hover_html = f'<span class="annexe-a-hover">Annexe&nbsp;B<span class="annexe-a-popup"><img src="{annexe_b_src}" style="width:900px; height:auto; text-decoration:none; border-bottom:none;"></span></span>'

    annexe_c_src = image_to_base64(ANNEXE_C_PATH)

    annexe_c_hover_html = "Annexe&nbsp;C"
    if annexe_c_src:
        annexe_c_hover_html = f'<span class="annexe-a-hover">Annexe&nbsp;C<span class="annexe-a-popup"><img src="{annexe_c_src}" style="width:900px; height:auto; text-decoration:none; border-bottom:none;"></span></span>'



    swot_data["facteurs_internes"] = render_swot_table(
        section_key="swot_internal",
        left_title="Forces",
        right_title="Faiblesses"
    )

    st.session_state["quick_save_after_section_i_clicked"] = render_quick_save_button("quick_save_after_section_i")
    
    st.markdown("<br>", unsafe_allow_html=True)
    st.divider()

    section_header("II- OPPORTUNITES ET MENACES – NIVEAU USJ")

    html_block(f"""
<div style="background-color:#ffffff; padding:24px 34px 10px 34px; border-radius:12px; border:1px solid #E0E0E0; box-shadow:0 2px 10px rgba(0,0,0,0.08); margin-bottom:6px;">
    <p style="font-size:19px; line-height:1.4; color:{USJ_BLUE}; font-weight:700; margin-bottom:10px;">
    Thématiques à prendre en considération pour répondre aux questions 3 et 4 :
    </p>

    <ul style="font-size:19px; line-height:1.45; color:{USJ_BLUE}; margin-top:0; margin-bottom:18px;">
        <li>Marché du travail et Associations professionnelles</li>
        <li>Concurrence avec les autres universités</li>
        <li>Intelligence artificielle</li>
        <li>Réputation et image</li>
        <li>Environnement politique et économique, émigration</li>
        <li>Autres</li>
    </ul>

    <p style="font-size:19px; line-height:1.5; color:{USJ_BLUE}; font-weight:700; margin-bottom:12px;">
    3. <span style="color:#C00000; font-style:italic;">Opportunités :</span>
    Quels sont les éléments, facteurs, tendances, pratiques externes qui pourraient consolider la position de l’Université ?
    </p>

    <p style="font-size:19px; line-height:1.5; color:{USJ_BLUE}; font-weight:700; margin-bottom:12px;">
    4. <span style="color:#C00000; font-style:italic;">Menaces :</span>
    Quels sont les éléments, facteurs, tendances, pratiques externes qui pourraient menacer l’évolution de l’Université ?
    </p>

    <p style="font-size:19px; line-height:1.5; color:#0070C0; font-style:italic; margin-bottom:0;">
    Nous vous remercions de bien vouloir compléter le tableau ci-dessous en indiquant au 
    <strong>maximum cinq opportunités et cinq menaces</strong>, avec un 
    <strong>maximum de 30 mots par ligne.</strong>
    </p>

</div>
""")

    swot_data["facteurs_externes"] = render_swot_table(
        section_key="swot_external",
        left_title="Opportunités",
        right_title="Menaces"
    )

    st.session_state["quick_save_after_section_ii_clicked"] = render_quick_save_button("quick_save_after_section_ii")

    return swot_data
    
def render_priorities_intro():
    
    html_block(f"""
<div style="background-color:#ffffff; padding:24px 34px 10px 34px; border-radius:12px; border:1px solid #E0E0E0; box-shadow:0 2px 10px rgba(0,0,0,0.08); margin-bottom:6px;">

    <p style="font-size:19px; line-height:1.5; color:{USJ_BLUE}; font-weight:700; margin-bottom:16px;">
    <span style="color:#C00000; font-style:italic;">Priorités :</span>
    Suite à l’analyse précédente, quelles sont les priorités qu’il faudrait intégrer au prochain plan stratégique ?
    </p>

    <p style="font-size:19px; line-height:1.5; color:#0070C0; font-style:italic; margin-bottom:0;">
    Nous vous remercions de bien vouloir compléter le tableau ci-dessous en indiquant au 
    <strong>maximum cinq priorités</strong>, avec un 
    <strong>maximum de 30 mots par ligne.</strong>
    </p>

</div>
""")

def render_priorities_table():
    priorities_initiatives = {}

    html_block(f"""
<div style="margin-top:10px; margin-bottom:0;">
    <div style="background-color:{USJ_BLUE}; color:white; padding:10px 14px; font-size:18px; font-weight:700; border:1px solid #595959;">
        PRIORITÉS – NIVEAU USJ
    </div>
</div>
""")

    rows_key = "priority_only_rows"
    if rows_key not in st.session_state:
        existing_indices = []
        for key, value in st.session_state.items():
            m = re.match(r"^priority_only_(\d+)$", str(key))
            if m and str(value or "").strip():
                existing_indices.append(int(m.group(1)))
        st.session_state[rows_key] = max(5, max(existing_indices, default=5))

    for i in range(1, st.session_state[rows_key] + 1):
        priorities_initiatives[f"priorite_{i}"] = word_limited_text_area(
            label=f"Priorité {i}",
            key=f"priority_only_{i}",
            height=75,
            max_words=None
        )

    if not st.session_state.get("read_only_submitted", False):
        if st.button("+", key="add_priority_only_row"):
            st.session_state[rows_key] += 1
            st.rerun()

    return priorities_initiatives

def render_pour_finir():
    pour_finir = {}
    read_only = st.session_state.get("read_only_submitted", False)

    html_block(f"""
<div style="background-color:#ffffff; padding:4px 0 2px 0; margin-bottom:2px;">
    <p style="font-size:18px; line-height:1.25; color:{USJ_RED}; font-weight:700; font-style:italic; margin-bottom:2px;">
    POUR FINIR. Nous vous remercions de compléter les phrases suivantes :
    </p>
</div>
""")

    phrases = [
        "Nous souhaitons que l’USJ soit reconnue pour …",
        "Nous souhaitons que les employeurs disent que les diplômés de l’USJ sont …",
        "L’USJ serait un meilleur lieu d’enseignement si …",
    ]

    for i, phrase in enumerate(phrases, start=1):
        rows_key = f"pour_finir_rows_{i}"
        if rows_key not in st.session_state:
            existing_indices = []
            for key, value in st.session_state.items():
                m = re.match(rf"^pour_finir_{i}_(\d+)$", str(key))
                if m and str(value or "").strip():
                    existing_indices.append(int(m.group(1)))
            st.session_state[rows_key] = max(2, max(existing_indices, default=1) + 1)

        col_label, col_boxes, col_empty = st.columns([260, 520, 1], gap="small")

        with col_label:
            html_block(f"""
<div class="pour-finir-screen-label" style="font-size:17px; line-height:1.35; color:{USJ_BLUE}; font-weight:700; margin-top:8px; white-space:nowrap;">
    &bull; {phrase}
</div>
""")

        with col_boxes:
            for j in range(st.session_state[rows_key]):
                key = f"pour_finir_{i}_{j}"
                current_value = st.session_state.get(key, "")
                pour_finir[key] = st.text_input(
                    label=key,
                    value=current_value,
                    key=key,
                    label_visibility="collapsed",
                    disabled=read_only,
                    on_change=lambda: st.session_state.update({"autosave_requested": True})
                )

            if not read_only:
                if st.button("+", key=f"add_pour_finir_{i}"):
                    st.session_state[rows_key] += 1
                    st.rerun()

    return pour_finir




def clean_word_text(value):
    if value is None:
        return ""
    return str(value).replace("\xa0", " ").strip()


def clean_word_cell(cell):
    if cell is None:
        return ""
    parts = []
    for p in cell.paragraphs:
        txt = clean_word_text(p.text)
        if txt:
            parts.append(txt)
    return "\n".join(parts).strip()


def slugify_for_code(value):
    value = unicodedata.normalize("NFKD", str(value or ""))
    value = "".join(ch for ch in value if not unicodedata.combining(ch))
    value = re.sub(r"[^A-Za-z0-9]+", "_", value).strip("_").upper()
    return value or "FOCUS_GROUP_WORD"


def parse_import_filename(filename):
    stem = Path(filename or "").stem.strip()
    stem = re.sub(r"\s+", " ", stem).strip()

    separators = [" - ", " – ", " — ", "__"]
    institution = ""
    responsable = ""

    for sep in separators:
        if sep in stem:
            parts = [p.strip() for p in stem.split(sep) if p.strip()]
            if len(parts) >= 2:
                institution = parts[0]
                responsable = " ".join(parts[1:])
                break

    if not institution:
        institution = stem

    return {
        "institution": institution,
        "responsable": responsable,
        "draft_code": "WORD_" + slugify_for_code(stem),
        "filename_stem": stem,
    }


def parse_focus_group_word_template(uploaded_docx):
    doc = Document(uploaded_docx)
    tables = doc.tables

    if len(tables) < 4:
        raise ValueError("Le fichier Word importé ne correspond pas au modèle Focus groupe attendu.")

    swot_internal = []
    for row in tables[0].rows[1:]:
        cells = row.cells
        force = clean_word_cell(cells[0]) if len(cells) > 0 else ""
        faiblesse = clean_word_cell(cells[1]) if len(cells) > 1 else ""
        if force.strip() or faiblesse.strip():
            swot_internal.append({"Forces": force, "Faiblesses": faiblesse})

    swot_external = []
    for row in tables[1].rows[1:]:
        cells = row.cells
        opportunite = clean_word_cell(cells[0]) if len(cells) > 0 else ""
        menace = clean_word_cell(cells[1]) if len(cells) > 1 else ""
        if opportunite.strip() or menace.strip():
            swot_external.append({"Opportunités": opportunite, "Menaces": menace})

    priorities_initiatives = {}
    p_index = 1
    for row in tables[2].rows[1:]:
        value = clean_word_cell(row.cells[0]) if row.cells else ""
        if value.strip():
            priorities_initiatives[f"priorite_{p_index}"] = value
            p_index += 1

    phrase_map = {
        "reconnue": 1,
        "employeurs": 2,
        "diplômés": 2,
        "diplomes": 2,
        "enseignement": 3,
    }
    counters = {1: 0, 2: 0, 3: 0}
    pour_finir = {}

    for row in tables[3].rows:
        cells = row.cells
        label = clean_word_cell(cells[0]) if len(cells) > 0 else ""
        value = clean_word_cell(cells[1]) if len(cells) > 1 else ""
        if not value.strip():
            continue

        label_norm = slugify_for_code(label).lower()
        phrase_id = None
        for fragment, idx in phrase_map.items():
            if slugify_for_code(fragment).lower() in label_norm:
                phrase_id = idx
                break
        if phrase_id is None:
            phrase_id = 1

        j = counters[phrase_id]
        pour_finir[f"pour_finir_{phrase_id}_{j}"] = value
        counters[phrase_id] += 1

    data = {
        "metadata": {},
        "introduction": {},
        "stakeholders": {"rows": []},
        "internal_analysis": {},
        "external_analysis": {},
        "swot_analysis": {
            "facteurs_internes": swot_internal,
            "facteurs_externes": swot_external,
        },
        "priorities_initiatives": priorities_initiatives,
        "pour_finir": pour_finir,
    }

    return data


def render_admin_word_importer():
    st.markdown("### Importer les réponses depuis Word")

    uploaded_docx = st.file_uploader(
        "Importer le fichier Word complété",
        type=["docx"],
        key="admin_word_import_docx"
    )

    detected = {"institution": "", "responsable": "", "draft_code": "", "filename_stem": ""}

    if uploaded_docx is not None:
        detected = parse_import_filename(uploaded_docx.name)
        col_inst, col_resp, col_file = st.columns(3)
        with col_inst:
            st.text_input("Focus groupe détecté depuis le nom du fichier", value=detected["institution"], disabled=True, key="admin_word_detected_institution")
        with col_resp:
            st.text_input("Participants détectés depuis le nom du fichier", value=detected["responsable"], disabled=True, key="admin_word_detected_responsable")
        with col_file:
            st.text_input("Nom du fichier", value=uploaded_docx.name, disabled=True, key="admin_word_detected_filename")
        st.caption("Format conseillé du nom du fichier : Focus groupe - Participants.docx")

    import_as_submitted = st.checkbox("Marquer comme Soumis", value=True, key="admin_word_import_submitted")

    if st.button("Importer ce fichier Word", key="admin_word_import_button"):
        if uploaded_docx is None:
            st.error("Veuillez d'abord sélectionner un fichier Word.")
            st.stop()

        try:
            detected = parse_import_filename(uploaded_docx.name)
            imported_data = parse_focus_group_word_template(uploaded_docx)
            statut = "Soumis" if import_as_submitted else "Brouillon"
            metadata = {
                "institution": detected["institution"],
                "responsable": detected["responsable"],
                "email": "",
                "response_date": datetime.now().strftime("%Y-%m-%d"),
                "statut": statut,
                "draft_code": detected["draft_code"],
                "import_source": "Word",
                "imported_filename": uploaded_docx.name,
                "imported_at": datetime.now().strftime("%Y-%m-%d %H:%M:%S"),
            }
            imported_data["metadata"] = metadata

            # IMPORTANT:
            # The Word file is saved as the ORIGINAL version in data_json.
            # Moderator/admin edits are saved later in admin_data_json only.
            # Therefore, the imported Word answers are preserved and are never overwritten
            # by the moderator modifications.
            save_response(metadata, imported_data)

            st.session_state["view_selected_form"] = True
            st.session_state["view_selected_draft_code"] = detected["draft_code"]
            st.session_state["current_draft_code"] = detected["draft_code"]
            st.session_state["read_only_submitted"] = True
            st.success(
                "Le fichier Word a été importé avec succès. "
                "La version Word originale est conservée séparément; "
                "les modifications admin seront enregistrées dans une version distincte."
            )
            st.rerun()
        except Exception as e:
            st.error(f"Import impossible : {e}")

def render_quick_save_button(key):
    if st.session_state.get("read_only_submitted", False):
        return False

    clicked = st.button(
        "Enregistrer et continuer plus tard",
        key=key
    )

    if clicked:
        st.session_state["last_quick_save_key"] = key

    if st.session_state.get("quick_save_success_key") == key:
        st.success(
            f"Vos réponses ont été enregistrées. Votre mot de passe pour reprendre plus tard : "
            f"{st.session_state.get('current_draft_code', '')}"
        )

    return clicked


def find_word_limit_errors(section_data, section_label, max_words):
    return []


def unlock_response_by_code(draft_code):
    conn = sqlite3.connect(DB_PATH)
    cur = conn.cursor()

    cur.execute("""
        UPDATE responses
        SET statut = 'Brouillon'
        WHERE draft_code = ?
    """, (draft_code.strip().upper(),))

    conn.commit()
    conn.close()


def delete_response_by_code(draft_code):
    conn = sqlite3.connect(DB_PATH)
    cur = conn.cursor()

    cur.execute("""
        DELETE FROM responses
        WHERE draft_code = ?
    """, (draft_code.strip().upper(),))

    conn.commit()
    conn.close()

def save_admin_version_by_code(draft_code, admin_data):
    conn = sqlite3.connect(DB_PATH)
    cur = conn.cursor()

    cur.execute(
        """
        UPDATE responses
        SET admin_data_json = ?
        WHERE draft_code = ?
        """,
        (
            json.dumps(admin_data, ensure_ascii=False),
            draft_code
        )
    )

    conn.commit()
    conn.close()



def safe_filename(value):
    value = str(value or "").strip()
    value = re.sub(r"[^A-Za-z0-9À-ÿ._-]+", "_", value)
    value = value.strip("_")
    return value or "groupe"


def _normalize_text_for_match(value):
    value = str(value or "").strip().lower()
    replacements = {
        "é": "e", "è": "e", "ê": "e", "ë": "e",
        "à": "a", "â": "a",
        "ù": "u", "û": "u",
        "î": "i", "ï": "i",
        "ô": "o",
        "ç": "c",
    }
    for old, new in replacements.items():
        value = value.replace(old, new)
    return value


def _clean_swot_answer(value):
    value = str(value or "").strip()
    value = re.sub(r"\s+", " ", value)
    value = value.strip("•*-–— ")
    return value


def _get_row_value_exact_or_safe(row, wanted_field):
    """
    Reads only the expected admin-edited answer field.
    This prevents the old long question labels from being placed under the wrong SWOT box.
    """
    if not isinstance(row, dict):
        return ""

    if wanted_field in row:
        return _clean_swot_answer(row.get(wanted_field, ""))

    wanted_norm = _normalize_text_for_match(wanted_field)

    # Safe fallback only for the same field family.
    for key, value in row.items():
        key_norm = _normalize_text_for_match(key)

        if wanted_norm == "forces" and key_norm in ["force", "forces"]:
            return _clean_swot_answer(value)

        if wanted_norm == "faiblesses" and key_norm in ["faiblesse", "faiblesses"]:
            return _clean_swot_answer(value)

        if wanted_norm == "opportunites" and ("opportunite" in key_norm or "opportunites" in key_norm):
            return _clean_swot_answer(value)

        if wanted_norm == "menaces" and ("menace" in key_norm or "menaces" in key_norm):
            return _clean_swot_answer(value)

    return ""


def _deduplicate_keep_order(values):
    cleaned = []
    seen = set()

    for value in values:
        value = _clean_swot_answer(value)

        if not value:
            continue

        # Do not inject old instruction/question text into the image.
        value_norm = _normalize_text_for_match(value)
        blocked_fragments = [
            "quels sont les elements",
            "maximum cinq",
            "merci de bien vouloir",
            "thematique",
            "opportunites : quels",
            "menaces : quels",
            "forces : quels",
            "faiblesses : quels",
        ]

        if any(fragment in value_norm for fragment in blocked_fragments):
            continue

        key = value_norm
        if key not in seen:
            cleaned.append(value)
            seen.add(key)

    return cleaned


def extract_admin_swot_values(admin_data):
    """
    Reads the live admin-edited answers only:
    - I - Forces et faiblesses: Forces, Faiblesses
    - II - Opportunités et menaces: Opportunités, Menaces
    """
    if not isinstance(admin_data, dict):
        admin_data = {}

    section_internal = admin_data.get("I - Forces et faiblesses", [])
    section_external = admin_data.get("II - Opportunités et menaces", [])

    def collect(section, field):
        values = []
        if isinstance(section, list):
            for row in section:
                values.append(_get_row_value_exact_or_safe(row, field))
        elif isinstance(section, dict):
            values.append(_get_row_value_exact_or_safe(section, field))
        return _deduplicate_keep_order(values)

    return {
        "forces": collect(section_internal, "Forces"),
        "faiblesses": collect(section_internal, "Faiblesses"),
        "opportunites": collect(section_external, "Opportunités"),
        "menaces": collect(section_external, "Menaces"),
    }


def _wrap_lines_for_box(text, max_chars=42, max_lines=3):
    text = _clean_swot_answer(text)
    wrapped = textwrap.wrap(
        text,
        width=max_chars,
        break_long_words=False,
        break_on_hyphens=False
    )

    if len(wrapped) > max_lines:
        wrapped = wrapped[:max_lines]
        wrapped[-1] = wrapped[-1].rstrip(" .,;:") + "…"

    return wrapped


def _draw_swot_box(ax, x, y, w, h, title, items, facecolor, edgecolor):
    box = FancyBboxPatch(
        (x, y),
        w,
        h,
        boxstyle="round,pad=0.035,rounding_size=0.18",
        linewidth=2.2,
        edgecolor=edgecolor,
        facecolor=facecolor,
        alpha=0.98,
        zorder=1
    )
    ax.add_patch(box)

    ax.text(
        x + 0.32,
        y + h - 0.35,
        title,
        fontsize=27,
        fontweight="bold",
        color=edgecolor,
        ha="left",
        va="center",
        zorder=2
    )

    ax.plot(
        [x + 0.32, x + w - 0.32],
        [y + h - 0.72, y + h - 0.72],
        color=edgecolor,
        linewidth=2.0,
        zorder=2
    )

    text_y = y + h - 1.05
    min_y = y + 0.34

    for item in _deduplicate_keep_order(items):
        lines = _wrap_lines_for_box(item, max_chars=45, max_lines=3)
        needed_height = 0.29 * len(lines) + 0.22

        if text_y - needed_height < min_y:
            break

        ax.text(
            x + 0.42,
            text_y,
            "•",
            fontsize=23,
            fontweight="bold",
            color=edgecolor,
            ha="left",
            va="top",
            zorder=2
        )

        line_y = text_y
        for line in lines:
            ax.text(
                x + 0.68,
                line_y,
                line,
                fontsize=17,
                color="#2B2B2B",
                ha="left",
                va="top",
                zorder=2
            )
            line_y -= 0.29

        text_y = line_y - 0.13


def create_swot_image_bytes(swot_values, group_name="groupe"):
    """
    Creates a clean SWOT image from the admin-modified answers.
    The title and group name are inside the image.
    """
    fig, ax = plt.subplots(figsize=(13.2, 8.0), dpi=180)
    ax.set_xlim(0, 10)
    ax.set_ylim(0, 7.25)
    ax.axis("off")

    ax.text(
        5,
        7.05,
        "Matrice SWOT - Niveau USJ",
        fontsize=22,
        fontweight="bold",
        color=USJ_BLUE,
        ha="center",
        va="center"
    )

    ax.text(
        5,
        6.78,
        f"Groupe : {group_name}",
        fontsize=12,
        color="#4F4F4F",
        ha="center",
        va="center"
    )

    boxes = [
        (0.35, 3.75, 4.25, 2.70, "FORCES", swot_values.get("forces", []), "#DDEFF7", USJ_BLUE),
        (5.40, 3.75, 4.25, 2.70, "FAIBLESSES", swot_values.get("faiblesses", []), "#FBE3C3", USJ_RED),
        (0.35, 0.45, 4.25, 2.70, "OPPORTUNITÉS", swot_values.get("opportunites", []), "#E2F2D3", "#2F6B2F"),
        (5.40, 0.45, 4.25, 2.70, "MENACES", swot_values.get("menaces", []), "#F4C6C4", USJ_RED),
    ]

    for box_args in boxes:
        _draw_swot_box(ax, *box_args)

    buffer = BytesIO()
    fig.savefig(buffer, format="png", dpi=180, bbox_inches="tight", facecolor="white", pad_inches=0.08)
    plt.close(fig)
    buffer.seek(0)
    return buffer.getvalue()



def _split_admin_items_for_display(items):
    cleaned_items = _deduplicate_keep_order(items)

    display_items = []
    for item in cleaned_items:
        item = _clean_swot_answer(item)
        if item:
            display_items.append(item)

    return display_items


def _swot_items_html(items, accent_color):
    display_items = _split_admin_items_for_display(items)

    if not display_items:
        return f"""
        <div class="swot-empty">
            Aucune réponse admin saisie pour cette rubrique.
        </div>
        """

    html_items = []

    for index, item in enumerate(display_items, start=1):
        safe_item = html_lib.escape(item)
        html_items.append(f"""
        <div class="swot-item" title="{safe_item}">
            <div class="swot-item-index" style="background:{accent_color};">{index}</div>
            <div class="swot-item-text">{safe_item}</div>
        </div>
        """)

    return "\n".join(html_items)


def _swot_card_html(title, subtitle, items, accent_color, background, icon):
    count = len(_split_admin_items_for_display(items))
    count_label = f"{count} élément" if count == 1 else f"{count} éléments"

    return f"""
    <div class="swot-card" style="--accent:{accent_color}; --bg:{background};">
        <div class="swot-card-top">
            <div class="swot-icon" style="background:{accent_color};">{icon}</div>
            <div>
                <div class="swot-title">{title}</div>
                <div class="swot-subtitle">{subtitle}</div>
            </div>
            <div class="swot-count">{count_label}</div>
        </div>

        <div class="swot-line"></div>

        <div class="swot-items">
            {_swot_items_html(items, accent_color)}
        </div>
    </div>
    """


def create_interactive_swot_html(swot_values, group_name="groupe"):
    """
    Creates an interactive, dynamic, professional SWOT matrix as HTML.
    It reads the admin-edited answers already extracted in swot_values.
    The cards expand/animate on hover, items show full text on hover, and a small
    JavaScript control allows the admin to focus on one quadrant.
    """
    safe_group_name = html_lib.escape(str(group_name or "groupe"))

    forces = swot_values.get("forces", [])
    faiblesses = swot_values.get("faiblesses", [])
    opportunites = swot_values.get("opportunites", [])
    menaces = swot_values.get("menaces", [])

    total_items = (
        len(_split_admin_items_for_display(forces)) +
        len(_split_admin_items_for_display(faiblesses)) +
        len(_split_admin_items_for_display(opportunites)) +
        len(_split_admin_items_for_display(menaces))
    )

    html = f"""
<!DOCTYPE html>
<html lang="fr">
<head>
<meta charset="UTF-8">
<style>
    * {{
        box-sizing: border-box;
    }}

    body {{
        margin: 0;
        padding: 0;
        font-family: Candara, Calibri, Arial, sans-serif;
        background: #ffffff;
        color: #1B2A41;
    }}

    .swot-shell {{
        width: 100%;
        min-height: auto;
        height: auto;
        padding: 28px 30px 36px 30px;
        background:
            radial-gradient(circle at 50% 47%, rgba(201,162,39,0.10) 0, rgba(201,162,39,0.00) 22%),
            linear-gradient(135deg, #FFFFFF 0%, #F7FAFD 48%, #FFFFFF 100%);
        border: 1px solid #E2E8F0;
        border-radius: 22px;
        box-shadow: 0 16px 38px rgba(0,31,91,0.12);
        overflow: visible !important;
        position: relative;
    }}

    .swot-header {{
        display: flex;
        justify-content: space-between;
        align-items: flex-start;
        gap: 18px;
        margin-bottom: 20px;
    }}

    .swot-heading {{
        flex: 1;
    }}

    .swot-heading h1 {{
        margin: 0;
        color: #001F5B;
        font-size: 34px;
        line-height: 1.05;
        letter-spacing: 0.2px;
        font-weight: 800;
    }}

    .swot-heading p {{
        margin: 7px 0 0 0;
        color: #4F5B6B;
        font-size: 17px;
        font-weight: 600;
    }}

    .swot-badge {{
        background: #001F5B;
        color: white;
        border-radius: 999px;
        padding: 10px 18px;
        font-size: 15px;
        font-weight: 800;
        box-shadow: 0 8px 18px rgba(0,31,91,0.18);
        white-space: nowrap;
    }}

    .swot-toolbar {{
        display: flex;
        gap: 10px;
        flex-wrap: wrap;
        margin-bottom: 20px;
    }}

    .swot-filter {{
        border: 1px solid #D6DEE8;
        background: white;
        color: #001F5B;
        font-family: Candara, Calibri, Arial, sans-serif;
        font-weight: 800;
        border-radius: 999px;
        padding: 8px 14px;
        cursor: pointer;
        transition: all 0.22s ease;
        box-shadow: 0 3px 10px rgba(0,0,0,0.04);
    }}

    .swot-filter:hover,
    .swot-filter.active {{
        background: #001F5B;
        color: white;
        transform: translateY(-1px);
        box-shadow: 0 8px 18px rgba(0,31,91,0.18);
    }}

    .swot-grid {{
        display: grid;
        grid-template-columns: 1fr 1fr;
        gap: 26px;
        position: relative;
        align-items: stretch;
        overflow: visible !important;
    }}

    .swot-card {{
        min-height: 300px;
        height: auto !important;
        background: var(--bg);
        border: 2.5px solid var(--accent);
        border-radius: 22px;
        padding: 22px 24px 26px 24px;
        box-shadow: 0 12px 24px rgba(0,0,0,0.08);
        transition: transform 0.22s ease, box-shadow 0.22s ease, opacity 0.22s ease;
        position: relative;
        overflow: visible !important;
    }}

    .swot-card::after {{
        content: "";
        position: absolute;
        width: 210px;
        height: 210px;
        right: -80px;
        bottom: -90px;
        border-radius: 50%;
        background: color-mix(in srgb, var(--accent) 10%, transparent);
        pointer-events: none;
        z-index: 0;
    }}

    .swot-card:hover {{
        transform: translateY(-5px) scale(1.01);
        box-shadow: 0 20px 36px rgba(0,31,91,0.15);
        z-index: 2;
    }}

    .swot-card.dimmed {{
        opacity: 0.25;
        transform: scale(0.985);
    }}

    .swot-card.focused {{
        transform: translateY(-5px) scale(1.015);
        box-shadow: 0 22px 42px rgba(0,31,91,0.20);
        z-index: 3;
    }}

    .swot-card-top {{
        display: grid;
        grid-template-columns: 58px 1fr auto;
        gap: 14px;
        align-items: center;
        position: relative;
        z-index: 1;
    }}

    .swot-icon {{
        width: 58px;
        height: 58px;
        border-radius: 18px;
        display: flex;
        align-items: center;
        justify-content: center;
        color: white;
        font-size: 30px;
        font-weight: 900;
        box-shadow: 0 10px 18px rgba(0,0,0,0.16);
    }}

    .swot-title {{
        color: var(--accent);
        font-size: 31px;
        font-weight: 900;
        letter-spacing: 0.5px;
        line-height: 1;
    }}

    .swot-subtitle {{
        color: #4F5B6B;
        font-size: 14px;
        font-weight: 700;
        margin-top: 5px;
    }}

    .swot-count {{
        color: var(--accent);
        background: rgba(255,255,255,0.72);
        border: 1px solid rgba(255,255,255,0.90);
        padding: 6px 10px;
        border-radius: 999px;
        font-size: 13px;
        font-weight: 900;
        white-space: nowrap;
    }}

    .swot-line {{
        height: 3px;
        background: var(--accent);
        border-radius: 999px;
        margin: 18px 0 16px 72px;
        position: relative;
        z-index: 1;
    }}

    .swot-items {{
        display: flex;
        flex-direction: column;
        gap: 10px;
        position: relative;
        z-index: 1;
    }}

    .swot-item {{
        display: grid;
        grid-template-columns: 32px 1fr;
        gap: 10px;
        align-items: start;
        min-height: 54px;
        height: auto !important;
        background: rgba(255,255,255,0.62);
        border: 1px solid rgba(255,255,255,0.80);
        border-radius: 14px;
        padding: 10px 12px;
        transition: all 0.20s ease;
        cursor: default;
        overflow: visible !important;
        white-space: normal !important;
        word-break: break-word;
    }}

    .swot-item:hover {{
        background: white;
        transform: translateX(4px);
        box-shadow: 0 8px 18px rgba(0,0,0,0.10);
    }}

    .swot-item-index {{
        width: 26px;
        height: 26px;
        color: white;
        border-radius: 50%;
        font-size: 14px;
        font-weight: 900;
        display: flex;
        align-items: center;
        justify-content: center;
        margin-top: 1px;
        flex-shrink: 0;
    }}

    .swot-item-text {{
        color: #202A35;
        font-size: 18px;
        line-height: 1.28;
        font-weight: 650;
        overflow-wrap: anywhere;
        white-space: normal !important;
        word-break: break-word;
    }}

    .swot-empty {{
        color: #6B7280;
        background: rgba(255,255,255,0.62);
        border: 1px dashed #B7C0CC;
        border-radius: 14px;
        padding: 16px;
        font-size: 17px;
        font-weight: 700;
        font-style: italic;
    }}


    .swot-footer {{
        margin-top: 18px;
        display: flex;
        justify-content: space-between;
        gap: 12px;
        color: #5B6777;
        font-size: 13px;
        font-weight: 700;
    }}

    @media (max-width: 900px) {{
        .swot-grid {{
            grid-template-columns: 1fr;
        }}

        .swot-card-top {{
            grid-template-columns: 50px 1fr;
        }}

        .swot-count {{
            grid-column: 2;
            width: fit-content;
        }}
    }}
</style>
</head>
<body>
<div class="swot-shell">
    <div class="swot-header">
        <div class="swot-heading">
            <h1>Matrice SWOT - Niveau USJ</h1>
            <p>Groupe : {safe_group_name}</p>
        </div>
        <div class="swot-badge">{total_items} éléments validés</div>
    </div>

    <div class="swot-toolbar">
        <button class="swot-filter active" onclick="filterCards('all', this)">Vue complète</button>
        <button class="swot-filter" onclick="filterCards('forces', this)">Forces</button>
        <button class="swot-filter" onclick="filterCards('faiblesses', this)">Faiblesses</button>
        <button class="swot-filter" onclick="filterCards('opportunites', this)">Opportunités</button>
        <button class="swot-filter" onclick="filterCards('menaces', this)">Menaces</button>
    </div>

    <div class="swot-grid">
        <div data-card="forces">
            {_swot_card_html("FORCES", "Atouts internes à consolider", forces, USJ_BLUE, "#DDEFF7", "✓")}
        </div>

        <div data-card="faiblesses">
            {_swot_card_html("FAIBLESSES", "Points internes à améliorer", faiblesses, USJ_RED, "#FBE3C3", "!")}
        </div>

        <div data-card="opportunites">
            {_swot_card_html("OPPORTUNITÉS", "Leviers externes à saisir", opportunites, "#2F6B2F", "#E2F2D3", "↗")}
        </div>

        <div data-card="menaces">
            {_swot_card_html("MENACES", "Risques externes à anticiper", menaces, USJ_RED, "#F4C6C4", "⚠")}
        </div>

    </div>


    </div>

<script>
function filterCards(target, button) {{
    const wrappers = document.querySelectorAll('[data-card]');
    const buttons = document.querySelectorAll('.swot-filter');

    buttons.forEach(btn => btn.classList.remove('active'));
    button.classList.add('active');

    wrappers.forEach(wrapper => {{
        const card = wrapper.querySelector('.swot-card');
        const key = wrapper.getAttribute('data-card');

        card.classList.remove('dimmed', 'focused');

        if (target !== 'all') {{
            if (key === target) {{
                card.classList.add('focused');
            }} else {{
                card.classList.add('dimmed');
            }}
        }}
    }});
}}
</script>
</body>
</html>
"""
    return html


def estimate_interactive_swot_height(swot_values):
    """
    Estimates iframe height so the interactive SWOT matrix expands instead of cutting content.
    """
    def item_lines(item):
        text = _clean_swot_answer(item)
        if not text:
            return 1
        return max(1, len(textwrap.wrap(text, width=58, break_long_words=False, break_on_hyphens=False)))

    def card_height(items):
        display_items = _split_admin_items_for_display(items)
        if not display_items:
            return 300

        items_height = 0
        for item in display_items:
            lines = item_lines(item)
            items_height += 58 + max(0, lines - 1) * 24 + 10

        return max(300, 145 + items_height + 26)

    top_row = max(
        card_height(swot_values.get("forces", [])),
        card_height(swot_values.get("faiblesses", []))
    )
    bottom_row = max(
        card_height(swot_values.get("opportunites", [])),
        card_height(swot_values.get("menaces", []))
    )

    # Header, toolbar, shell padding, grid gap, footer, and safety buffer.
    return int(max(820, 165 + top_row + bottom_row + 26 + 65 + 70))



def _swot_print_static_html(swot_values, group_name="groupe"):
    def items_to_li(items):
        cleaned = _split_admin_items_for_display(items)
        if not cleaned:
            return "<li>Aucune réponse admin saisie.</li>"
        return "".join(f"<li>{html_lib.escape(str(item))}</li>" for item in cleaned)

    safe_group = html_lib.escape(str(group_name or "groupe"))

    return f"""
<div class="swot-print-only">
    <div class="swot-print-shell">
        <div class="swot-print-title">Matrice SWOT - Niveau USJ</div>
        <div class="swot-print-group">Groupe : {safe_group}</div>
        <div class="swot-print-grid">
            <div class="swot-print-card" style="--accent:{USJ_BLUE}; --bg:#DDEFF7;">
                <h3>FORCES</h3>
                <ul>{items_to_li(swot_values.get('forces', []))}</ul>
            </div>
            <div class="swot-print-card" style="--accent:{USJ_RED}; --bg:#FBE3C3;">
                <h3>FAIBLESSES</h3>
                <ul>{items_to_li(swot_values.get('faiblesses', []))}</ul>
            </div>
            <div class="swot-print-card" style="--accent:#2F6B2F; --bg:#E2F2D3;">
                <h3>OPPORTUNITÉS</h3>
                <ul>{items_to_li(swot_values.get('opportunites', []))}</ul>
            </div>
            <div class="swot-print-card" style="--accent:{USJ_RED}; --bg:#F4C6C4;">
                <h3>MENACES</h3>
                <ul>{items_to_li(swot_values.get('menaces', []))}</ul>
            </div>
        </div>
    </div>
</div>
"""

def render_swot_image_download_block(updated_admin_data, selected_row):
    swot_values = extract_admin_swot_values(updated_admin_data)

    group_name = selected_row.get("respondent_name", "") if hasattr(selected_row, "get") else ""
    group_unit = selected_row.get("respondent_unit", "") if hasattr(selected_row, "get") else ""
    draft_code = selected_row.get("draft_code", "") if hasattr(selected_row, "get") else ""

    display_group_name = " - ".join(
        [part for part in [str(group_name).strip(), str(group_unit).strip()] if part]
    )

    if not display_group_name:
        display_group_name = str(draft_code).strip() or "groupe"

    has_values = any(swot_values[key] for key in swot_values)

    safe_code = safe_filename(str(draft_code))
    button_key = f"generate_interactive_swot_{safe_code}"
    state_key = f"show_interactive_swot_{safe_code}"
    html_file_name = f"SWOT_interactive_{safe_filename(display_group_name)}.html"

    st.markdown("---")

    if state_key not in st.session_state:
        st.session_state[state_key] = False

    if st.button("Générer la matrice SWOT interactive", key=button_key):
        st.session_state[state_key] = True

    if not has_values:
        st.warning("Aucune réponse admin n’est disponible pour générer la matrice SWOT interactive.")
        return

    # Always include the printable SWOT matrix in the print/PDF version, even if the admin
    # did not open the interactive matrix on screen.
    html_block(_swot_print_static_html(swot_values, display_group_name))

    if not st.session_state[state_key]:
        return

    interactive_html = create_interactive_swot_html(
        swot_values=swot_values,
        group_name=display_group_name
    )

    html_block('<div class="swot-screen-only">')

    components.html(
        interactive_html,
        height=estimate_interactive_swot_height(swot_values),
        scrolling=False
    )

    html_block('</div>')

    st.download_button(
        label="Télécharger la matrice SWOT interactive en HTML",
        data=interactive_html.encode("utf-8"),
        file_name=html_file_name,
        mime="text/html",
        key=f"download_interactive_swot_html_{safe_code}"
    )

def trigger_admin_autosave():
    st.session_state["admin_autosave_requested"] = True


def offline_text_backup_safe():
    components.html("""
<script>
(function() {
    const PREFIX = "USJ_SAFE_BACKUP_";

    function getKey(el) {
        const parent = el.closest('[class*="st-key-"]');
        if (!parent) return null;
        const keyClass = Array.from(parent.classList).find(c => c.startsWith("st-key-"));
        if (!keyClass) return null;
        return PREFIX + keyClass;
    }

    function attachBackup() {
        const doc = window.parent.document;
        const fields = doc.querySelectorAll("textarea, input[type='text']");

        fields.forEach(function(el) {
            if (el.dataset.usjBackupAttached === "1") return;

            const key = getKey(el);
            if (!key) return;

            el.dataset.usjBackupAttached = "1";

            el.addEventListener("input", function() {
                localStorage.setItem(key, el.value);
            });
        });
    }

    setTimeout(attachBackup, 1500);
    setInterval(attachBackup, 5000);
})();
</script>
""", height=0)

def main():
    st.set_page_config(
        page_title=APP_TITLE,
        page_icon="📋",
        layout="wide",
        initial_sidebar_state="collapsed"
    )

    apply_usj_style()
    keep_session_alive()
    offline_text_backup_safe()
    
    if not st.session_state.get("show_login_code", False):
        html_block("""
    <style>
    .st-key-login_draft_code input {
        -webkit-text-security: disc !important;
    }
    </style>
    """)
    
    init_db()

    st.session_state.setdefault("n_autres_rows", 1)
    st.session_state.setdefault("access_granted", False)
    st.session_state.setdefault("current_draft_code", "")
    st.session_state.setdefault("admin_mode", False)
    st.session_state.setdefault("read_only_submitted", False)

    render_first_page_header()

    ADMIN_CODE = "USJ-ADMIN-2032"
    # TEMP_ADMIN_CODE_MIREILLE = "MIREILLE-ADMIN-TEST"

    if st.session_state.get("current_draft_code", "").strip().upper() in [ADMIN_CODE]:
        st.session_state["access_granted"] = True
        st.session_state["admin_mode"] = True

    # Admin preview mode: open the imported Word answers in the normal form layout,
    # but keep it strictly read-only so the original imported data_json is never overwritten.
    if st.session_state.get("view_selected_form", False):
        selected_view_code = st.session_state.get("view_selected_draft_code", "").strip().upper()
        if selected_view_code:
            selected_view_data = load_existing_draft_by_code(selected_view_code)
            if selected_view_data:
                preload_draft_into_session(selected_view_data)
                st.session_state["current_draft_code"] = selected_view_code
                st.session_state["access_granted"] = True
                st.session_state["read_only_submitted"] = True

    if st.session_state.get("admin_mode", False) and not st.session_state.get("view_selected_form", False):
        html_block('<div class="admin-screen-only"><h2 style="color:#001F5B; margin-top:0;">Gestion des réponses</h2></div>')

        

        if not DB_PATH.exists():
            st.error("Database not found in this app environment.")
            st.stop()

        with st.container(key="admin_word_import_block"):
            render_admin_word_importer()
        
        st.markdown("---")

        df = load_responses()

        def build_admin_export(df):
            export_rows = []

            for _, row in df.iterrows():
                original_data = json.loads(row["data_json"]) if row["data_json"] else {}
                admin_data = json.loads(row["admin_data_json"]) if pd.notna(row.get("admin_data_json", None)) and row.get("admin_data_json") else {}
        
                export_rows.append({
                    "id": row.get("id", ""),
                    "submitted_at": row.get("submitted_at", ""),
                    "draft_code": row.get("draft_code", ""),
                    "statut": row.get("statut", ""),
                    "groupe": row.get("respondent_unit", ""),
                    "participants": row.get("respondent_name", ""),
                    "original_answers_json": json.dumps(original_data, ensure_ascii=False),
                    "admin_modified_answers_json": json.dumps(admin_data, ensure_ascii=False),
                })
        
            return pd.DataFrame(export_rows)

        if df.empty:
            st.warning("Database exists, but no responses are saved yet. You can import a Word file above.")
            st.stop()

        html_block(f'<div class="admin-screen-only" style="background:#E9F8EF; padding:12px 18px; border-radius:8px; margin-bottom:18px;">{len(df)} réponse(s) enregistrée(s).</div>')


      

        st.markdown("---")
        
        admin_df = df.copy()
        admin_df["display_label"] = (
            admin_df["draft_code"].fillna("") + " | " +
            admin_df["respondent_name"].fillna("") + " | " +
            admin_df["respondent_unit"].fillna("") + " | " +
            admin_df["statut"].fillna("")
        )

        admin_df["fg_order"] = admin_df["draft_code"].astype(str).str.extract(r"FG(\d+)", expand=False).fillna(999).astype(int)
        admin_df = admin_df.sort_values(["fg_order", "draft_code"]).reset_index(drop=True)
        
        selected_response = st.selectbox(
            "Choisir un sous groupe",
            options=admin_df["display_label"].tolist()
        )

        selected_draft_code = selected_response.split(" | ")[0].strip()

        if st.button("Voir les réponses originales importées dans le formulaire", key="view_selected_form_button"):
            st.session_state["view_selected_form"] = True
            st.session_state["view_selected_draft_code"] = selected_draft_code
            st.session_state["current_draft_code"] = selected_draft_code
            st.session_state["read_only_submitted"] = True
            st.rerun()

        if st.button("Effacer les réponses du groupe sélectionné", key="delete_selected_group_button"):
            delete_response_by_code(selected_draft_code)
            st.success(f"Les réponses du groupe {selected_draft_code} ont été effacées.")
            st.rerun()

        if st.button("Autoriser ce groupe à modifier ses réponses", key="unlock_selected_group_button"):
            unlock_response_by_code(selected_draft_code)
            st.success(f"Le groupe {selected_draft_code} peut maintenant modifier ses réponses.")
            st.rerun()

        
        selected_row = admin_df[admin_df["draft_code"] == selected_draft_code].iloc[0]

        print_group_name = " - ".join(
            [part for part in [str(selected_row.get("respondent_name", "")).strip(), str(selected_row.get("respondent_unit", "")).strip()] if part]
        ) or selected_draft_code

        print_logo_src = image_to_base64(LOGO_PATH)
        print_logo_html = f'<img src="{print_logo_src}" alt="USJ">' if print_logo_src else ""

        html_block(
            f'<div class="admin-print-cover-header">'
            f'{print_logo_html}'
            f'<div class="admin-print-cover-title">{html_lib.escape(print_group_name)}</div>'
            f'</div>'
        )

        html_block(f'<div class="admin-print-title">{html_lib.escape(print_group_name)}</div>')

        original_data = json.loads(selected_row["data_json"]) if selected_row["data_json"] else {}

        if pd.notna(selected_row.get("admin_data_json", None)) and selected_row.get("admin_data_json"):
            admin_data = json.loads(selected_row["admin_data_json"])
        else:
            admin_data = {}


        section_map = {
            "I - Forces et faiblesses": ("swot_analysis", "facteurs_internes"),
            "II - Opportunités et menaces": ("swot_analysis", "facteurs_externes"),
            "III - Priorités": ("priorities_initiatives", None),
            "IV - Conclusion": ("pour_finir", None),
        }

        updated_all_admin_data = dict(admin_data) if isinstance(admin_data, dict) else {}

        def render_admin_title_bar(title, color):
            st.markdown(
                f'''
<div style="
background-color:#2F4F7F;
color:white;
padding:10px 14px;
font-size:18px;
font-weight:700;
text-align:center;
border-radius:6px;
margin-top:4px;
margin-bottom:6px;
box-shadow:0 2px 6px rgba(0,0,0,0.05);
">
{title}
</div>
''',
                unsafe_allow_html=True
            )


        def clean_admin_display_label(label):
            label = str(label or "")
            m = re.match(r"^priorite_(\d+)$", label)
            if m:
                return f"Priorité {m.group(1)}"
            m = re.match(r"^priority_only_(\d+)$", label)
            if m:
                return f"Priorité {m.group(1)}"
            m = re.match(r"^pour_finir_(\d+)_(\d+)$", label)
            if m:
                return f"Réponse {m.group(2)}"
            return label

        def render_original_answer_box(value):
            safe_value = html_lib.escape(str(value or "")).replace("\n", "<br>")
            if not safe_value.strip():
                safe_value = "&nbsp;"

            st.markdown(
                f'''
<div class="admin-original-answer-box" style="
background-color:#E3DED9;
padding:12px 16px;
border-radius:0px;
border:1.5px solid #595959;
margin-bottom:12px;
min-height:95px;
font-size:16px;
line-height:1.35;
color:#000000;
box-sizing:border-box;
">
{safe_value}
</div>
''',
                unsafe_allow_html=True
            )
      
        def render_admin_edit_box(label, value, key, height=95):
            return st.text_area(
                label=label,
                value=str(value) if value else "",
                height=height,
                key=key,
                label_visibility="collapsed",
                on_change=trigger_admin_autosave
            )

        def get_existing_admin_section(section_label, original_section):
            existing = updated_all_admin_data.get(section_label)
            if existing in [None, "", [], {}]:
                return original_section
            return existing

        def render_list_section(section_label, original_section):
            existing_admin_section = get_existing_admin_section(section_label, original_section)
            updated_admin_section = []

            if not isinstance(original_section, list):
                original_section = []

            if section_label == "I - Forces et faiblesses":
                field_names = ["Forces", "Faiblesses"]
            elif section_label == "II - Opportunités et menaces":
                field_names = ["Opportunités", "Menaces"]
            else:
                field_names = []
                for row in original_section:
                    if isinstance(row, dict):
                        for key in row.keys():
                            if key not in field_names:
                                field_names.append(key)

            def get_value_from_row(row, expected_field):
                if not isinstance(row, dict):
                    return ""

                if expected_field in row:
                    return row.get(expected_field, "")

                expected_lower = expected_field.lower()

                for key, value in row.items():
                    key_lower = str(key).lower()
                    if expected_lower in key_lower:
                        return value

                if expected_field == "Opportunités":
                    for key, value in row.items():
                        if "opportun" in str(key).lower():
                            return value

                if expected_field == "Menaces":
                    for key, value in row.items():
                        if "menace" in str(key).lower():
                            return value

                return ""

            def get_admin_value(saved_admin_row, expected_field, original_value):
                if not isinstance(saved_admin_row, dict):
                    return original_value

                def fallback_if_empty(value):
                    return value if str(value or "").strip() else original_value

                if expected_field in saved_admin_row:
                    return fallback_if_empty(saved_admin_row.get(expected_field, ""))

                expected_lower = expected_field.lower()

                for key, value in saved_admin_row.items():
                    key_lower = str(key).lower()
                    if expected_lower in key_lower:
                        return fallback_if_empty(value)

                if expected_field == "Opportunités":
                    for key, value in saved_admin_row.items():
                        if "opportun" in str(key).lower():
                            return fallback_if_empty(value)

                if expected_field == "Menaces":
                    for key, value in saved_admin_row.items():
                        if "menace" in str(key).lower():
                            return fallback_if_empty(value)

                return original_value

            admin_rows_key = f"admin_rows_{selected_draft_code}_{section_label}"

            if admin_rows_key not in st.session_state:
                admin_existing_len = len(existing_admin_section) if isinstance(existing_admin_section, list) else 0
                original_len = len(original_section) if isinstance(original_section, list) else 0
                st.session_state[admin_rows_key] = max(5, admin_existing_len, original_len)
            
            number_of_rows = st.session_state[admin_rows_key]

            

            col_left, col_right = st.columns(2)

            for field_index, field_name in enumerate(field_names):
                admin_col = col_left if field_index == 0 else col_right

                with admin_col:
                    render_admin_title_bar(field_name, USJ_RED)
                
                    for i in range(1, number_of_rows + 1):
                        row = original_section[i - 1] if i <= len(original_section) else {}
                        original_value = get_value_from_row(row, field_name)
                
                        saved_admin_row = {}
                        if (
                            isinstance(existing_admin_section, list)
                            and len(existing_admin_section) >= i
                            and isinstance(existing_admin_section[i - 1], dict)
                        ):
                            saved_admin_row = existing_admin_section[i - 1]
                
                        admin_value = get_admin_value(saved_admin_row, field_name, original_value)
                
                        while len(updated_admin_section) < i:
                            updated_admin_section.append({})
                
                        updated_admin_section[i - 1][field_name] = render_admin_edit_box(
                            label=f"{section_label}_{field_name}_{i}",
                            value=admin_value,
                            key=f"admin_edit_{selected_draft_code}_{section_label}_{field_name}_{i}",
                            height=95
                        )
                
                    if st.button(
                        "+",
                        key=f"add_admin_row_{selected_draft_code}_{section_label}_{field_name}"
                    ):
                        st.session_state[admin_rows_key] += 1
                        st.rerun()

            return updated_admin_section
            
        def render_dict_section(section_label, original_section):
            existing_admin_section = get_existing_admin_section(section_label, original_section)
            updated_admin_section = {}
        
            if not isinstance(original_section, dict):
                original_section = {}
        
            if not isinstance(existing_admin_section, dict):
                existing_admin_section = {}
        
            priority_rows_key = f"admin_priority_rows_{selected_draft_code}"
        
            if section_label == "III - Priorités":
                if priority_rows_key not in st.session_state:
                    st.session_state[priority_rows_key] = max(5, len(original_section), len(existing_admin_section))
        
                number_of_priority_rows = st.session_state[priority_rows_key]
        
                for i in range(1, number_of_priority_rows + 1):
                    key = f"priorite_{i}"
                    original_value = original_section.get(key, "")
                    saved_admin_value = existing_admin_section.get(key, original_value)
                    admin_value = saved_admin_value if str(saved_admin_value or "").strip() else original_value
        
                    updated_admin_section[key] = render_admin_edit_box(
                        label=f"{section_label}_{key}",
                        value=admin_value,
                        key=f"admin_edit_{selected_draft_code}_{section_label}_{key}",
                        height=95
                    )
        
                if st.button(
                    "+",
                    key=f"add_admin_priority_{selected_draft_code}"
                ):
                    st.session_state[priority_rows_key] += 1
                    st.rerun()
        
                return updated_admin_section
        
            if not original_section:
                st.info("Aucune réponse saisie pour cette section.")
                return original_section
        
            for key, original_value in original_section.items():
                saved_admin_value = existing_admin_section.get(key, original_value)
                admin_value = saved_admin_value if str(saved_admin_value or "").strip() else original_value
        
                render_admin_title_bar(clean_admin_display_label(key), USJ_RED)
        
                updated_admin_section[key] = render_admin_edit_box(
                    label=f"{section_label}_{key}",
                    value=admin_value,
                    key=f"admin_edit_{selected_draft_code}_{section_label}_{key}",
                    height=95
                )
        
            return updated_admin_section

        def render_conclusion_section(section_label, original_section):
            existing_admin_section = get_existing_admin_section(section_label, original_section)
            updated_admin_section = {}
        
            if not isinstance(original_section, dict):
                original_section = {}
        
            if not isinstance(existing_admin_section, dict):
                existing_admin_section = {}
        
            phrases = [
                "Nous souhaitons que l’USJ soit reconnue pour …",
                "Nous souhaitons que les employeurs disent que les diplômés de l’USJ sont …",
                "L’USJ serait un meilleur lieu d’enseignement si …",
            ]
        
            for i, phrase in enumerate(phrases, start=1):
                rows_key = f"admin_conclusion_rows_{selected_draft_code}_{i}"
        
                if rows_key not in st.session_state:
                    existing_count = max(
                        [
                            int(k.split("_")[-1])
                            for k in existing_admin_section.keys()
                            if k.startswith(f"pour_finir_{i}_") and k.split("_")[-1].isdigit()
                        ],
                        default=1
                    )
                    st.session_state[rows_key] = max(2, existing_count + 1)
        
                number_of_rows = st.session_state[rows_key]
        
                if i > 1:
                    html_block('<div class="admin-conclusion-no-page-break"></div>')
        
                st.markdown(
                    f"""
        <div style="
        font-size:18px;
        font-weight:700;
        color:{USJ_BLUE};
        margin-top:12px;
        margin-bottom:8px;
        ">
        &bull; {phrase}
        </div>
        """,
                    unsafe_allow_html=True
                )
        
                for j in range(number_of_rows):
                    key = f"pour_finir_{i}_{j}"
        
                    original_value = original_section.get(key, "")
                    saved_admin_value = existing_admin_section.get(key, original_value)
                    admin_value = saved_admin_value if str(saved_admin_value or "").strip() else original_value
        
                    html_block('<div class="admin-answer-row-wrapper">')
        
                    updated_admin_section[key] = render_admin_edit_box(
                        label=f"{section_label}_{key}",
                        value=admin_value,
                        key=f"admin_edit_{selected_draft_code}_{section_label}_{key}",
                        height=95
                    )
        
                    html_block('</div>')
        
                if st.button(
                    "+",
                    key=f"add_admin_conclusion_row_{selected_draft_code}_{i}"
                ):
                    st.session_state[rows_key] += 1
                    st.rerun()
        
            return updated_admin_section
        
        for section_index, (section_label, (main_key, sub_key)) in enumerate(section_map.items()):
            if section_index > 0:
                html_block('<div class="admin-print-page-break"></div>')

            st.markdown("---")
            section_header(section_label)

            if sub_key:
                original_section = original_data.get(main_key, {}).get(sub_key, [])
            else:
                original_section = original_data.get(main_key, {})

            if section_label == "IV - Conclusion":
                updated_all_admin_data[section_label] = render_conclusion_section(section_label, original_section)
            elif isinstance(original_section, list):
                updated_all_admin_data[section_label] = render_list_section(section_label, original_section)
            elif isinstance(original_section, dict):
                updated_all_admin_data[section_label] = render_dict_section(section_label, original_section)
            else:
                col_original_answer, col_admin_answer = st.columns(2)

                existing_admin_section = get_existing_admin_section(section_label, original_section)

                with col_original_answer:
                    render_admin_title_bar("Réponse originale", USJ_BLUE)
                    render_original_answer_box(original_section)

                with col_admin_answer:
                    render_admin_title_bar("Version admin", USJ_RED)
                    updated_all_admin_data[section_label] = render_admin_edit_box(
                        label=section_label,
                        value=existing_admin_section,
                        key=f"admin_edit_{selected_draft_code}_{section_label}",
                        height=95
                    )

            if st.button(
                "Enregistrer les réponses modifiées",
                key=f"save_admin_section_{section_index}_{selected_draft_code}",
                use_container_width=True
            ):
                save_admin_version_by_code(selected_draft_code, updated_all_admin_data)
                st.success("Les réponses modifiées ont été enregistrées.")
                st.rerun()

            if section_label == "II - Opportunités et menaces":
                render_swot_image_download_block(updated_all_admin_data, selected_row)

        st.markdown("---")

        html_block('<div class="admin-action-row-fix">')

        col_admin_print, col_admin_save, col_admin_spacer = st.columns(
            [1.35, 1.35, 2.30],
            gap="large",
            vertical_alignment="center"
        )

        with col_admin_print:
            components.html(
                """
                <!DOCTYPE html>
                <html>
                <head>
                <meta charset="UTF-8">
                <style>
                    html, body {
                        margin: 0;
                        padding: 0;
                        width: 100%;
                        height: 58px;
                        overflow: hidden;
                        background: transparent;
                        font-family: Candara, Calibri, Arial, sans-serif;
                        display: flex;
                        align-items: center;
                    }
                    button {
                        width: 100%;
                        min-width: 100%;
                        max-width: 100%;
                        height: 42px;
                        min-height: 42px;
                        max-height: 42px;
                        background-color: #8B1538;
                        color: #ffffff;
                        border: 1px solid #8B1538;
                        border-radius: 8px;
                        padding: 10px 22px;
                        margin: 0;
                        font-family: Candara, Calibri, Arial, sans-serif;
                        font-size: 18px;
                        font-weight: 800;
                        line-height: 1;
                        cursor: pointer;
                        display: flex;
                        align-items: center;
                        justify-content: center;
                        white-space: nowrap;
                        box-sizing: border-box;
                    }
                    button:hover {
                        background-color: #761130;
                        border-color: #761130;
                    }
                </style>
                </head>
                <body>
                    <button type="button" onclick="try { window.parent.print(); } catch(e) { try { window.top.print(); } catch(e2) { window.print(); } }">
                        Imprimer / Enregistrer en PDF
                    </button>
                </body>
                </html>
                """,
                height=58,
                scrolling=False
            )

       
        # Auto-save admin modifications on every rerun
        if st.session_state.get("admin_autosave_requested", False):
            save_admin_version_by_code(selected_draft_code, updated_all_admin_data)
            st.session_state["admin_autosave_requested"] = False

        df = load_responses()

        
        with col_admin_save:
            if st.button(
                "Enregistrer toutes les versions admin",
                key="save_admin_all_button",
                use_container_width=True
            ):
                save_admin_version_by_code(selected_draft_code, updated_all_admin_data)
                st.session_state["admin_saved_message"] = True
                st.rerun()

        if st.session_state.pop("admin_saved_message", False):
            st.success("Modifications admin enregistrées.")

        export_df = build_admin_export(df)

        csv_data = export_df.to_csv(index=False, encoding="utf-8-sig")

        excel_buffer = BytesIO()
        with pd.ExcelWriter(excel_buffer, engine="openpyxl") as writer:
            export_df.to_excel(writer, index=False, sheet_name="Export_Admin")

        col_export_csv, col_export_xlsx, col_export_empty = st.columns(
            [1.35, 1.35, 2.30],
            gap="large",
            vertical_alignment="center"
        )

        with col_export_csv:
            st.download_button(
                label="Télécharger CSV",
                data=csv_data,
                file_name="export_reponses_focus_groupes_admin.csv",
                mime="text/csv",
                use_container_width=True,
                key="download_export_csv"
            )

        with col_export_xlsx:
            st.download_button(
                label="Télécharger Excel",
                data=excel_buffer.getvalue(),
                file_name="export_reponses_focus_groupes_admin.xlsx",
                mime="application/vnd.openxmlformats-officedocument.spreadsheetml.sheet",
                use_container_width=True,
                key="download_export_excel"
            )

        html_block('</div>')

        st.stop()

    if not st.session_state["access_granted"]:
        # Login row: wide password field + compact button, like the reference layout.
        # The last empty column prevents the button from stretching across the page.
        col_code, col_button, col_empty = st.columns([4.7, 1.2, 1.25], gap="small")

        def submit_login_code():
            st.session_state["enter_form_clicked"] = True


        with col_code:
            login_code = st.text_input(
                "Mot de passe reçu par email",
                placeholder="",
                key="login_draft_code",
                on_change=submit_login_code
            )

        with col_button:
            st.markdown("<br>", unsafe_allow_html=True)

            enter_form = (
                st.button(
                    "Accéder au rapport",
                    key="access_report_button",
                    use_container_width=True
                )
                or st.session_state.pop("enter_form_clicked", False)
            )

        if enter_form:
            cleaned_code = login_code.strip().upper()
            
            if not cleaned_code:
                st.warning("Veuillez saisir un code personnel de reprise avant d’accéder au formulaire.")
                return

            if cleaned_code in [ADMIN_CODE]:
                st.session_state["access_granted"] = True
                st.session_state["admin_mode"] = True
                st.session_state["current_draft_code"] = cleaned_code
                st.rerun()

            if cleaned_code not in AUTHORIZED_TEST_CODES:
                st.error("Merci d'utiliser le mot de passe reçu par email.")
                return

            draft = load_existing_draft_by_code(cleaned_code)

            if draft:
                preload_draft_into_session(draft)
                st.session_state["current_draft_code"] = cleaned_code
                st.session_state["access_granted"] = True

                if draft.get("loaded_statut") == "Soumis":
                    st.session_state["read_only_submitted"] = True
                else:
                    st.session_state["read_only_submitted"] = False

                st.success("Vos réponses enregistrées ont été chargées.")
                st.rerun()

            else:
                st.session_state["current_draft_code"] = cleaned_code
                st.session_state["access_granted"] = True
                st.session_state["read_only_submitted"] = False

                st.session_state["institution"] = AUTHORIZED_TEST_CODES[cleaned_code]["institution"]
                st.session_state["responsable"] = AUTHORIZED_TEST_CODES[cleaned_code]["responsable"]
            
                st.info("Nouveau formulaire ouvert. Vous pouvez commencer à remplir vos réponses.")
                st.rerun()

        html_block('</div>')
        st.stop()

    mode = "Saisir une réponse"


    if mode == "Saisir une réponse":
        with st.container():

            if st.session_state.get("view_selected_form", False):
                if st.button("Retour à l'administration", key="back_to_admin_from_form"):
                    st.session_state["view_selected_form"] = False
                    st.session_state["admin_mode"] = True
                    st.session_state["access_granted"] = True
                    st.rerun()
                st.info(
                    "Affichage en lecture seule de la version originale importée depuis Word. "
                    "Les modifications des modérateurs se font dans la vue Admin et sont enregistrées séparément."
                )

            st.markdown("## Informations générales")

            col1, col2, col3 = st.columns(3)

            with col1:
                focus_group_options = [
                    "Sous groupe 1",
                    "Sous groupe 2",
                    "Sous groupe 3",
                    "Sous groupe 4",
                    "Sous groupe 5",
                ]


                institution_default = st.session_state.get("institution", "Sous groupe 1") or "Sous groupe 1"

                institution = st.text_input(
                    "Focus groupe",
                    value=institution_default,
                    key="institution_locked",
                    disabled=True
                )

            with col2:
                responsable = st.text_input(
                    "Nom des participants",
                    value=st.session_state.get("responsable", ""),
                    key="responsable",
                    placeholder="Nom 1, Nom 2, Nom 3...",
                    on_change=lambda: st.session_state.update({"autosave_requested": True})
                )

            with col3:
                st.text_input(
                    "Date",
                    value="2026-06-12",
                    disabled=True
                )
                response_date = "2026-06-12"

            st.markdown("---")

            section_header("Introduction")
            render_fixed_introduction()

            st.divider()
    
       
        # Section II removed for Focus Group version
        stakeholder_rows = []
        quick_save_after_stakeholders = False

        # Section III removed for Focus Group version
        internal_analysis = {}
        quick_save_after_internal = False

        # Section IV removed for Focus Group version
        external_analysis = {}
        quick_save_after_external = False

        section_header("I- FORCES ET FAIBLESSES – NIVEAU USJ")
        render_swot_intro()
        swot_analysis = render_swot_analysis()
        quick_save_after_swot = False

        st.divider()

        section_header("III - PRIORITES – Niveau USJ")
        render_priorities_intro()
        priorities_initiatives = render_priorities_table()
        quick_save_after_priorities = render_quick_save_button("quick_save_after_priorities")

        st.divider()

        section_header("IV- CONCLUSION")
        pour_finir = render_pour_finir()

        st.markdown("<br>", unsafe_allow_html=True)

        if st.session_state.get("read_only_submitted", False):
            st.info(
                "Le rapport de votre focus groupe a déjà été envoyé. "
                "Les modifications ne sont plus possibles."
            )
            save_draft = False
            submit_final = False

            st.markdown("<br>", unsafe_allow_html=True)
            render_print_icon_button()

        else:
            col_save_final, col_save_empty = st.columns(
                [1.25, 2.75],
                vertical_alignment="center"
            )

            with col_save_final:
                save_draft = st.button(
                    "Enregistrer et continuer plus tard",
                    key="save_draft_button",
                    use_container_width=True
                )
            

            st.markdown('<hr class="final-action-line">', unsafe_allow_html=True)

            col_submit_final, col_print_final, col_right_final = st.columns(
                [1.25, 1.25, 1.50],
                vertical_alignment="center"
            )

            with col_submit_final:
                submit_final = st.button(
                    "Envoyer la version finale\u00A0uniquement",
                    key="submit_final_button",
                    type="primary",
                    use_container_width=True
                )

            with col_print_final:
                print_icon_src = image_to_base64(PRINT_ICON_PATH)

                if print_icon_src:
                    components.html(
                        f"""
                        <div style="
                            height:150px;
                            display:flex;
                            align-items:center;
                            justify-content:center;
                            overflow:visible;
                            padding:0;
                            margin:0;
                        ">
                            <button onclick="window.parent.print()" title="Imprimer / Enregistrer en PDF" style="
                                background-color:transparent;
                                border:none;
                                cursor:pointer;
                                padding:0;
                                margin:0;
                                width:130px;
                                height:130px;
                                display:flex;
                                align-items:center;
                                justify-content:center;
                            ">
                                <img src="{print_icon_src}" alt="Imprimer / Enregistrer en PDF" style="
                                    width:130px;
                                    height:130px;
                                    object-fit:contain;
                                    display:block;
                                ">
                            </button>
                        </div>
                        """,
                        height=155
                    )

        quick_save_clicked = any([
            quick_save_after_stakeholders,
            quick_save_after_internal,
            quick_save_after_external,
            quick_save_after_swot,
            quick_save_after_priorities,
            st.session_state.get("quick_save_after_section_i_clicked", False),
            st.session_state.get("quick_save_after_section_ii_clicked", False),
        ])


        autosave_requested = st.session_state.get("autosave_requested", False)

        if st.session_state.get("view_selected_form", False):
            autosave_requested = False

        if (not st.session_state.get("view_selected_form", False)) and (save_draft or submit_final or quick_save_clicked or autosave_requested):
        
            word_limit_errors = []

            word_limit_errors.extend(
                find_word_limit_errors(
                    internal_analysis,
                    "Section III - Analyse interne",
                    max_words=500
                )
            )

            word_limit_errors.extend(
                find_word_limit_errors(
                    external_analysis,
                    "Section IV - Analyse externe",
                    max_words=500
                )
            )

            word_limit_errors.extend(
                find_word_limit_errors(
                    swot_analysis,
                    "Section V - Analyse SWOT",
                    max_words=30
                )
            )

            word_limit_errors.extend(
                find_word_limit_errors(
                    priorities_initiatives,
                    "Section VI - Priorités stratégiques et initiatives",
                    max_words=30
                )
            )

            if word_limit_errors:
                st.error(
                    "Certaines réponses dépassent la limite autorisée. "
                    "Merci de les réduire avant l’enregistrement."
                )

                for error in word_limit_errors:
                    st.warning(error)

                st.stop()

            statut = "Soumis" if submit_final else "Brouillon"

            metadata = {
                "institution": institution,
                "responsable": responsable,
                "email": "",
                "response_date": str(response_date),
                "statut": statut,
                "draft_code": st.session_state.get("current_draft_code", ""),
            }

            data = {
                "metadata": metadata,
                "introduction": {},
                "stakeholders": {
                    "rows": stakeholder_rows,
                },
                "internal_analysis": {},
                "external_analysis": {},
                "swot_analysis": swot_analysis,
                "priorities_initiatives": priorities_initiatives,
                "pour_finir": pour_finir,
            }

            try:
                draft_code = save_response(metadata, data)
                st.session_state["autosave_requested"] = False
                st.session_state["current_draft_code"] = draft_code

                if quick_save_clicked:
                    st.session_state["quick_save_success_key"] = st.session_state.get("last_quick_save_key", "")
                    st.rerun()

                if save_draft or quick_save_clicked:
                    st.success(
                        f"Vos réponses ont été enregistrées. Votre mot de passe pour reprendre plus tard : {draft_code}"
                    )

                if submit_final:
                    st.session_state["read_only_submitted"] = True
                    st.success("Merci.\nVos réponses ont été enregistrées.")
                    st.rerun()

            except ValueError as e:
                st.error(str(e))

if __name__ == "__main__":
    main()
