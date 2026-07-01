import re
import html
import base64
import difflib
from pathlib import Path
from datetime import datetime
from io import BytesIO

import pandas as pd
import streamlit as st
import streamlit.components.v1 as components

from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_CELL_VERTICAL_ALIGNMENT
from docx.oxml import OxmlElement
from docx.oxml.ns import qn

APP_TITLE = "PLAN STRATÉGIQUE USJ 2032"
SHEET_NAME = "All answers 29-6-2026"

USJ_BLUE = "#001F5B"
USJ_BLUE_2 = "#1F3C88"
USJ_RED = "#8B1538"
USJ_LIGHT_BLUE = "#EAF2F8"
USJ_TEXT = "#1B2A41"
LOGO_PATH = "LOGO NEW UAQ.png"


def clean(value):
    if pd.isna(value):
        return ""
    return str(value).strip()


def normalize(value):
    value = clean(value).lower()
    replacements = {
        "é": "e", "è": "e", "ê": "e", "ë": "e",
        "à": "a", "â": "a",
        "ù": "u", "û": "u",
        "ç": "c",
        "î": "i", "ï": "i",
        "ô": "o",
    }
    for old, new in replacements.items():
        value = value.replace(old, new)
    return value


def esc(value):
    return html.escape(clean(value)).replace("\n", "<br>")


def safe_filename(value):
    value = clean(value)
    value = re.sub(r"[^A-Za-z0-9À-ÿ._-]+", "_", value)
    return value.strip("_") or "rapport"


def html_block(content):
    st.markdown(content, unsafe_allow_html=True)


def image_to_base64(image_path):
    image_path = Path(image_path)
    if not image_path.exists():
        return ""
    suffix = image_path.suffix.lower().replace(".", "")
    mime_type = "png" if suffix == "png" else suffix
    with open(image_path, "rb") as img_file:
        encoded = base64.b64encode(img_file.read()).decode()
    return f"data:image/{mime_type};base64,{encoded}"


PRINT_CSS = f"""
body {{
    font-family: Candara, Calibri, Arial, sans-serif;
    background:white;
    color:{USJ_TEXT};
    margin:0;
    padding:0;
}}

.print-report {{
    background:white;
    border:1px solid #D0D6E0;
    border-radius:12px;
    padding:28px 34px;
}}

.report-page {{
    background:white;
    border:none;
    border-radius:0;
    padding:0;
    margin:0;
    width:auto;
    min-height:auto;
    box-sizing:border-box;
}}

.report-page + .report-page {{
    margin-top:28px;
}}

.usj-main-header {{
    display:flex;
    justify-content:space-between;
    align-items:flex-start;
    border-bottom:2px solid #D0D6E0;
    padding-bottom:18px;
    margin-bottom:28px;
}}

.usj-main-header h1 {{
    font-size:42px;
    margin:0 0 6px 0;
    color:{USJ_BLUE};
    font-weight:800;
}}

.usj-main-header p {{
    font-size:18px;
    font-weight:700;
    color:{USJ_BLUE_2};
    margin:0;
}}

.usj-logo-box {{
    display:flex;
    justify-content:flex-end;
    align-items:flex-start;
}}

.usj-logo {{
    width:170px;
    height:auto;
    object-fit:contain;
}}

.group-title {{
    text-align:center;
    font-size:28px;
    font-weight:800;
    color:{USJ_RED};
    margin:24px 0;
}}

.names {{
    text-align:center;
    font-size:21px;
    font-weight:700;
    color:#111;
    margin-bottom:32px;
}}

.section-header {{
    background-color:{USJ_LIGHT_BLUE};
    padding:12px 18px;
    border-radius:10px;
    border-left:7px solid {USJ_BLUE};
    margin-top:24px;
    margin-bottom:16px;
    page-break-inside:avoid;
    box-shadow:0 2px 10px rgba(0,0,0,0.06);
}}

.section-header h2 {{
    font-size:28px;
    margin:0;
    color:{USJ_BLUE};
}}

.two-cols {{
    display:grid;
    grid-template-columns:1fr 1fr;
    gap:20px;
}}

.col-title {{
    background:white;
    color:{USJ_BLUE};
    padding:10px 14px;
    text-align:center;
    font-size:21px;
    font-weight:900;
    border-radius:6px;
    margin-bottom:10px;
}}

.answer-box {{
    background:#E3DED9;
    border:1.5px solid #595959;
    padding:12px 16px;
    margin-bottom:10px;
    font-size:18px;
    line-height:1.35;
    color:#000;
    page-break-inside:avoid;
    break-inside:avoid;
}}

.answer-number {{
    color:{USJ_RED};
    font-weight:800;
    margin-right:6px;
}}

.conclusion-label {{
    color:{USJ_BLUE};
    font-size:19px;
    font-weight:800;
    margin:28px 0 10px 4px;
}}

.conclusion-box {{
    background:#E3DED9;
    border:1.5px solid #595959;
    padding:14px 18px;
    margin-bottom:18px;
    min-height:72px;
    font-size:20px;
    line-height:1.35;
    color:#000;
    page-break-inside:avoid;
    break-inside:avoid;
}}

.swot-section-wrapper {{
    break-inside:avoid;
    page-break-inside:avoid;
}}

.swot-matrix-title {{
    background-color:{USJ_LIGHT_BLUE};
    padding:12px 18px;
    border-radius:10px;
    border-left:7px solid {USJ_BLUE};
    margin-top:24px;
    margin-bottom:16px;
    color:{USJ_BLUE};
    font-size:28px;
    font-weight:800;
    page-break-after:avoid;
    break-after:avoid;
}}

.swot-grid {{
    display:grid;
    grid-template-columns:1fr 1fr;
    gap:18px;
    margin-top:14px;
    page-break-before:avoid;
    break-before:avoid;
}}

.swot-card {{
    border-radius:14px;
    padding:16px;
    border:2px solid var(--accent);
    background:var(--bg);
    page-break-inside:avoid;
    break-inside:avoid;
}}

.swot-card h3 {{
    margin:0 0 10px 0;
    color:var(--accent);
    font-size:24px;
    font-weight:900;
    border-bottom:2px solid var(--accent);
    padding-bottom:8px;
}}

.swot-card ul {{
    margin:0;
    padding-left:22px;
}}

.swot-card li {{
    margin-bottom:8px;
    font-size:16px;
    line-height:1.3;
}}

.print-page-break {{
    page-break-before:always;
}}

.print-button {{
    display:inline-block;
    background:#8B1538;
    color:white;
    border:1px solid #8B1538;
    border-radius:8px;
    padding:10px 22px;
    font-family:Candara,Calibri,Arial,sans-serif;
    font-size:18px;
    font-weight:800;
    cursor:pointer;
    margin:0 12px 20px 0;
    text-decoration:none;
}}

.print-actions {{
    margin-bottom:20px;
}}

@media print {{
    @page {{
        size:A4 portrait;
        margin:8mm 9mm;
    }}

    body {{
        margin:0;
    }}

    .print-button {{
        display:none;
    }}

    .print-report {{
        border:none;
        padding:0;
        background:white;
    }}

    .report-page {{
        border:none;
        border-radius:0;
        padding:0;
        margin:0;
        width:auto;
        min-height:auto;
        page-break-after:always;
        break-after:page;
    }}

    .report-page:last-child {{
        page-break-after:auto;
        break-after:auto;
    }}

    .usj-main-header h1 {{
        font-size:22px;
        line-height:1.05;
        margin-bottom:1mm;
    }}

    .usj-main-header p {{
        font-size:10px;
    }}

    .usj-logo {{
        width:36mm;
        max-width:36mm;
        height:auto;
    }}

    .group-title {{
        font-size:18px;
        margin:8mm 0 4mm 0;
    }}

    .names {{
        font-size:12px;
        margin-bottom:8mm;
    }}

    .section-header {{
        margin-top:5mm;
        margin-bottom:3mm;
        padding:7px 11px;
        box-shadow:none;
    }}

    .section-header h2 {{
        font-size:16px;
    }}

    .swot-section-wrapper {{
        break-inside:avoid;
        page-break-inside:avoid;
    }}

    .swot-matrix-title {{
        font-size:16px;
        margin-top:5mm;
        margin-bottom:3mm;
        padding:7px 11px;
        page-break-after:avoid;
        break-after:avoid;
    }}

    .col-title {{
        font-size:13px;
        padding:6px;
        color:{USJ_BLUE};
        background:white;
        font-weight:900;
    }}

    .answer-box,
    .conclusion-box {{
        font-size:10.5px;
        line-height:1.2;
        padding:6px;
        margin-bottom:2mm;
        page-break-inside:avoid;
        break-inside:avoid;
    }}

    .conclusion-label {{
        font-size:11px;
        margin:4mm 0 2mm 0;
    }}

    .swot-card h3 {{
        font-size:13px;
    }}

    .swot-card li {{
        font-size:9.5px;
        line-height:1.15;
    }}
}}
"""


APP_CSS = f"""
<style>
html, body, .stApp {{
    font-family: Candara, Calibri, Arial, sans-serif !important;
    color:{USJ_TEXT};
}}

header, footer, #MainMenu {{
    display:none !important;
}}

.block-container {{
    max-width:100% !important;
    padding-left:3rem !important;
    padding-right:3rem !important;
}}

h1, h2, h3 {{
    color:{USJ_BLUE} !important;
    font-weight:800 !important;
}}

.stButton button,
.stDownloadButton button {{
    background-color:#0070C0 !important;
    color:white !important;
    border-radius:8px !important;
    border:1px solid #0070C0 !important;
    font-weight:800 !important;
    font-size:18px !important;
}}

.stButton button p,
.stDownloadButton button p {{
    color:white !important;
}}
</style>
"""


def get_col(df, possible_names):
    cols = {normalize(c): c for c in df.columns}
    for name in possible_names:
        n = normalize(name)
        if n in cols:
            return cols[n]
    return None


def get_answers(df_group, section_contains=None, category_contains=None):
    temp = df_group.copy()

    if section_contains:
        temp = temp[temp["section_norm"].str.contains(section_contains, na=False)]

    if category_contains:
        temp = temp[temp["category_norm"].str.contains(category_contains, na=False)]

    return [clean(x) for x in temp["Final_Answer"].tolist() if clean(x)]


def answer_boxes(answers):
    if not answers:
        return '<div class="answer-box">Aucune réponse disponible.</div>'

    return "".join(
        f"""
        <div class="answer-box">
            <span class="answer-number">{i}.</span>{esc(answer)}
        </div>
        """
        for i, answer in enumerate(answers, start=1)
    )


def swot_card(title, answers, accent, bg):
    if not answers:
        li = "<li>Aucune réponse disponible.</li>"
    else:
        li = "".join(f"<li>{esc(x)}</li>" for x in answers)

    return f"""
    <div class="swot-card" style="--accent:{accent}; --bg:{bg};">
        <h3>{title}</h3>
        <ul>{li}</ul>
    </div>
    """


def swot_matrix_html(forces, faiblesses, opportunites, menaces):
    return f"""
    <div class="swot-section-wrapper">
        <div class="swot-matrix-title">
            Matrice SWOT - Niveau USJ
        </div>

        <div class="swot-grid">
            {swot_card("FORCES", forces, USJ_BLUE, "#DDEFF7")}
            {swot_card("FAIBLESSES", faiblesses, USJ_RED, "#FBE3C3")}
            {swot_card("OPPORTUNITÉS", opportunites, "#2F6B2F", "#E2F2D3")}
            {swot_card("MENACES", menaces, USJ_RED, "#F4C6C4")}
        </div>
    </div>
    """


def conclusion_html(df_group):
    conclusion_rows = df_group[
        df_group["section_norm"].str.contains("conclusion", na=False)
    ].copy()

    if conclusion_rows.empty:
        return '<div class="conclusion-box">Aucune réponse disponible.</div>'

    blocks = ""
    question_order = []

    for question in conclusion_rows["question"].tolist():
        question = clean(question)
        if question and question not in question_order:
            question_order.append(question)

    for question in question_order:
        temp = conclusion_rows[conclusion_rows["question"] == question]
        answers = [clean(x) for x in temp["Final_Answer"].tolist() if clean(x)]

        blocks += f"""
        <div class="conclusion-label">• {esc(question)}</div>
        """

        if not answers:
            blocks += '<div class="conclusion-box"></div>'
        else:
            for answer in answers:
                blocks += f"""
                <div class="conclusion-box">{esc(answer)}</div>
                """

    return blocks


def build_one_report_html(df_group, participant_type, title_label, hide_names):
    names = sorted(set(clean(x) for x in df_group["participants"].dropna() if clean(x)))

    names_html = ""
    if not hide_names and names:
        names_html = f"""
        <div class="names">
            Nom des participants : {esc("; ".join(names))}
        </div>
        """

    logo_src = image_to_base64(LOGO_PATH)
    logo_html = f'<img src="{logo_src}" class="usj-logo">' if logo_src else ""

    forces = get_answers(df_group, section_contains="forces", category_contains="force")
    faiblesses = get_answers(df_group, section_contains="forces", category_contains="faiblesse")
    opportunites = get_answers(df_group, section_contains="opportunites", category_contains="opportun")
    menaces = get_answers(df_group, section_contains="opportunites", category_contains="menace")
    priorites = get_answers(df_group, section_contains="prior")

    return f"""
<div class="print-report">

    <div class="report-page">
        <div class="usj-main-header">
            <div>
                <h1>PLAN STRATÉGIQUE USJ 2032</h1>
                <p>Focus groupe</p>
            </div>

            <div class="usj-logo-box">
                {logo_html}
            </div>
        </div>

        <div class="group-title">{esc(title_label)}</div>

        {names_html}

        <div class="section-header">
            <h2>I - Forces et faiblesses</h2>
        </div>

        <div class="two-cols">
            <div>
                <div class="col-title">Forces</div>
                {answer_boxes(forces)}
            </div>
            <div>
                <div class="col-title">Faiblesses</div>
                {answer_boxes(faiblesses)}
            </div>
        </div>
    </div>

    <div class="report-page">
        <div class="section-header">
            <h2>II - Opportunités et menaces</h2>
        </div>

        <div class="two-cols">
            <div>
                <div class="col-title">Opportunités</div>
                {answer_boxes(opportunites)}
            </div>
            <div>
                <div class="col-title">Menaces</div>
                {answer_boxes(menaces)}
            </div>
        </div>

        {swot_matrix_html(forces, faiblesses, opportunites, menaces)}
    </div>

    <div class="report-page">
        <div class="section-header">
            <h2>III - Priorités</h2>
        </div>

        <div>
            {answer_boxes(priorites)}
        </div>
    </div>

    <div class="report-page">
        <div class="section-header">
            <h2>IV - Conclusion</h2>
        </div>

        <div>
            {conclusion_html(df_group)}
        </div>
    </div>

</div>
"""


def build_full_html(html_report, selected_type, selected_label, docx_base64="", docx_filename="rapport.docx"):
    word_button = ""
    if docx_base64:
        word_button = f"""
<a class="print-button"
   href="data:application/vnd.openxmlformats-officedocument.wordprocessingml.document;base64,{docx_base64}"
   download="{esc(docx_filename)}">
   Télécharger Word
</a>
"""

    return f"""
<!DOCTYPE html>
<html>
<head>
<meta charset="UTF-8">
<title>{esc(selected_type)} - {esc(selected_label)}</title>
<style>
{PRINT_CSS}
</style>
</head>
<body>
<div class="print-actions">
    <button class="print-button" onclick="window.print()">Imprimer / Enregistrer en PDF</button>
    {word_button}
</div>
{html_report}
</body>
</html>
"""



def set_cell_shading(cell, fill):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:fill"), fill.replace("#", ""))


def set_cell_border(cell, color="595959", size="8"):
    tc_pr = cell._tc.get_or_add_tcPr()
    borders = tc_pr.first_child_found_in("w:tcBorders")
    if borders is None:
        borders = OxmlElement("w:tcBorders")
        tc_pr.append(borders)
    for edge in ("top", "left", "bottom", "right"):
        tag = "w:{}".format(edge)
        element = borders.find(qn(tag))
        if element is None:
            element = OxmlElement(tag)
            borders.append(element)
        element.set(qn("w:val"), "single")
        element.set(qn("w:sz"), size)
        element.set(qn("w:space"), "0")
        element.set(qn("w:color"), color.replace("#", ""))


def clear_cell(cell):
    cell.text = ""
    if not cell.paragraphs:
        cell.add_paragraph()
    return cell.paragraphs[0]


def set_paragraph_font(paragraph, size=10.5, bold=False, color="000000"):
    for run in paragraph.runs:
        run.font.name = "Candara"
        run._element.rPr.rFonts.set(qn("w:eastAsia"), "Candara")
        run.font.size = Pt(size)
        run.font.bold = bold
        run.font.color.rgb = RGBColor.from_string(color.replace("#", ""))


def add_word_header(document, participant_type):
    table = document.add_table(rows=1, cols=2)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.autofit = False
    table.columns[0].width = Inches(5.1)
    table.columns[1].width = Inches(1.5)

    for row in table.rows:
        for cell in row.cells:
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.TOP
            set_cell_border(cell, "FFFFFF", "0")

    left = table.cell(0, 0)
    right = table.cell(0, 1)

    p = clear_cell(left)

    r = p.add_run("PLAN STRATÉGIQUE USJ 2032")
    r.bold = True
    r.font.name = "Candara"
    r._element.rPr.rFonts.set(qn("w:eastAsia"), "Candara")
    r.font.size = Pt(20)
    r.font.color.rgb = RGBColor(0, 31, 91)
    
    p2 = left.add_paragraph()
    r2 = p2.add_run(f"Focus groupe - {participant_type}")
    r2.bold = True
    r2.font.name = "Candara"
    r2._element.rPr.rFonts.set(qn("w:eastAsia"), "Candara")
    r2.font.size = Pt(16)
    r2.font.color.rgb = RGBColor(31, 60, 136)

    p_logo = clear_cell(right)
    p_logo.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    if Path(LOGO_PATH).exists():
        p_logo.add_run().add_picture(
            LOGO_PATH,
            width=Inches(1.73),
            height=Inches(0.74)
        )

    line = document.add_paragraph()
    line.paragraph_format.space_after = Pt(10)
    line_run = line.add_run("_")
    line_run.font.color.rgb = RGBColor(208, 214, 224)
    line_run.font.size = Pt(1)


def add_section_header(document, text):
    table = document.add_table(rows=1, cols=1)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    cell = table.cell(0, 0)
    set_cell_shading(cell, USJ_LIGHT_BLUE)
    set_cell_border(cell, USJ_LIGHT_BLUE, "6")
    cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
    p = clear_cell(cell)
    p.paragraph_format.left_indent = Inches(0.05)
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.space_after = Pt(0)
    r = p.add_run(text)
    r.bold = True
    r.font.name = "Candara"
    r._element.rPr.rFonts.set(qn("w:eastAsia"), "Candara")
    r.font.size = Pt(16)
    r.font.color.rgb = RGBColor(0, 31, 91)
    document.add_paragraph().paragraph_format.space_after = Pt(2)


def add_answer_to_cell(cell, number, answer, font_size=10.5):
    set_cell_shading(cell, "#E3DED9")
    set_cell_border(cell, "595959", "8")
    cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.TOP
    p = clear_cell(cell)
    p.paragraph_format.space_after = Pt(0)
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.left_indent = Inches(0.02)
    r1 = p.add_run(f"{number}. ")
    r1.bold = True
    r1.font.name = "Candara"
    r1._element.rPr.rFonts.set(qn("w:eastAsia"), "Candara")
    r1.font.size = Pt(font_size)
    r1.font.color.rgb = RGBColor(139, 21, 56)
    r2 = p.add_run(answer)
    r2.font.name = "Candara"
    r2._element.rPr.rFonts.set(qn("w:eastAsia"), "Candara")
    r2.font.size = Pt(font_size)
    r2.font.color.rgb = RGBColor(0, 0, 0)


def add_single_answer_box(document, number, answer, font_size=10.5, min_height=None):
    table = document.add_table(rows=1, cols=1)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    cell = table.cell(0, 0)
    if min_height:
        table.rows[0].height = min_height
    add_answer_to_cell(cell, number, answer, font_size=font_size)
    document.add_paragraph().paragraph_format.space_after = Pt(1)


def add_two_column_answers(document, left_title, left_answers, right_title, right_answers):
    max_len = max(len(left_answers), len(right_answers), 1)

    table = document.add_table(rows=max_len + 1, cols=2)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.autofit = False
    table.style = "Table Grid"
    table.columns[0].width = Inches(3.15)
    table.columns[1].width = Inches(3.15)

    # Repeat the header row if the table continues on another page.
    header_tr_pr = table.rows[0]._tr.get_or_add_trPr()
    tbl_header = OxmlElement("w:tblHeader")
    tbl_header.set(qn("w:val"), "true")
    header_tr_pr.append(tbl_header)

    for i, title in enumerate([left_title, right_title]):
        cell = table.cell(0, i)
        set_cell_border(cell, "595959", "8")
        cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
        p = clear_cell(cell)
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p.paragraph_format.space_before = Pt(0)
        p.paragraph_format.space_after = Pt(0)
        r = p.add_run(title)
        r.bold = True
        r.font.name = "Candara"
        r._element.rPr.rFonts.set(qn("w:eastAsia"), "Candara")
        r.font.size = Pt(13)
        r.font.color.rgb = RGBColor(0, 31, 91)

    for i in range(max_len):
        left_text = left_answers[i] if i < len(left_answers) else ""
        right_text = right_answers[i] if i < len(right_answers) else ""

        for col_idx, text in enumerate([left_text, right_text]):
            cell = table.cell(i + 1, col_idx)
            set_cell_border(cell, "595959", "8")
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.TOP
            p = clear_cell(cell)
            p.paragraph_format.space_after = Pt(0)
            p.paragraph_format.space_before = Pt(0)
            p.paragraph_format.left_indent = Inches(0.02)

            if text:
                r = p.add_run(text)
                r.font.name = "Candara"
                r._element.rPr.rFonts.set(qn("w:eastAsia"), "Candara")
                r.font.size = Pt(10.2)
                r.font.color.rgb = RGBColor(0, 0, 0)

    document.add_paragraph().paragraph_format.space_after = Pt(2)


def add_swot_cell(cell, title, answers, accent_hex, bg_hex):
    set_cell_shading(cell, bg_hex)
    set_cell_border(cell, accent_hex, "12")
    p = clear_cell(cell)
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    r = p.add_run(title)
    r.bold = True
    r.font.name = "Candara"
    r._element.rPr.rFonts.set(qn("w:eastAsia"), "Candara")
    r.font.size = Pt(12)
    r.font.color.rgb = RGBColor.from_string(accent_hex.replace("#", ""))
    
    for answer in answers or ["Aucune réponse disponible."]:
        p_item = cell.add_paragraph(style=None)
        p_item.paragraph_format.left_indent = Inches(0.12)
        p_item.paragraph_format.space_after = Pt(2)
    
        answer_clean = answer.replace("\r\n", "\n").replace("\r", "\n").strip()

        if "\n" in answer_clean:
            parts = [x.strip() for x in answer_clean.split("\n") if x.strip()]
        else:
            parts = re.split(r"(?<=[.!?])\s+(?=[A-ZÉÈÊÀÂÎÏÔÙÛÇ])", answer_clean)
            parts = [x.strip() for x in parts if x.strip()]
    
        if parts:
            run = p_item.add_run("• " + parts[0])
            run.font.name = "Candara"
            run._element.rPr.rFonts.set(qn("w:eastAsia"), "Candara")
            run.font.size = Pt(9.2)
            run.font.color.rgb = RGBColor(0, 0, 0)
    
            for extra_part in parts[1:]:
                run_break = p_item.add_run()
                run_break.add_break()
    
                run_extra = p_item.add_run(extra_part)
                run_extra.font.name = "Candara"
                run_extra._element.rPr.rFonts.set(qn("w:eastAsia"), "Candara")
                run_extra.font.size = Pt(9.2)
                run_extra.font.color.rgb = RGBColor(0, 0, 0)
    


def add_swot_matrix_word(document, forces, faiblesses, opportunites, menaces):
    add_section_header(document, "Matrice SWOT - Niveau USJ")
    table = document.add_table(rows=2, cols=2)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.autofit = False
    table.columns[0].width = Inches(3.15)
    table.columns[1].width = Inches(3.15)
    add_swot_cell(table.cell(0, 0), "FORCES", forces, "#001F5B", "#DDEFF7")
    add_swot_cell(table.cell(0, 1), "FAIBLESSES", faiblesses, "#8B1538", "#FBE3C3")
    add_swot_cell(table.cell(1, 0), "OPPORTUNITÉS", opportunites, "#2F6B2F", "#E2F2D3")
    add_swot_cell(table.cell(1, 1), "MENACES", menaces, "#8B1538", "#F4C6C4")
    document.add_paragraph().paragraph_format.space_after = Pt(2)


def lowercase_first_letter(text):
    text = clean(text)
    if not text:
        return text
    return text[0].lower() + text[1:]


def add_conclusion_word(document, df_group):
    conclusion_rows = df_group[df_group["section_norm"].str.contains("conclusion", na=False)].copy()

    if conclusion_rows.empty:
        table = document.add_table(rows=1, cols=1)
        table.alignment = WD_TABLE_ALIGNMENT.CENTER
        cell = table.cell(0, 0)
        set_cell_border(cell, "595959", "8")
        p = clear_cell(cell)
        r = p.add_run("Aucune réponse disponible.")
        r.font.name = "Candara"
        r._element.rPr.rFonts.set(qn("w:eastAsia"), "Candara")
        r.font.size = Pt(10.5)
        r.font.color.rgb = RGBColor(0, 0, 0)
        return

    question_order = []
    for question in conclusion_rows["question"].tolist():
        question = clean(question)
        if question and question not in question_order:
            question_order.append(question)

    for question in question_order:
        p_q = document.add_paragraph()
        p_q.paragraph_format.space_before = Pt(4)
        p_q.paragraph_format.space_after = Pt(3)

        r_q = p_q.add_run("• " + question)
        r_q.bold = True
        r_q.font.name = "Candara"
        r_q._element.rPr.rFonts.set(qn("w:eastAsia"), "Candara")
        r_q.font.size = Pt(11)
        r_q.font.color.rgb = RGBColor(0, 31, 91)

        answers = [
            lowercase_first_letter(x)
            for x in conclusion_rows[conclusion_rows["question"] == question]["Final_Answer"].tolist()
            if clean(x)
        ]

        if not answers:
            answers = [""]

        table = document.add_table(rows=len(answers), cols=1)
        table.alignment = WD_TABLE_ALIGNMENT.CENTER
        table.autofit = True

        for i, answer in enumerate(answers):
            cell = table.cell(i, 0)
            set_cell_border(cell, "595959", "8")

            p_a = clear_cell(cell)
            p_a.paragraph_format.space_after = Pt(0)
            p_a.paragraph_format.space_before = Pt(0)

            r_a = p_a.add_run(answer)
            r_a.font.name = "Candara"
            r_a._element.rPr.rFonts.set(qn("w:eastAsia"), "Candara")
            r_a.font.size = Pt(10.5)
            r_a.font.color.rgb = RGBColor(0, 0, 0)

        document.add_paragraph().paragraph_format.space_after = Pt(4)

def get_focus_group_date(participant_type):
    dates = {
        "Teachers": "4 juin 2026",
        "Enseignants": "4 juin 2026",
        "Students": "5 juin 2026",
        "Étudiants": "5 juin 2026",
        "Directors": "10 juin 2026",
        "Directeurs": "10 juin 2026",
        "PSG": "12 juin 2026",
        "Anciens": "15 juin 2026",
    }
    return dates.get(participant_type, "")

def add_field(paragraph, field_code):
    run = paragraph.add_run()
    fld_char_begin = OxmlElement("w:fldChar")
    fld_char_begin.set(qn("w:fldCharType"), "begin")

    instr_text = OxmlElement("w:instrText")
    instr_text.set(qn("xml:space"), "preserve")
    instr_text.text = field_code

    fld_char_separate = OxmlElement("w:fldChar")
    fld_char_separate.set(qn("w:fldCharType"), "separate")

    fld_char_end = OxmlElement("w:fldChar")
    fld_char_end.set(qn("w:fldCharType"), "end")

    run._r.append(fld_char_begin)
    run._r.append(instr_text)
    run._r.append(fld_char_separate)
    run._r.append(fld_char_end)

def build_word_docx(df_group, participant_type, title_label, hide_names):
    document = Document()

    section = document.sections[0]
    section.page_width = Inches(8.27)
    section.page_height = Inches(11.69)
    section.top_margin = Inches(0.32)
    section.bottom_margin = Inches(0.32)
    section.left_margin = Inches(0.35)
    section.right_margin = Inches(0.35)

    section.footer_distance = Inches(0.1)

    footer = section.footer
    footer_table = footer.add_table(rows=1, cols=3, width=Inches(7.57))
    footer_table.alignment = WD_TABLE_ALIGNMENT.CENTER
    footer_table.autofit = False
    
    footer_table.columns[0].width = Inches(2.5)
    footer_table.columns[1].width = Inches(2.6)
    footer_table.columns[2].width = Inches(2.5)
    
    date_text = get_focus_group_date(participant_type)
    
    footer_table.cell(0, 0).text = date_text
    footer_table.cell(0, 1).text = f"Focus groupe - {participant_type}"
    
    p_page = footer_table.cell(0, 2).paragraphs[0]
    p_page.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    add_field(p_page, "PAGE")
    p_page.add_run("/")
    add_field(p_page, "NUMPAGES")
    
    for i, cell in enumerate(footer_table.rows[0].cells):
        set_cell_border(cell, "FFFFFF", "0")
    
        for p in cell.paragraphs:
            if i == 0:
                p.alignment = WD_ALIGN_PARAGRAPH.LEFT
            elif i == 1:
                p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            else:
                p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    
            for r in p.runs:
                r.font.name = "Candara"
                r._element.rPr.rFonts.set(qn("w:eastAsia"), "Candara")
                r.font.size = Pt(9)
                r.font.bold = True
                r.font.color.rgb = RGBColor(0, 31, 91)
    
    styles = document.styles
    styles["Normal"].font.name = "Candara"
    styles["Normal"]._element.rPr.rFonts.set(qn("w:eastAsia"), "Candara")
    styles["Normal"].font.size = Pt(10.5)

    names = sorted(set(clean(x) for x in df_group["participants"].dropna() if clean(x)))

    forces = get_answers(df_group, section_contains="forces", category_contains="force")
    faiblesses = get_answers(df_group, section_contains="forces", category_contains="faiblesse")
    opportunites = get_answers(df_group, section_contains="opportunites", category_contains="opportun")
    menaces = get_answers(df_group, section_contains="opportunites", category_contains="menace")
    priorites = get_answers(df_group, section_contains="prior")

    add_word_header(document, participant_type)

    title = document.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    title.paragraph_format.space_before = Pt(4)
    title.paragraph_format.space_after = Pt(8)
    run = title.add_run(title_label)
    run.bold = True
    run.font.name = "Candara"
    run._element.rPr.rFonts.set(qn("w:eastAsia"), "Candara")
    run.font.size = Pt(16)
    run.font.color.rgb = RGBColor(139, 21, 56)

    if names and not hide_names:
        p_names = document.add_paragraph()
        p_names.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p_names.paragraph_format.space_after = Pt(10)
        r_names = p_names.add_run("Nom des participants : " + "; ".join(names))
        r_names.bold = True
        r_names.font.name = "Candara"
        r_names._element.rPr.rFonts.set(qn("w:eastAsia"), "Candara")
        r_names.font.size = Pt(12)
        r_names.font.color.rgb = RGBColor(0, 0, 0)

    add_section_header(document, "I - Forces et faiblesses")
    add_two_column_answers(document, "Forces", forces, "Faiblesses", faiblesses)

    document.add_page_break()
    add_section_header(document, "II - Opportunités et menaces")
    add_two_column_answers(document, "Opportunités", opportunites, "Menaces", menaces)
    
    document.add_page_break()
    add_swot_matrix_word(document, forces, faiblesses, opportunites, menaces)

    document.add_page_break()
    add_section_header(document, "III - Priorités")
    priority_table = document.add_table(rows=max(len(priorites), 1), cols=1)
    priority_table.alignment = WD_TABLE_ALIGNMENT.CENTER
    priority_table.autofit = False
    priority_table.columns[0].width = Inches(7.2)
    
    if priorites:
        for i, answer in enumerate(priorites):
            cell = priority_table.cell(i, 0)
            set_cell_border(cell, "595959", "8")
            p = clear_cell(cell)
            p.paragraph_format.space_after = Pt(0)
    
            run = p.add_run(answer)
            run.font.name = "Candara"
            run._element.rPr.rFonts.set(qn("w:eastAsia"), "Candara")
            run.font.size = Pt(10.5)
            run.font.color.rgb = RGBColor(0, 0, 0)
    else:
        cell = priority_table.cell(0, 0)
        set_cell_border(cell, "595959", "8")
        p = clear_cell(cell)
        run = p.add_run("Aucune réponse disponible.")
        run.font.name = "Candara"
        run._element.rPr.rFonts.set(qn("w:eastAsia"), "Candara")
        run.font.size = Pt(10.5)
        run.font.color.rgb = RGBColor(0, 0, 0)
    

    document.add_page_break()
    add_section_header(document, "IV - Conclusion")
    add_conclusion_word(document, df_group)

    buffer = BytesIO()
    document.save(buffer)
    buffer.seek(0)
    return buffer.getvalue()


# ============================================================
# AI ASSISTED SWOT THEMATIC CLASSIFICATION
# ============================================================

THEME_TAXONOMY = {
    "PSG": {
        "internal": [
            "Soutenabilité financière",
            "Gouvernance et Leadership (gestion, relation, représentation, etc.)",
            "Stratégie académique et qualité d’enseignement",
            "Recherche et Innovation",
            "Ressources documentaires et Environnement digital",
            "Succès des étudiants (recrutement, accompagnement, services de support, employabilité, etc.)",
            "Ressources humaines",
            "Stratégie et mobilité internationales",
            "Mission sociétale",
            "Espace et infrastructures",
            "Environnement de travail",
            "Diversité et inclusion",
            "Développement Durable (ODD)",
            "Autre",
        ],
        "external": [
            "Marché du travail et Associations professionnelles",
            "Concurrence avec les autres universités",
            "Intelligence artificielle",
            "Réputation et image",
            "Environnement politique et économique, émigration",
            "Autres",
        ],
    },
    "Anciens": {
        "internal": [
            "Soutenabilité financière",
            "Gouvernance et Leadership (Gestion, relation, représentation, etc.)",
            "Stratégie académique et qualité d’enseignement",
            "Recherche et Innovation",
            "Succès des étudiants (recrutement, accompagnement, services de support, employabilité, etc.)",
            "Ressources humaines",
            "Mission sociétale",
            "Espace et infrastructures",
            "Autre",
        ],
        "external": [
            "Marché du travail et Associations professionnelles",
            "Concurrence avec les autres universités",
            "Intelligence artificielle",
            "Réputation et image",
            "Environnement politique et économique, émigration",
            "Autres",
        ],
    },
    "Directeurs": {
        "internal": [
            "Soutenabilité financière",
            "Gouvernance et Leadership (gestion, relation, représentation, etc.)",
            "Stratégie académique et qualité d’enseignement",
            "Recherche et Innovation",
            "Ressources documentaires et Environnement digital",
            "Succès des étudiants (recrutement, accompagnement, services de support, employabilité, etc.)",
            "Ressources humaines",
            "Stratégie et mobilité internationales",
            "Mission sociétale",
            "Espace et infrastructures",
            "Environnement de travail",
            "Diversité et inclusion",
            "Développement Durable (ODD)",
            "Autre",
        ],
        "external": [
            "Marché du travail et relations avec les employeurs",
            "Concurrence avec les autres universités",
            "Intelligence artificielle",
            "Réputation et image",
            "Environnement politique et économique, émigration",
            "Autres",
        ],
    },
    "Étudiants": {
        "internal": [
            "Qualité de l’enseignement (Accompagnement, Orientation, Tutorat, etc.)",
            "Relation avec l’administration et les enseignants",
            "Représentation des étudiants",
            "Ressources documentaires, numériques et informatiques",
            "Espace et infrastructure",
            "Vie universitaire (Vie étudiante, Sport, Aide psychologique, etc.)",
            "Diversité et inclusion",
            "Développement durable (ODD)",
            "Mission sociétale",
            "Mobilité internationale",
            "Aide financière",
            "Insertion professionnelle",
            "Autres",
        ],
        "external": [
            "Marché du travail et relations avec les employeurs",
            "Concurrence avec les autres universités",
            "Intelligence artificielle",
            "Réputation et image",
            "Environnement politique et économique, émigration",
            "Autres",
        ],
    },
    "Enseignants": {
        "internal": [
            "Soutenabilité financière",
            "Gouvernance et Leadership (gestion, relation, représentation, etc.)",
            "Stratégie académique et qualité d’enseignement",
            "Recherche et Innovation",
            "Ressources documentaires et Environnement digital",
            "Succès des étudiants (recrutement, accompagnement, services de support, employabilité, etc.)",
            "Ressources humaines",
            "Stratégie et mobilité internationales",
            "Mission sociétale",
            "Espace et infrastructures",
            "Environnement de travail",
            "Diversité et inclusion",
            "Développement Durable (ODD)",
            "Autre",
        ],
        "external": [
            "Marché du travail et Associations professionnelles",
            "Concurrence avec les autres universités",
            "Intelligence artificielle",
            "Réputation et image",
            "Environnement politique et économique, émigration",
            "Autres",
        ],
    },
}

PARTICIPANT_TYPE_ALIASES = {
    "psg": "PSG",
    "anciens": "Anciens",
    "directeurs": "Directeurs",
    "directors": "Directeurs",
    "etudiants": "Étudiants",
    "étudiants": "Étudiants",
    "students": "Étudiants",
    "enseignants": "Enseignants",
    "teachers": "Enseignants",
}


STOPWORDS_FR = {
    "le", "la", "les", "un", "une", "des", "du", "de", "d", "l", "et", "ou", "en",
    "au", "aux", "avec", "sans", "pour", "par", "dans", "sur", "ce", "cet", "cette",
    "ces", "qui", "que", "quoi", "dont", "est", "sont", "être", "avoir", "plus",
    "moins", "très", "tres", "comme", "nous", "vous", "ils", "elles", "il", "elle",
    "ne", "pas", "se", "sa", "son", "ses", "leur", "leurs", "a", "à", "of", "the",
    "and", "or", "to", "in", "for", "with", "by", "on",
    # Generic words that must not make two answers look like duplicates
    "renforce", "renforcee", "renforcer", "davantage", "ameliorer", "meilleur",
    "meilleure", "systeme", "niveau", "mise", "jour", "usj", "universite"
}


def canonical_participant_type(value):
    raw = clean(value)
    norm = normalize(raw)
    return PARTICIPANT_TYPE_ALIASES.get(norm, raw)


def official_themes_for(participant_type, swot_element, include_autre=False):
    participant_type = canonical_participant_type(participant_type)
    family = "internal" if swot_element in ["Forces", "Faiblesses"] else "external"
    themes = THEME_TAXONOMY.get(participant_type, {}).get(family, [])
    if include_autre:
        return themes
    return [t for t in themes if normalize(t) not in ["autre", "autres"]]


def detect_swot_element(row):
    section_n = normalize(row.get("section", ""))
    category_n = normalize(row.get("category", ""))

    if "faiblesse" in section_n or "faiblesse" in category_n:
        return "Faiblesses"
    if "force" in section_n or "force" in category_n:
        return "Forces"
    if "menace" in section_n or "menace" in category_n:
        return "Menaces"
    if "opportun" in section_n or "opportun" in category_n:
        return "Opportunités"
    return ""


def tokenize_for_similarity(text):
    text = normalize(text)
    tokens = re.findall(r"[a-z0-9]{3,}", text)
    return [t for t in tokens if t not in STOPWORDS_FR]


THEME_KEYWORD_RULES = {
    "soutenabilite financiere": [
        "budget", "financement", "financier", "financiere", "finance", "frais", "scolarite", "bourse", "cout", "couts", "ressources financieres", "modele economique", "soutenabilite", "viabilite"
    ],
    "gouvernance et leadership": [
        "gouvernance", "leadership", "direction", "doyen", "doyens", "directeur", "directeurs", "decision", "decisions", "pilotage", "procedure", "procedures", "administration", "centralisation", "decentralisation", "strategie", "representation", "responsabilisation", "accountability", "transparence", "autonomie"
    ],
    "strategie academique et qualite d enseignement": [
        "programme", "programmes", "formation", "formations", "enseignement", "academique", "pedagogie", "pedagogique", "curriculum", "cours", "accreditation", "accreditations", "diplome", "diplomes", "qualite", "evaluation", "apprentissage", "langue", "anglais", "francais", "trilinguisme"
    ],
    "qualite de l enseignement": [
        "enseignement", "cours", "enseignant", "enseignants", "pedagogie", "pedagogique", "qualite", "accompagnement", "orientation", "tutorat", "apprentissage", "evaluation", "formation", "langue", "anglais", "francais"
    ],
    "recherche et innovation": [
        "recherche", "innovation", "innovant", "publications", "publication", "scopus", "laboratoire", "laboratoires", "projet scientifique", "projets scientifiques", "these", "doctorat", "brevet", "brevets", "valorisation"
    ],
    "ressources documentaires et environnement digital": [
        "bibliotheque", "documentaire", "documentaires", "ressources numeriques", "digital", "digitale", "numerique", "numeriques", "plateforme", "plateformes", "informatique", "informatiques", "technologie", "technologique", "systeme", "systemes", "outils", "logiciel", "logiciels", "data", "donnees"
    ],
    "ressources documentaires numeriques et informatiques": [
        "bibliotheque", "documentaire", "documentaires", "ressources numeriques", "digital", "digitale", "numerique", "numeriques", "plateforme", "plateformes", "informatique", "informatiques", "technologie", "technologique", "systeme", "systemes", "outils", "logiciel", "logiciels", "data", "donnees"
    ],
    "succes des etudiants": [
        "etudiant", "etudiants", "recrutement", "inscription", "inscriptions", "accompagnement", "orientation", "support", "service", "services", "employabilite", "insertion", "stage", "stages", "reussite", "parcours", "diplomes", "alumni"
    ],
    "ressources humaines": [
        "personnel", "personnels", "enseignant", "enseignants", "administratif", "administratifs", "rh", "recrutement", "charge", "charges", "competence", "competences", "formation du personnel", "evaluation", "modernisation du personnel", "ressources humaines"
    ],
    "strategie et mobilite internationales": [
        "international", "internationale", "internationales", "mobilite", "echange", "echanges", "partenariat", "partenariats", "etranger", "etrangers", "diaspora", "francophone", "francophonie", "anglais", "langues"
    ],
    "mission societale": [
        "societe", "societal", "societale", "communautaire", "communaute", "social", "sociale", "service", "citoyen", "citoyenne", "engagement", "impact", "liban", "reconstruction", "responsabilite sociale"
    ],
    "espace et infrastructures": [
        "campus", "batiment", "batiments", "salle", "salles", "locaux", "infrastructure", "infrastructures", "espace", "espaces", "equipement", "equipements", "laboratoire", "laboratoires", "maintenance", "beyrouth"
    ],
    "espace et infrastructure": [
        "campus", "batiment", "batiments", "salle", "salles", "locaux", "infrastructure", "infrastructures", "espace", "espaces", "equipement", "equipements", "laboratoire", "laboratoires", "maintenance", "beyrouth"
    ],
    "environnement de travail": [
        "environnement de travail", "conditions de travail", "climat", "bien etre", "collaboration", "communication interne", "charge de travail", "horaires", "processus administratif", "processus administratifs"
    ],
    "diversite et inclusion": [
        "diversite", "inclusion", "inclusive", "egalite", "equite", "handicap", "genre", "discrimination", "accessibilite", "individualiste", "individualistes"
    ],
    "developpement durable": [
        "developpement durable", "odd", "durable", "durabilite", "environnement", "ecologie", "transition energetique", "energie", "responsabilite environnementale"
    ],
    "relation avec l administration et les enseignants": [
        "administration", "enseignants", "enseignant", "relation", "communication", "disponibilite", "ecoute", "procedures", "bureau", "service administratif"
    ],
    "representation des etudiants": [
        "representation", "representants", "delegue", "delegues", "voix etudiante", "participation", "consultation", "conseil etudiant"
    ],
    "vie universitaire": [
        "vie universitaire", "vie etudiante", "sport", "activites", "clubs", "campus", "aide psychologique", "psychologique", "animation", "experience etudiante"
    ],
    "mobilite internationale": [
        "mobilite", "international", "internationale", "echange", "echanges", "erasmus", "partenariat", "partenariats", "etranger", "voyage"
    ],
    "aide financiere": [
        "aide financiere", "bourse", "bourses", "frais", "scolarite", "paiement", "soutien financier", "financement"
    ],
    "insertion professionnelle": [
        "insertion", "professionnelle", "emploi", "employabilite", "stage", "stages", "carriere", "entreprise", "entreprises", "marche du travail", "metier", "metiers"
    ],
    "marche du travail et associations professionnelles": [
        "marche du travail", "emploi", "employeur", "employeurs", "entreprise", "entreprises", "association", "associations", "professionnel", "professionnelle", "professionnelles", "metier", "metiers", "competences", "stage", "stages"
    ],
    "marche du travail et relations avec les employeurs": [
        "marche du travail", "emploi", "employeur", "employeurs", "entreprise", "entreprises", "relation employeurs", "professionnel", "professionnelle", "competences", "stage", "stages", "carriere"
    ],
    "concurrence avec les autres universites": [
        "concurrence", "concurrent", "concurrentes", "universites", "autres universites", "competition", "competitive", "attractivite", "frais", "scolarite", "programmes flexibles", "facilites financieres", "classe a", "americaines"
    ],
    "intelligence artificielle": [
        "intelligence artificielle", "ia", "ai", "chatgpt", "automatisation", "technologie generative", "generative", "data science", "analyse", "digitalisation", "transformation numerique"
    ],
    "reputation et image": [
        "reputation", "image", "classement", "classements", "visibilite", "notoriete", "communication", "attractivite", "perception", "rayonnement", "jeunes", "ecoles", "recteurs", "recruteurs"
    ],
    "environnement politique et economique emigration": [
        "politique", "economique", "economie", "crise", "emigration", "emigrer", "instabilite", "instabilites", "liban", "inflation", "securite", "guerre", "familles", "parents", "banque", "bancaire", "immobilier", "financiere"
    ],
}


def strong_keyword_score(answer, theme):
    answer_norm = normalize(answer)
    theme_norm = normalize(theme)
    tokens = set(tokenize_for_similarity(answer))
    best_words = []
    for key, words in THEME_KEYWORD_RULES.items():
        if key in theme_norm or theme_norm in key:
            best_words = words
            break
    if not best_words:
        return 0.0
    hits = 0
    strong_hits = 0
    for w in best_words:
        nw = normalize(w)
        if " " in nw:
            if nw in answer_norm:
                hits += 2
                strong_hits += 1
        else:
            if nw in tokens or nw in answer_norm:
                hits += 1
    # A single clear domain keyword should matter, but multiple hits should dominate.
    return min(1.0, (hits / 5.0) + (0.12 * strong_hits))


def simple_similarity_score(answer, theme):
    answer_tokens = set(tokenize_for_similarity(answer))
    theme_tokens = set(tokenize_for_similarity(theme))

    if not answer_tokens or not theme_tokens:
        return 0.0

    overlap = len(answer_tokens & theme_tokens) / max(len(theme_tokens), 1)
    jaccard = len(answer_tokens & theme_tokens) / max(len(answer_tokens | theme_tokens), 1)

    answer_norm = normalize(answer)
    theme_norm = normalize(theme)
    phrase_bonus = 0.30 if theme_norm in answer_norm else 0.0
    keyword_score = strong_keyword_score(answer, theme)

    # Conservative but stronger institutional classifier:
    # keyword/domain evidence is prioritized over generic sentence similarity.
    return min(1.0, (0.55 * keyword_score) + (0.20 * overlap) + (0.15 * jaccard) + phrase_bonus)


def classify_theme(answer, participant_type, swot_element):
    themes = official_themes_for(participant_type, swot_element, include_autre=False)
    if not themes:
        return "", 0.0

    scores = [(theme, simple_similarity_score(answer, theme)) for theme in themes]
    scores = sorted(scores, key=lambda x: x[1], reverse=True)
    best_theme, best_score = scores[0]

    if best_score <= 0:
        return themes[0], 0.0

    return best_theme, round(best_score * 100, 1)


def read_corrected_excel(uploaded_file):
    try:
        df = pd.read_excel(uploaded_file, sheet_name=SHEET_NAME, engine="openpyxl")
    except Exception as e:
        st.error(f"Impossible de lire la feuille : {SHEET_NAME}")
        st.exception(e)
        st.stop()

    df = df.dropna(how="all")

    type_col = get_col(df, ["Respondent_Type", "Type participant", "Type de participants"])
    group_col = get_col(df, ["groupe", "Sous groupe", "Subgroup"])
    names_col = get_col(df, ["participants", "Participants", "Nom participants"])
    section_col = get_col(df, ["section"])
    category_col = get_col(df, ["category", "catégorie", "categorie"])
    answer_col = get_col(df, ["Final_Answer", "Final Answer", "Réponse finale", "Reponse finale"])

    missing = []
    if not type_col:
        missing.append("Respondent_Type")
    if not group_col:
        missing.append("groupe")
    if not names_col:
        missing.append("participants")
    if not section_col:
        missing.append("section")
    if not category_col:
        missing.append("category")
    if not answer_col:
        missing.append("Final_Answer")

    if missing:
        st.error("Colonnes manquantes : " + ", ".join(missing))
        st.stop()

    df = df.rename(columns={
        type_col: "Respondent_Type",
        group_col: "groupe",
        names_col: "participants",
        section_col: "section",
        category_col: "category",
        answer_col: "Final_Answer"
    })

    df["question"] = df["category"]

    for col in ["Respondent_Type", "groupe", "participants", "section", "category", "question", "Final_Answer"]:
        df[col] = df[col].apply(clean)

    df = df[(df["Respondent_Type"] != "") & (df["groupe"] != "")]
    df["section_norm"] = df["section"].apply(normalize)
    df["question_norm"] = df["question"].apply(normalize)
    df["category_norm"] = df["category"].apply(normalize)
    df["swot_element"] = df.apply(detect_swot_element, axis=1)

    return df


def canonical_duplicate_tokens(text):
    """Tokens for duplicate detection.
    This is intentionally stricter than general semantic similarity:
    two answers are duplicates only when they share the same core topic/object,
    not only the same sentence structure.
    """
    tokens = tokenize_for_similarity(text)
    canonical = []
    synonym_map = {
        "alumni": "anciens",
        "ancien": "anciens",
        "anciens": "anciens",
        "diplome": "diplomes",
        "diplomes": "diplomes",
        "diplômé": "diplomes",
        "diplômés": "diplomes",
        "enseignant": "enseignants",
        "enseignants": "enseignants",
        "professeur": "enseignants",
        "professeurs": "enseignants",
        "personnel": "personnel",
        "personnels": "personnel",
        "administratif": "administratifs",
        "administratifs": "administratifs",
        "gouvernance": "gouvernance",
        "integration": "integration",
        "intégration": "integration",
        "ia": "intelligence_artificielle",
        "ai": "intelligence_artificielle",
        "intelligence": "intelligence_artificielle",
        "artificielle": "intelligence_artificielle",
        "technologie": "technologie",
        "technologies": "technologie",
        "anglais": "anglais",
        "emploi": "emploi",
        "employabilite": "emploi",
        "employabilité": "emploi",
        "stage": "stage",
        "stages": "stage",
    }
    for t in tokens:
        canonical.append(synonym_map.get(t, t))
    return [t for t in canonical if t not in STOPWORDS_FR]


def strict_duplicate_score(answer_a, answer_b):
    a_norm = normalize(answer_a)
    b_norm = normalize(answer_b)

    if not a_norm or not b_norm:
        return 0.0, "Vide"

    if a_norm == b_norm:
        return 1.0, "Doublon exact"

    seq_ratio = difflib.SequenceMatcher(None, a_norm, b_norm).ratio()

    tokens_a = set(canonical_duplicate_tokens(answer_a))
    tokens_b = set(canonical_duplicate_tokens(answer_b))

    if not tokens_a or not tokens_b:
        return seq_ratio, "Texte trop court"

    shared = tokens_a & tokens_b
    union = tokens_a | tokens_b
    jaccard = len(shared) / max(len(union), 1)
    containment = len(shared) / max(min(len(tokens_a), len(tokens_b)), 1)

    # Require a common core topic/object. This prevents false positives such as:
    # "anglais renforcé" vs "technologie renforcée".
    has_core_overlap = len(shared) >= 2 or (len(shared) == 1 and min(len(tokens_a), len(tokens_b)) <= 3)

    if not has_core_overlap:
        return max(seq_ratio * 0.35, jaccard), "Même structure possible, mais sujet différent"

    if containment >= 0.80 and seq_ratio >= 0.55:
        return max(containment, seq_ratio), "Même idée avec formulation différente"

    if jaccard >= 0.62 and seq_ratio >= 0.50:
        return max(jaccard, seq_ratio), "Doublon probable"

    if seq_ratio >= 0.88 and containment >= 0.55:
        return seq_ratio, "Très proche lexicalement"

    return max(jaccard, containment * 0.75, seq_ratio * 0.50), "Non retenu"


def detect_duplicate_pairs(df):
    rows = df[df["Final_Answer"].apply(clean) != ""].copy().reset_index(drop=True)
    rows = rows[rows["swot_element"].isin(["Forces", "Faiblesses", "Opportunités", "Menaces"])]
    rows = rows.reset_index(drop=True)
    pairs = []

    for i in range(len(rows)):
        for j in range(i + 1, len(rows)):
            # Compare only inside the same SWOT element. A Force and a Priority/Threat are not duplicates.
            if rows.loc[i, "swot_element"] != rows.loc[j, "swot_element"]:
                continue

            score, reason = strict_duplicate_score(
                rows.loc[i, "Final_Answer"],
                rows.loc[j, "Final_Answer"]
            )

            # Fixed strict rule. No slider.
            if score >= 0.80:
                pairs.append({
                    "Élément SWOT": rows.loc[i, "swot_element"],
                    "Sous-groupe 1": rows.loc[i, "groupe"],
                    "Réponse 1": rows.loc[i, "Final_Answer"],
                    "Sous-groupe 2": rows.loc[j, "groupe"],
                    "Réponse 2": rows.loc[j, "Final_Answer"],
                    "Score doublon": round(score * 100, 1),
                    "Raison": reason,
                    "Validation expert": "À vérifier",
                })

    return pd.DataFrame(pairs)


def build_ai_classification_dataframe(df_scope):
    rows = []
    for idx, row in df_scope.iterrows():
        swot_element = row["swot_element"]
        if swot_element not in ["Forces", "Faiblesses", "Opportunités", "Menaces"]:
            continue

        answer = clean(row["Final_Answer"])
        if not answer:
            continue

        suggested_theme, confidence = classify_theme(
            answer=answer,
            participant_type=row["Respondent_Type"],
            swot_element=swot_element
        )

        rows.append({
            "row_id": idx,
            "Respondent_Type": row["Respondent_Type"],
            "groupe": row["groupe"],
            "participants": row["participants"],
            "swot_element": swot_element,
            "question": row["question"],
            "Final_Answer": answer,
            "AI_Suggested_Theme": suggested_theme,
            "Confidence": confidence,
            "Validated_Theme": suggested_theme,
            "New_Theme": "",
        })

    return pd.DataFrame(rows)


def render_mapping_validation(classified_df):
    if classified_df.empty:
        st.warning("Aucune réponse SWOT à classifier.")
        return classified_df

    validated_rows = []

    for swot_element in ["Forces", "Faiblesses", "Opportunités", "Menaces"]:
        temp = classified_df[classified_df["swot_element"] == swot_element].copy()
        if temp.empty:
            continue

        st.markdown(f"### {swot_element}")

        for _, row in temp.iterrows():
            row_key = f"{row['row_id']}_{swot_element}"
            themes = official_themes_for(row["Respondent_Type"], swot_element, include_autre=False)
            suggested = row["AI_Suggested_Theme"]
            default_index = themes.index(suggested) if suggested in themes else 0

            with st.expander(
                f"{row['groupe']} | Confiance: {row['Confidence']}% | Thème proposé: {suggested}",
                expanded=False
            ):
                st.markdown("**Réponse corrigée**")
                st.write(row["Final_Answer"])

                selected_theme = st.selectbox(
                    "Thème validé par l'expert",
                    options=themes,
                    index=default_index,
                    key=f"theme_{row_key}"
                )

                new_theme = st.text_input(
                    "Ajouter un nouveau thème si nécessaire",
                    value="",
                    key=f"new_theme_{row_key}"
                )

                final_theme = clean(new_theme) if clean(new_theme) else selected_theme

                validated_row = row.to_dict()
                validated_row["Validated_Theme"] = final_theme
                validated_row["New_Theme"] = clean(new_theme)
                validated_rows.append(validated_row)

    return pd.DataFrame(validated_rows)


def build_theme_synthesis(validated_df):
    if validated_df.empty:
        return pd.DataFrame()

    rows = []
    grouped = validated_df.groupby(["Respondent_Type", "swot_element", "Validated_Theme"], dropna=False)

    for (ptype, swot_element, theme), group in grouped:
        answers = [clean(x) for x in group["Final_Answer"].tolist() if clean(x)]
        short_points = " | ".join(answers[:3])
        rows.append({
            "Type de participants": ptype,
            "Élément SWOT": swot_element,
            "Thème": theme,
            "Nombre de réponses": len(answers),
            "Synthèse automatique courte": short_points,
        })

    return pd.DataFrame(rows).sort_values(
        ["Type de participants", "Élément SWOT", "Nombre de réponses"],
        ascending=[True, True, False]
    )


def load_validated_mapping_file(uploaded_mapping):
    if uploaded_mapping is None:
        return pd.DataFrame()
    try:
        df_map = pd.read_excel(uploaded_mapping, engine="openpyxl")
    except Exception:
        uploaded_mapping.seek(0)
        df_map = pd.read_csv(uploaded_mapping)
    return df_map


def prepare_training_dataset(validated_df):
    if validated_df is None or validated_df.empty:
        return pd.DataFrame()

    text_col = get_col(validated_df, ["Final_Answer", "answer", "Réponse", "Reponse", "text"])
    label_col = get_col(validated_df, ["Validated_Theme", "Thème validé", "Theme valide", "theme", "label"])
    swot_col = get_col(validated_df, ["swot_element", "Élément SWOT", "Element SWOT", "SWOT"])
    type_col = get_col(validated_df, ["Respondent_Type", "Type de participants", "Type participant"])

    if not text_col or not label_col:
        return pd.DataFrame()

    out = pd.DataFrame({
        "text": validated_df[text_col].apply(clean),
        "label": validated_df[label_col].apply(clean),
    })
    out["swot_element"] = validated_df[swot_col].apply(clean) if swot_col else ""
    out["Respondent_Type"] = validated_df[type_col].apply(clean) if type_col else ""
    out = out[(out["text"] != "") & (out["label"] != "")]
    return out.reset_index(drop=True)


def train_test_theme_algorithms(train_df, test_size=0.25, random_state=42):
    try:
        from sklearn.model_selection import train_test_split
        from sklearn.pipeline import Pipeline, FeatureUnion
        from sklearn.feature_extraction.text import TfidfVectorizer
        from sklearn.linear_model import LogisticRegression, SGDClassifier
        from sklearn.naive_bayes import MultinomialNB
        from sklearn.svm import LinearSVC
        from sklearn.metrics import accuracy_score, f1_score, classification_report
    except Exception as e:
        return pd.DataFrame(), pd.DataFrame(), f"scikit-learn n'est pas installé. Ajoutez scikit-learn dans requirements.txt. Détail: {e}"

    if train_df.empty or train_df["label"].nunique() < 2:
        return pd.DataFrame(), pd.DataFrame(), "Il faut au moins deux thèmes validés différents pour entraîner/tester un modèle."

    counts = train_df["label"].value_counts()
    if len(train_df) < 8 or counts.min() < 2:
        return pd.DataFrame(), pd.DataFrame(), "Données insuffisantes: il faut idéalement au moins 2 réponses validées par thème et au moins 8 réponses au total."

    X = train_df["text"].astype(str)
    y = train_df["label"].astype(str)
    stratify = y if counts.min() >= 2 else None

    X_train, X_test, y_train, y_test = train_test_split(
        X, y, test_size=test_size, random_state=random_state, stratify=stratify
    )

    features = FeatureUnion([
        ("word", TfidfVectorizer(ngram_range=(1, 2), min_df=1, max_df=0.95)),
        ("char", TfidfVectorizer(analyzer="char_wb", ngram_range=(3, 5), min_df=1)),
    ])

    algorithms = {
        "TF-IDF + Logistic Regression": LogisticRegression(max_iter=2000, class_weight="balanced"),
        "TF-IDF + Linear SVM": LinearSVC(class_weight="balanced"),
        "TF-IDF + SGD Linear SVM": SGDClassifier(loss="hinge", class_weight="balanced", random_state=random_state),
        "TF-IDF + Naive Bayes": MultinomialNB(),
    }

    rows = []
    pred_rows = []

    for name, clf in algorithms.items():
        try:
            model = Pipeline([
                ("features", features),
                ("clf", clf),
            ])
            model.fit(X_train, y_train)
            pred = model.predict(X_test)
            acc = accuracy_score(y_test, pred)
            f1 = f1_score(y_test, pred, average="macro", zero_division=0)
            rows.append({
                "Algorithme": name,
                "Accuracy test": round(acc, 3),
                "Macro F1 test": round(f1, 3),
                "Train n": len(X_train),
                "Test n": len(X_test),
                "Nombre de thèmes": y.nunique(),
            })
            for true_label, pred_label, txt in zip(y_test, pred, X_test):
                pred_rows.append({
                    "Algorithme": name,
                    "Réponse test": txt,
                    "Thème réel": true_label,
                    "Thème prédit": pred_label,
                    "Correct": true_label == pred_label,
                })
        except Exception as e:
            rows.append({
                "Algorithme": name,
                "Accuracy test": None,
                "Macro F1 test": None,
                "Train n": len(X_train),
                "Test n": len(X_test),
                "Nombre de thèmes": y.nunique(),
                "Erreur": str(e),
            })

    results = pd.DataFrame(rows).sort_values(["Macro F1 test", "Accuracy test"], ascending=False, na_position="last")
    predictions = pd.DataFrame(pred_rows)
    return results, predictions, ""



def build_classification_with_trained_model(df_scope, model):
    rows = []
    for idx, row in df_scope.iterrows():
        swot_element = row["swot_element"]
        if swot_element not in ["Forces", "Faiblesses", "Opportunités", "Menaces"]:
            continue

        answer = clean(row["Final_Answer"])
        if not answer:
            continue

        suggested_theme = model.predict([answer])[0]
        confidence = ""

        try:
            if hasattr(model.named_steps.get("clf"), "predict_proba"):
                proba = model.predict_proba([answer])[0]
                confidence = round(float(max(proba)) * 100, 1)
            elif hasattr(model.named_steps.get("clf"), "decision_function"):
                scores = model.decision_function([answer])
                import numpy as np
                scores = np.array(scores).reshape(-1)
                if len(scores) > 0:
                    exp_scores = np.exp(scores - scores.max())
                    confidence = round(float(exp_scores.max() / exp_scores.sum()) * 100, 1)
        except Exception:
            confidence = ""

        rows.append({
            "row_id": idx,
            "Respondent_Type": row["Respondent_Type"],
            "groupe": row["groupe"],
            "participants": row["participants"],
            "swot_element": swot_element,
            "question": row["question"],
            "Final_Answer": answer,
            "AI_Suggested_Theme": suggested_theme,
            "Confidence": confidence,
            "Validated_Theme": suggested_theme,
            "New_Theme": "",
        })

    return pd.DataFrame(rows)


def train_final_theme_model(train_df, algorithm_name="TF-IDF + Linear SVM", random_state=42):
    try:
        from sklearn.pipeline import Pipeline, FeatureUnion
        from sklearn.feature_extraction.text import TfidfVectorizer
        from sklearn.linear_model import LogisticRegression, SGDClassifier
        from sklearn.naive_bayes import MultinomialNB
        from sklearn.svm import LinearSVC
    except Exception as e:
        return None, f"scikit-learn n'est pas installé. Ajoutez scikit-learn dans requirements.txt. Détail: {e}"

    if train_df.empty or train_df["label"].nunique() < 2:
        return None, "Il faut au moins deux thèmes validés différents pour entraîner le modèle final."

    features = FeatureUnion([
        ("word", TfidfVectorizer(ngram_range=(1, 2), min_df=1, max_df=0.95)),
        ("char", TfidfVectorizer(analyzer="char_wb", ngram_range=(3, 5), min_df=1)),
    ])

    algorithms = {
        "TF-IDF + Logistic Regression": LogisticRegression(max_iter=2000, class_weight="balanced"),
        "TF-IDF + Linear SVM": LinearSVC(class_weight="balanced"),
        "TF-IDF + SGD Linear SVM": SGDClassifier(loss="hinge", class_weight="balanced", random_state=random_state),
        "TF-IDF + Naive Bayes": MultinomialNB(),
    }

    clf = algorithms.get(algorithm_name, LinearSVC(class_weight="balanced"))
    model = Pipeline([("features", features), ("clf", clf)])
    model.fit(train_df["text"].astype(str), train_df["label"].astype(str))
    return model, ""



# ============================================================
# PROFESSIONAL AI ENGINE - PRETRAINED / NO EXPERT TRAINING REQUIRED
# ============================================================

THEME_CONTEXT = {
    "Soutenabilité financière": "budget financement frais scolarité ressources financières modèle économique bourses coûts viabilité financière",
    "Gouvernance et Leadership": "gouvernance leadership management décision direction pilotage organisation représentation relation procédures transparence responsabilité",
    "Stratégie académique et qualité d’enseignement": "programmes formation pédagogie qualité enseignement curriculum accréditation excellence académique méthodes cours apprentissage",
    "Recherche et Innovation": "recherche innovation publications laboratoires projets scientifiques valorisation créativité transfert technologie",
    "Ressources documentaires et Environnement digital": "bibliothèque ressources numériques plateformes digitales systèmes informatiques documentation accès digitalisation outils technologiques",
    "Ressources documentaires, numériques et informatiques": "bibliothèque ressources numériques informatique plateformes documentation accès digital technologies outils numériques",
    "Succès des étudiants": "étudiants recrutement accompagnement orientation soutien services employabilité insertion suivi réussite parcours étudiant",
    "Ressources humaines": "personnel enseignants administratifs recrutement compétences charge travail formation évaluation ressources humaines",
    "Stratégie et mobilité internationales": "international mobilité partenariats étrangers échanges coopération internationale rayonnement international",
    "Mission sociétale": "société engagement service communauté responsabilité sociale impact social mission citoyenne territoire",
    "Espace et infrastructures": "campus bâtiments salles laboratoires espaces équipements infrastructures locaux maintenance accessibilité",
    "Espace et infrastructure": "campus bâtiments salles laboratoires espaces équipements infrastructures locaux maintenance accessibilité",
    "Environnement de travail": "conditions travail climat organisation charge bien-être collaboration communication interne environnement professionnel",
    "Diversité et inclusion": "diversité inclusion égalité accessibilité handicap genre équité discrimination intégration",
    "Développement Durable": "développement durable ODD environnement écologie transition énergétique durabilité responsabilité environnementale",
    "Qualité de l’enseignement": "qualité enseignement accompagnement orientation tutorat pédagogie cours enseignants apprentissage évaluation suivi académique",
    "Relation avec l’administration et les enseignants": "administration enseignants relation communication disponibilité écoute procédures services étudiants",
    "Représentation des étudiants": "représentation étudiants délégués participation gouvernance voix étudiante consultation implication",
    "Vie universitaire": "vie étudiante sport activités campus aide psychologique animation clubs expérience étudiante",
    "Mobilité internationale": "mobilité internationale échanges Erasmus partenariats étudiants étrangers voyage coopération",
    "Aide financière": "aide financière bourses frais scolarité soutien social paiement financement étudiants",
    "Insertion professionnelle": "emploi stage employabilité marché travail carrière entreprises compétences professionnelles insertion",
    "Marché du travail et Associations professionnelles": "marché travail employeurs entreprises associations professionnelles métiers compétences emploi stages secteur professionnel",
    "Marché du travail et relations avec les employeurs": "marché travail employeurs entreprises emploi stages insertion besoins compétences relation professionnelle",
    "Concurrence avec les autres universités": "concurrence universités compétition établissements attractivité inscriptions frais programmes alternatives",
    "Intelligence artificielle": "intelligence artificielle IA AI ChatGPT automatisation technologie transformation numérique outils génératifs",
    "Réputation et image": "réputation image visibilité marque notoriété classement communication attractivité perception rayonnement",
    "Environnement politique et économique, émigration": "politique économie crise émigration instabilité contexte pays inflation sécurité situation nationale",
}


def enrich_theme_text(theme):
    theme_clean = clean(theme)
    norm_theme = normalize(theme_clean)
    context = ""
    for key, value in THEME_CONTEXT.items():
        if normalize(key) in norm_theme or norm_theme in normalize(key):
            context = value
            break
    return f"{theme_clean}. {context}".strip()


@st.cache_resource(show_spinner=False)
def load_multilingual_sentence_model():
    try:
        from sentence_transformers import SentenceTransformer
        return SentenceTransformer("paraphrase-multilingual-MiniLM-L12-v2")
    except Exception:
        return None


def cosine_from_vectors(a, b):
    import numpy as np
    a = np.asarray(a)
    b = np.asarray(b)
    denom = (np.linalg.norm(a) * np.linalg.norm(b))
    if denom == 0:
        return 0.0
    return float(np.dot(a, b) / denom)


def semantic_theme_scores(answer, participant_type, swot_element, engine_name):
    themes = official_themes_for(participant_type, swot_element, include_autre=False)
    if not themes:
        return []

    answer_text = clean(answer)
    theme_texts = [enrich_theme_text(t) for t in themes]

    if "Sentence Transformer" in engine_name:
        model = load_multilingual_sentence_model()
        if model is not None:
            try:
                embeddings = model.encode([answer_text] + theme_texts, normalize_embeddings=True)
                answer_vec = embeddings[0]
                theme_vecs = embeddings[1:]
                scores = [cosine_from_vectors(answer_vec, v) for v in theme_vecs]
                return sorted(zip(themes, scores), key=lambda x: x[1], reverse=True)
            except Exception:
                pass

    # Offline professional fallback: hybrid word + char TF-IDF + theme context.
    try:
        from sklearn.feature_extraction.text import TfidfVectorizer
        from sklearn.metrics.pairwise import cosine_similarity
        from sklearn.pipeline import FeatureUnion

        corpus = [answer_text] + theme_texts
        features = FeatureUnion([
            ("word", TfidfVectorizer(ngram_range=(1, 2), min_df=1, strip_accents="unicode")),
            ("char", TfidfVectorizer(analyzer="char_wb", ngram_range=(3, 5), min_df=1, strip_accents="unicode")),
        ])
        matrix = features.fit_transform(corpus)
        sims = cosine_similarity(matrix[0:1], matrix[1:]).flatten().tolist()
        keyword_scores = [simple_similarity_score(answer_text, t) for t in themes]
        # Use lexical/domain evidence as the anchor; semantic similarity is only supportive.
        scores = [(0.45 * keyword_scores[i]) + (0.35 * sims[i]) + (0.20 * strong_keyword_score(answer_text, themes[i])) for i in range(len(themes))]
        return sorted(zip(themes, scores), key=lambda x: x[1], reverse=True)
    except Exception:
        scores = [simple_similarity_score(answer_text, t) for t in themes]
        return sorted(zip(themes, scores), key=lambda x: x[1], reverse=True)


def classify_theme_with_engine(answer, participant_type, swot_element, engine_name):
    scores = semantic_theme_scores(answer, participant_type, swot_element, engine_name)
    if not scores:
        return "", 0.0, ""

    best_theme, best_score = scores[0]
    second_score = scores[1][1] if len(scores) > 1 else 0
    margin = max(best_score - second_score, 0)

    # Calibrated confidence: high only when the best theme is clearly separated.
    confidence = min(99.0, max(5.0, 35 + (best_score * 55) + (margin * 65)))
    if margin < 0.025:
        confidence = min(confidence, 58.0)
    elif margin < 0.06:
        confidence = min(confidence, 72.0)
    top3 = " | ".join([f"{t} ({round(s * 100, 1)})" for t, s in scores[:3]])
    return best_theme, round(confidence, 1), top3


def build_classification_with_engine(df_scope, engine_name):
    rows = []
    for idx, row in df_scope.iterrows():
        swot_element = row["swot_element"]
        if swot_element not in ["Forces", "Faiblesses", "Opportunités", "Menaces"]:
            continue

        answer = clean(row["Final_Answer"])
        if not answer:
            continue

        suggested_theme, confidence, top3 = classify_theme_with_engine(
            answer=answer,
            participant_type=row["Respondent_Type"],
            swot_element=swot_element,
            engine_name=engine_name,
        )

        rows.append({
            "row_id": idx,
            "Respondent_Type": row["Respondent_Type"],
            "groupe": row["groupe"],
            "participants": row["participants"],
            "swot_element": swot_element,
            "question": row["question"],
            "Final_Answer": answer,
            "AI_Suggested_Theme": suggested_theme,
            "Confidence": confidence,
            "Top_3_Themes": top3,
            "Validated_Theme": suggested_theme,
            "New_Theme": "",
        })
    return pd.DataFrame(rows)




def available_ai_engines():
    engines = []
    if load_multilingual_sentence_model() is not None:
        engines.append("Sentence Transformer multilingue pré-entraîné - recommandé")
    engines.extend([
        "Hybrid semantic TF-IDF offline - fallback robuste",
        "Keyword-assisted TF-IDF - contrôle lexical",
    ])
    return engines


def compare_ai_engines(df_scope):
    engines = available_ai_engines()
    engine_outputs = []

    for engine in engines:
        temp = build_classification_with_engine(df_scope, engine)
        if temp.empty:
            continue
        temp = temp[[
            "row_id", "Respondent_Type", "groupe", "participants", "swot_element",
            "question", "Final_Answer", "AI_Suggested_Theme", "Confidence", "Top_3_Themes"
        ]].copy()
        temp = temp.rename(columns={
            "AI_Suggested_Theme": f"Theme_{safe_filename(engine)}",
            "Confidence": f"Confidence_{safe_filename(engine)}",
            "Top_3_Themes": f"Top3_{safe_filename(engine)}",
        })
        engine_outputs.append((engine, temp))

    if not engine_outputs:
        return pd.DataFrame(), pd.DataFrame(), []

    comparison = engine_outputs[0][1]
    for _, temp in engine_outputs[1:]:
        comparison = comparison.merge(
            temp.drop(columns=["Respondent_Type", "groupe", "participants", "swot_element", "question", "Final_Answer"]),
            on="row_id",
            how="outer"
        )

    summary_rows = []
    for engine, _ in engine_outputs:
        theme_col = f"Theme_{safe_filename(engine)}"
        conf_col = f"Confidence_{safe_filename(engine)}"
        summary_rows.append({
            "Moteur": engine,
            "Réponses classifiées": comparison[theme_col].notna().sum() if theme_col in comparison else 0,
            "Confiance moyenne": round(pd.to_numeric(comparison[conf_col], errors="coerce").mean(), 1) if conf_col in comparison else None,
            "Confiance médiane": round(pd.to_numeric(comparison[conf_col], errors="coerce").median(), 1) if conf_col in comparison else None,
            "Faible confiance (<60%)": int((pd.to_numeric(comparison[conf_col], errors="coerce") < 60).sum()) if conf_col in comparison else 0,
        })

    theme_cols = [f"Theme_{safe_filename(engine)}" for engine, _ in engine_outputs]
    if len(theme_cols) >= 2:
        def agreement(row):
            vals = [clean(row.get(c, "")) for c in theme_cols if clean(row.get(c, ""))]
            return "Accord" if vals and len(set(vals)) == 1 else "Désaccord"
        comparison["Accord_entre_moteurs"] = comparison.apply(agreement, axis=1)
    else:
        comparison["Accord_entre_moteurs"] = "Un seul moteur"

    summary = pd.DataFrame(summary_rows)
    if len(theme_cols) >= 2:
        agreement_rate = round((comparison["Accord_entre_moteurs"] == "Accord").mean() * 100, 1) if len(comparison) else 0
        summary.loc[len(summary)] = {
            "Moteur": "Accord global entre moteurs",
            "Réponses classifiées": len(comparison),
            "Confiance moyenne": agreement_rate,
            "Confiance médiane": None,
            "Faible confiance (<60%)": None,
        }
    return comparison, summary, engines


def evaluate_engines_against_validated(comparison_df, validated_file):
    if validated_file is None or comparison_df.empty:
        return pd.DataFrame()
    try:
        if validated_file.name.lower().endswith(".csv"):
            valid = pd.read_csv(validated_file)
        else:
            valid = pd.read_excel(validated_file)
    except Exception as e:
        st.error("Impossible de lire le fichier validé.")
        st.exception(e)
        return pd.DataFrame()

    possible_label_cols = ["Validated_Theme", "Theme_validé", "Theme valide", "Thème validé", "Final_Theme", "theme"]
    label_col = None
    for c in possible_label_cols:
        if c in valid.columns:
            label_col = c
            break
    if label_col is None:
        st.error("Le fichier validé doit contenir une colonne Validated_Theme ou Thème validé.")
        return pd.DataFrame()

    if "row_id" in valid.columns:
        merged = comparison_df.merge(valid[["row_id", label_col]], on="row_id", how="inner")
    else:
        key_cols = [c for c in ["Respondent_Type", "groupe", "swot_element", "Final_Answer"] if c in valid.columns and c in comparison_df.columns]
        if not key_cols:
            st.error("Le fichier validé doit contenir row_id, ou les colonnes Respondent_Type, groupe, swot_element et Final_Answer.")
            return pd.DataFrame()
        merged = comparison_df.merge(valid[key_cols + [label_col]], on=key_cols, how="inner")

    results = []
    true_labels = merged[label_col].apply(clean)
    for col in [c for c in merged.columns if c.startswith("Theme_")]:
        pred = merged[col].apply(clean)
        ok = pred == true_labels
        results.append({
            "Moteur": col.replace("Theme_", "").replace("_", " "),
            "N validé": len(merged),
            "Accuracy": round(ok.mean() * 100, 1) if len(merged) else 0,
            "Erreurs": int((~ok).sum()),
        })
    return pd.DataFrame(results)

def render_model_management(df_scope, selected_type="Tous"):
    st.markdown("""
    <div class="ai-card pro-card">
        <h3>1. Sélection du moteur IA</h3>
        <p>Le système utilise un moteur linguistique pré-entraîné. L'expert ne doit pas entraîner le modèle manuellement. Les thèmes officiels restent inchangés; le modèle compare chaque réponse aux thèmes autorisés pour le type de participants et l'élément SWOT.</p>
    </div>
    """, unsafe_allow_html=True)

    transformer_available = load_multilingual_sentence_model() is not None
    engine_options = available_ai_engines()

    c1, c2 = st.columns([1.4, 1])
    with c1:
        engine_name = st.selectbox("Moteur de classification", engine_options, key="ai_engine_name")
    with c2:
        st.metric("Moteur transformer disponible", "Oui" if transformer_available else "Non")

    swot_rows = df_scope[df_scope["swot_element"].isin(["Forces", "Faiblesses", "Opportunités", "Menaces"])]
    available_theme_count = 0
    for ptype in sorted(df_scope["Respondent_Type"].dropna().unique().tolist()):
        for el in ["Forces", "Faiblesses", "Opportunités", "Menaces"]:
            available_theme_count += len(official_themes_for(ptype, el, include_autre=False))

    k1, k2, k3 = st.columns(3)
    k1.metric("Réponses SWOT", len(swot_rows))
    k2.metric("Thèmes officiels disponibles", available_theme_count)
    k3.metric("Mode", "Pré-entraîné")

    st.markdown("""
    <div class="ai-info-box">
        <b>Important :</b> ce module ne prétend pas créer un modèle supervisé à partir de zéro. Il exploite un modèle linguistique pré-entraîné sur de grands corpus multilingues, puis l'adapte à votre taxonomie officielle par similarité sémantique contrôlée. La validation humaine reste disponible pour corriger les cas ambigus et ajouter un nouveau thème lorsque nécessaire.
    </div>
    """, unsafe_allow_html=True)

    if st.button("Activer ce moteur IA", key="activate_ai_engine"):
        st.session_state["active_ai_engine"] = engine_name
        st.success(f"Moteur IA actif : {engine_name}")

    active = st.session_state.get("active_ai_engine", "")
    if active:
        st.success(f"Moteur actuellement utilisé : {active}")
    else:
        st.warning("Veuillez activer un moteur IA avant de lancer le mapping, les doublons ou la synthèse.")


def ai_platform_page(show_header=True, df_preloaded=None):
    html_block(APP_CSS)

    html_block(f"""
    <style>
    .ai-hero {{
        background: linear-gradient(135deg, #001F5B 0%, #123C7C 55%, #8B1538 100%);
        color: white;
        border-radius: 22px;
        padding: 30px 34px;
        margin: 8px 0 24px 0;
        box-shadow: 0 16px 38px rgba(0,31,91,0.22);
    }}
    .ai-hero h2 {{
        color: white !important;
        margin: 0 0 8px 0;
        font-size: 32px;
        font-weight: 900;
        letter-spacing: 0.2px;
    }}
    .ai-hero p {{
        color: #EAF2F8;
        margin: 0;
        font-size: 16px;
        line-height: 1.5;
    }}
    .ai-card {{
        background: white;
        border: 1px solid #D0D6E0;
        border-radius: 18px;
        padding: 18px 22px;
        margin: 14px 0 18px 0;
        box-shadow: 0 8px 24px rgba(0,0,0,0.055);
    }}
    .pro-card h3 {{
        color: {USJ_BLUE} !important;
        font-size: 24px;
        margin: 0 0 8px 0;
    }}
    .ai-info-box {{
        background: #EAF2F8;
        color: #003B7A;
        border-left: 6px solid {USJ_BLUE};
        border-radius: 12px;
        padding: 14px 18px;
        margin: 14px 0;
        font-size: 15px;
    }}
    .ai-warning {{
        background: #FFF7E6;
        border-left: 6px solid #F59F00;
        border-radius: 12px;
        padding: 14px 18px;
        color: #7A4B00;
        font-weight: 700;
        margin: 12px 0;
    }}
    .ai-kpi {{
        background: white;
        border: 1px solid #D0D6E0;
        border-radius: 16px;
        padding: 16px 18px;
        box-shadow: 0 8px 20px rgba(0,0,0,0.045);
    }}
    .ai-kpi-number {{
        color: {USJ_RED};
        font-size: 28px;
        font-weight: 900;
    }}
    .ai-kpi-label {{
        color: #64748B;
        font-size: 13px;
        font-weight: 700;
        margin-top: 3px;
    }}
    </style>
    <div class="ai-hero">
        <h2>Plateforme IA supervisée - Analyse thématique stratégique</h2>
        <p>Workflow institutionnel : moteur IA pré-entraîné, mapping contrôlé sur les thèmes officiels, validation humaine, détection stricte des doublons et synthèse par thème.</p>
    </div>
    """)

    if df_preloaded is None:
        uploaded_file = st.file_uploader(
            "Uploader le fichier Excel corrigé pour l'analyse IA",
            type=["xlsx", "xls"],
            key="ai_excel_upload"
        )
        if uploaded_file is None:
            st.stop()
        df = read_corrected_excel(uploaded_file)
    else:
        df = df_preloaded.copy()

    # IA module works on the full uploaded dataset, not on a selected subset.
    selected_type = "Tous les types de participants"
    selected_subgroup = "Tous les sous-groupes"
    df_scope = df.copy()

    st.markdown("""
    <div class="ai-info-box">
        <b>Périmètre IA :</b> l'analyse IA est appliquée à toutes les réponses importées, tous types de participants et tous sous-groupes confondus. La taxonomie officielle reste adaptée ligne par ligne selon le type de participant et l'élément SWOT.
    </div>
    """, unsafe_allow_html=True)

    swot_count = len(df_scope[df_scope["swot_element"].isin(["Forces", "Faiblesses", "Opportunités", "Menaces"])])
    k1, k2, k3 = st.columns(3)
    with k1:
        st.markdown(f'<div class="ai-kpi"><div class="ai-kpi-number">{len(df_scope)}</div><div class="ai-kpi-label">Réponses dans le périmètre</div></div>', unsafe_allow_html=True)
    with k2:
        st.markdown(f'<div class="ai-kpi"><div class="ai-kpi-number">{swot_count}</div><div class="ai-kpi-label">Réponses SWOT analysables</div></div>', unsafe_allow_html=True)
    with k3:
        active_engine = st.session_state.get("active_ai_engine", "Non")
        st.markdown(f'<div class="ai-kpi"><div class="ai-kpi-number">{"Oui" if active_engine != "Non" else "Non"}</div><div class="ai-kpi-label">Moteur IA actif</div></div>', unsafe_allow_html=True)

    tab_model, tab_compare, tab_mapping, tab_duplicates, tab_synthesis = st.tabs(
        ["1. Moteur IA global", "2. Comparaison moteurs", "3. Mapping SWOT global", "4. Doublons globaux", "5. Synthèse globale"]
    )

    with tab_model:
        render_model_management(df_scope, selected_type)

    with tab_compare:
        st.subheader("2. Comparaison des moteurs IA sur toutes les réponses")
        st.caption("Cette étape applique tous les moteurs disponibles au même fichier Excel. Sans fichier validé par expert, elle mesure l'accord entre moteurs et les niveaux de confiance. Avec un fichier validé, elle calcule l'accuracy réelle.")

        if st.button("Comparer les moteurs IA", key="compare_ai_engines"):
            comparison_df, summary_df, engines = compare_ai_engines(df_scope)
            st.session_state["engine_comparison_df"] = comparison_df
            st.session_state["engine_summary_df"] = summary_df

        summary_df = st.session_state.get("engine_summary_df", pd.DataFrame())
        comparison_df = st.session_state.get("engine_comparison_df", pd.DataFrame())

        if not summary_df.empty:
            st.markdown("#### Synthèse comparative")
            st.dataframe(summary_df, use_container_width=True)

        if not comparison_df.empty:
            st.markdown("#### Détail des prédictions par moteur")
            display_cols = [c for c in comparison_df.columns if c in ["Respondent_Type", "groupe", "swot_element", "Final_Answer", "Accord_entre_moteurs"] or c.startswith("Theme_") or c.startswith("Confidence_")]
            st.dataframe(comparison_df[display_cols], use_container_width=True)

            output = BytesIO()
            with pd.ExcelWriter(output, engine="openpyxl") as writer:
                comparison_df.to_excel(writer, index=False, sheet_name="Comparaison moteurs")
                summary_df.to_excel(writer, index=False, sheet_name="Synthese")
            output.seek(0)
            st.download_button(
                "Télécharger la comparaison des moteurs",
                data=output.getvalue(),
                file_name="comparaison_moteurs_IA.xlsx",
                mime="application/vnd.openxmlformats-officedocument.spreadsheetml.sheet"
            )

            st.markdown("#### Évaluation avec un fichier validé par expert")
            validated_file = st.file_uploader(
                "Uploader un fichier validé contenant row_id et Validated_Theme",
                type=["xlsx", "csv"],
                key="validated_engine_eval_upload"
            )
            if validated_file is not None:
                eval_df = evaluate_engines_against_validated(comparison_df, validated_file)
                if not eval_df.empty:
                    st.dataframe(eval_df, use_container_width=True)
        else:
            st.caption("Cliquez sur « Comparer les moteurs IA » pour produire la comparaison.")

    with tab_mapping:
        st.subheader("3. Mapping SWOT supervisé - toutes les données")
        engine = st.session_state.get("active_ai_engine", "")
        if not engine:
            st.markdown('<div class="ai-warning">Activez d’abord un moteur IA dans l’onglet 1.</div>', unsafe_allow_html=True)
        else:
            st.caption(f"Moteur actif : {engine}")
            if st.button("Classifier les réponses SWOT", key="run_ai_mapping_engine"):
                st.session_state["classified_df"] = build_classification_with_engine(df_scope, engine)

            classified_df = st.session_state.get("classified_df", pd.DataFrame())
            if not classified_df.empty:
                preview_cols = ["groupe", "swot_element", "Final_Answer", "AI_Suggested_Theme", "Confidence", "Top_3_Themes"]
                st.dataframe(classified_df[preview_cols], use_container_width=True)

                validated_df = render_mapping_validation(classified_df)
                st.session_state["validated_mapping_df"] = validated_df

                if not validated_df.empty:
                    export_cols = [
                        "Respondent_Type", "groupe", "participants", "swot_element", "question",
                        "Final_Answer", "AI_Suggested_Theme", "Confidence", "Top_3_Themes",
                        "Validated_Theme", "New_Theme"
                    ]
                    export_df = validated_df[export_cols].copy()
                    output = BytesIO()
                    with pd.ExcelWriter(output, engine="openpyxl") as writer:
                        export_df.to_excel(writer, index=False, sheet_name="Validated Mapping")
                    output.seek(0)
                    st.download_button(
                        "Télécharger le mapping validé",
                        data=output.getvalue(),
                        file_name=f"mapping_valide_{safe_filename(selected_type)}_{safe_filename(selected_subgroup)}.xlsx",
                        mime="application/vnd.openxmlformats-officedocument.spreadsheetml.sheet"
                    )
            else:
                st.caption("Cliquez sur « Classifier les réponses SWOT » pour lancer le mapping.")

    with tab_duplicates:
        st.subheader("4. Détection stricte des doublons - toutes les données")
        engine = st.session_state.get("active_ai_engine", "")
        if not engine:
            st.markdown('<div class="ai-warning">Activez d’abord un moteur IA dans l’onglet 1.</div>', unsafe_allow_html=True)
        else:
            st.caption("Règle stricte : même élément SWOT + même sujet central. Une structure grammaticale similaire avec sujet différent n’est pas retenue.")
            if st.button("Détecter les doublons", key="detect_duplicates"):
                dup_df = detect_duplicate_pairs(df_scope)
                st.session_state["duplicates_df"] = dup_df

            dup_df = st.session_state.get("duplicates_df", pd.DataFrame())
            if not dup_df.empty:
                st.dataframe(dup_df, use_container_width=True)
                st.download_button(
                    "Télécharger les doublons détectés",
                    data=dup_df.to_csv(index=False).encode("utf-8-sig"),
                    file_name=f"doublons_{safe_filename(selected_type)}.csv",
                    mime="text/csv"
                )
            else:
                st.caption("Aucun doublon affiché pour le moment.")

    with tab_synthesis:
        st.subheader("5. Synthèse automatique par thème validé")
        validated_df = st.session_state.get("validated_mapping_df", pd.DataFrame())
        if validated_df.empty:
            st.warning("Veuillez d'abord produire/valider le mapping SWOT dans l'onglet 2.")
        else:
            synthesis_df = build_theme_synthesis(validated_df)
            st.dataframe(synthesis_df, use_container_width=True)
            output = BytesIO()
            with pd.ExcelWriter(output, engine="openpyxl") as writer:
                synthesis_df.to_excel(writer, index=False, sheet_name="Synthese")
            output.seek(0)
            st.download_button(
                "Télécharger la synthèse",
                data=output.getvalue(),
                file_name=f"synthese_{safe_filename(selected_type)}_{safe_filename(selected_subgroup)}.xlsx",
                mime="application/vnd.openxmlformats-officedocument.spreadsheetml.sheet"
            )


# ============================================================
# PROFESSIONAL AI ENGINE OVERRIDE
# ============================================================
# This block intentionally overrides the previous experimental AI workflow.
# It does NOT modify the print preview / Word / PDF functions above.

ENGINE_RULES = "Expert keyword rules - contrôle thématique"
ENGINE_TFIDF = "TF-IDF semantic theme matching - comparaison lexicale"
ENGINE_HYBRID = "Hybrid consensus model - recommandé"
ENGINE_TRANSFORMER = "Sentence Transformer multilingual - si disponible"


THEME_CONTEXT_ENRICHED = {
    "soutenabilite financiere": "budget financement frais scolarité bourses coût modèle économique viabilité financière ressources financières",
    "gouvernance et leadership": "gouvernance leadership direction administration décisions procédures transparence stratégie représentation responsabilisation autonomie pilotage",
    "strategie academique et qualite d enseignement": "programmes formation enseignement pédagogie académique qualité cours curriculum diplôme accréditation apprentissage langues anglais français trilinguisme",
    "qualite de l enseignement": "enseignement cours pédagogie accompagnement orientation tutorat enseignants apprentissage évaluation formation qualité langues",
    "recherche et innovation": "recherche innovation publications laboratoire doctorat thèse brevets valorisation scientifique projets scientifiques",
    "ressources documentaires et environnement digital": "bibliothèque ressources documentaires numérique digital informatique technologie systèmes plateformes logiciels données outils",
    "ressources documentaires numeriques et informatiques": "bibliothèque ressources documentaires numériques informatiques digital plateformes logiciels technologie systèmes outils données",
    "succes des etudiants": "étudiants recrutement inscriptions accompagnement services support employabilité insertion stage stages réussite parcours diplômés",
    "ressources humaines": "ressources humaines personnel personnels enseignants administratifs recrutement compétences charge travail formation évaluation RH",
    "strategie et mobilite internationales": "international mobilité échanges partenariats étranger diaspora francophonie langues anglais internationalisation",
    "mission societale": "mission sociétale engagement social société communauté responsabilité sociale impact citoyen Liban reconstruction",
    "espace et infrastructures": "campus bâtiments locaux salles infrastructures espaces équipements laboratoires maintenance Beyrouth hôtel-dieu",
    "espace et infrastructure": "campus bâtiment locaux salles infrastructure espaces équipements laboratoires maintenance",
    "environnement de travail": "conditions de travail environnement travail climat collaboration communication interne charge horaires processus administratif",
    "diversite et inclusion": "diversité inclusion égalité accessibilité handicap genre équité équitable",
    "developpement durable odd": "développement durable ODD environnement écologie durable responsabilité environnementale",
    "relation avec l administration et les enseignants": "administration enseignants relation communication écoute coordination service étudiant procédures administratives",
    "representation des etudiants": "représentation étudiants conseil étudiant participation voix des étudiants délégués implication",
    "vie universitaire": "vie universitaire vie étudiante sport activités clubs campus animation aide psychologique expérience étudiante",
    "mobilite internationale": "mobilité internationale échanges erasmus partenariats étranger voyage",
    "aide financiere": "aide financière bourses frais scolarité paiement soutien financier financement",
    "insertion professionnelle": "insertion professionnelle emploi employabilité stage stages carrière entreprises marché travail métier métiers",
    "marche du travail et associations professionnelles": "marché travail emploi employeurs entreprises associations professionnelles compétences stages métiers carrière",
    "marche du travail et relations avec les employeurs": "marché travail emploi employeurs entreprises relations employeurs professionnel compétences stages carrière",
    "concurrence avec les autres universites": "concurrence autres universités compétition attractivité frais scolarité programmes flexibles universités américaines classe A",
    "intelligence artificielle": "intelligence artificielle IA AI ChatGPT automatisation data science analyse générative transformation numérique",
    "reputation et image": "réputation image classement classements visibilité notoriété communication attractivité perception rayonnement jeunes écoles recruteurs",
    "environnement politique et economique emigration": "politique économique crise émigration instabilité Liban inflation sécurité guerre familles parents banque immobilier finance",
}


def theme_context(theme):
    t = clean(theme)
    nt = normalize(t)
    for key, value in THEME_CONTEXT_ENRICHED.items():
        if key in nt or nt in key:
            return value
    return " ".join(tokenize_for_similarity(t))


def theme_document(theme):
    return f"{clean(theme)}. {theme_context(theme)}"


def tfidf_scores_against_themes(answer, themes):
    try:
        from sklearn.feature_extraction.text import TfidfVectorizer
        from sklearn.metrics.pairwise import cosine_similarity
        from sklearn.pipeline import FeatureUnion
    except Exception:
        # Fallback when scikit-learn is unavailable.
        return [simple_similarity_score(answer, t) for t in themes]

    corpus = [clean(answer)] + [theme_document(t) for t in themes]
    try:
        features = FeatureUnion([
            ("word", TfidfVectorizer(ngram_range=(1, 2), min_df=1, strip_accents="unicode", lowercase=True)),
            ("char", TfidfVectorizer(analyzer="char_wb", ngram_range=(3, 5), min_df=1, strip_accents="unicode", lowercase=True)),
        ])
        matrix = features.fit_transform(corpus)
        return cosine_similarity(matrix[0:1], matrix[1:]).flatten().tolist()
    except Exception:
        return [simple_similarity_score(answer, t) for t in themes]


def keyword_rule_scores(answer, themes):
    return [strong_keyword_score(answer, t) for t in themes]


def lexical_overlap_scores(answer, themes):
    answer_tokens = set(tokenize_for_similarity(answer))
    scores = []
    for theme in themes:
        theme_tokens = set(tokenize_for_similarity(theme_document(theme)))
        if not answer_tokens or not theme_tokens:
            scores.append(0.0)
        else:
            overlap = len(answer_tokens & theme_tokens) / max(len(theme_tokens), 1)
            containment = len(answer_tokens & theme_tokens) / max(min(len(answer_tokens), len(theme_tokens)), 1)
            scores.append(min(1.0, 0.6 * overlap + 0.4 * containment))
    return scores


def transformer_scores(answer, themes):
    model = load_multilingual_sentence_model()
    if model is None:
        return None
    try:
        embs = model.encode([clean(answer)] + [theme_document(t) for t in themes], normalize_embeddings=True)
        answer_vec = embs[0]
        return [cosine_from_vectors(answer_vec, v) for v in embs[1:]]
    except Exception:
        return None


def score_theme_candidates(answer, participant_type, swot_element, engine_name):
    themes = official_themes_for(participant_type, swot_element, include_autre=False)
    if not themes:
        return []

    kw = keyword_rule_scores(answer, themes)
    tfidf = tfidf_scores_against_themes(answer, themes)
    lex = lexical_overlap_scores(answer, themes)

    if engine_name == ENGINE_RULES:
        scores = [(0.75 * kw[i]) + (0.25 * lex[i]) for i in range(len(themes))]
    elif engine_name == ENGINE_TFIDF:
        scores = [(0.80 * tfidf[i]) + (0.20 * lex[i]) for i in range(len(themes))]
    elif engine_name == ENGINE_TRANSFORMER:
        tr = transformer_scores(answer, themes)
        if tr is None:
            scores = [(0.55 * tfidf[i]) + (0.30 * kw[i]) + (0.15 * lex[i]) for i in range(len(themes))]
        else:
            scores = [(0.65 * tr[i]) + (0.25 * kw[i]) + (0.10 * lex[i]) for i in range(len(themes))]
    else:
        # Recommended hybrid: domain keywords prevent illogical mappings; TF-IDF helps with wording variants.
        scores = [(0.42 * kw[i]) + (0.38 * tfidf[i]) + (0.20 * lex[i]) for i in range(len(themes))]

    ranked = sorted(zip(themes, scores, kw, tfidf, lex), key=lambda x: x[1], reverse=True)
    return ranked


def calibrated_confidence(ranked):
    if not ranked:
        return 0.0
    best = ranked[0][1]
    second = ranked[1][1] if len(ranked) > 1 else 0
    margin = max(best - second, 0)
    conf = 25 + best * 60 + margin * 80
    if best < 0.18:
        conf = min(conf, 45)
    if margin < 0.025:
        conf = min(conf, 55)
    elif margin < 0.055:
        conf = min(conf, 68)
    return round(max(5, min(conf, 97)), 1)


def classify_theme_with_engine(answer, participant_type, swot_element, engine_name):
    ranked = score_theme_candidates(answer, participant_type, swot_element, engine_name)
    if not ranked:
        return "", 0.0, "", "Aucun thème officiel disponible"
    best_theme = ranked[0][0]
    conf = calibrated_confidence(ranked)
    top3 = " | ".join([f"{t} ({round(s * 100, 1)})" for t, s, _, _, _ in ranked[:3]])
    evidence = []
    if ranked[0][2] > 0:
        evidence.append("mots-clés thématiques")
    if ranked[0][3] > 0:
        evidence.append("similarité TF-IDF")
    if ranked[0][4] > 0:
        evidence.append("recouvrement lexical")
    evidence_txt = ", ".join(evidence) if evidence else "évidence faible"
    return best_theme, conf, top3, evidence_txt


def available_ai_engines():
    engines = [ENGINE_RULES, ENGINE_TFIDF, ENGINE_HYBRID]
    if load_multilingual_sentence_model() is not None:
        engines.append(ENGINE_TRANSFORMER)
    return engines


def build_classification_with_engine(df_scope, engine_name):
    rows = []
    for idx, row in df_scope.iterrows():
        swot_element = row.get("swot_element", "")
        if swot_element not in ["Forces", "Faiblesses", "Opportunités", "Menaces"]:
            continue
        answer = clean(row.get("Final_Answer", ""))
        if not answer:
            continue
        theme, conf, top3, evidence = classify_theme_with_engine(
            answer=answer,
            participant_type=row.get("Respondent_Type", ""),
            swot_element=swot_element,
            engine_name=engine_name,
        )
        rows.append({
            "row_id": idx,
            "Respondent_Type": row.get("Respondent_Type", ""),
            "groupe": row.get("groupe", ""),
            "participants": row.get("participants", ""),
            "swot_element": swot_element,
            "question": row.get("question", ""),
            "Final_Answer": answer,
            "AI_Suggested_Theme": theme,
            "Confidence": conf,
            "Top_3_Themes": top3,
            "Evidence": evidence,
            "Validation_Status": "À valider" if conf < 70 else "Suggestion forte",
            "Validated_Theme": theme,
            "New_Theme": "",
        })
    return pd.DataFrame(rows)


def compare_ai_engines(df_scope):
    engines = available_ai_engines()
    outputs = []
    for engine in engines:
        pred = build_classification_with_engine(df_scope, engine)
        if pred.empty:
            continue
        small = pred[[
            "row_id", "Respondent_Type", "groupe", "participants", "swot_element", "question", "Final_Answer",
            "AI_Suggested_Theme", "Confidence", "Top_3_Themes", "Evidence", "Validation_Status"
        ]].copy()
        suffix = safe_filename(engine)
        small = small.rename(columns={
            "AI_Suggested_Theme": f"Theme_{suffix}",
            "Confidence": f"Confidence_{suffix}",
            "Top_3_Themes": f"Top3_{suffix}",
            "Evidence": f"Evidence_{suffix}",
            "Validation_Status": f"Status_{suffix}",
        })
        outputs.append((engine, small))

    if not outputs:
        return pd.DataFrame(), pd.DataFrame(), engines

    comparison = outputs[0][1]
    base_cols = ["Respondent_Type", "groupe", "participants", "swot_element", "question", "Final_Answer"]
    for _, small in outputs[1:]:
        comparison = comparison.merge(small.drop(columns=base_cols), on="row_id", how="outer")

    theme_cols = [f"Theme_{safe_filename(engine)}" for engine, _ in outputs]
    conf_cols = [f"Confidence_{safe_filename(engine)}" for engine, _ in outputs]

    def agreement_count(row):
        vals = [clean(row.get(c, "")) for c in theme_cols if clean(row.get(c, ""))]
        if not vals:
            return 0
        return max(vals.count(v) for v in set(vals))

    def consensus_theme(row):
        vals = [clean(row.get(c, "")) for c in theme_cols if clean(row.get(c, ""))]
        if not vals:
            return ""
        return max(set(vals), key=vals.count)

    comparison["Consensus_Theme"] = comparison.apply(consensus_theme, axis=1)
    comparison["Agreement_Count"] = comparison.apply(agreement_count, axis=1)
    comparison["Agreement_Rate"] = round(comparison["Agreement_Count"] / max(len(theme_cols), 1) * 100, 1)
    comparison["Consensus_Status"] = comparison["Agreement_Rate"].apply(lambda x: "Accord fort" if x >= 67 else "Désaccord - validation prioritaire")
    comparison["Max_Confidence"] = comparison[conf_cols].apply(lambda row: pd.to_numeric(row, errors="coerce").max(), axis=1)

    summary_rows = []
    for engine, _ in outputs:
        suffix = safe_filename(engine)
        tcol = f"Theme_{suffix}"
        ccol = f"Confidence_{suffix}"
        summary_rows.append({
            "Moteur": engine,
            "Réponses classifiées": int(comparison[tcol].notna().sum()),
            "Confiance moyenne": round(pd.to_numeric(comparison[ccol], errors="coerce").mean(), 1),
            "Confiance médiane": round(pd.to_numeric(comparison[ccol], errors="coerce").median(), 1),
            "Faible confiance (<70%)": int((pd.to_numeric(comparison[ccol], errors="coerce") < 70).sum()),
        })
    summary_rows.append({
        "Moteur": "Consensus entre moteurs",
        "Réponses classifiées": len(comparison),
        "Confiance moyenne": round(comparison["Agreement_Rate"].mean(), 1),
        "Confiance médiane": round(comparison["Agreement_Rate"].median(), 1),
        "Faible confiance (<70%)": int((comparison["Agreement_Rate"] < 67).sum()),
    })
    return comparison, pd.DataFrame(summary_rows), engines


def build_consensus_classification(df_scope, selected_engine=ENGINE_HYBRID):
    comparison, _, engines = compare_ai_engines(df_scope)
    if comparison.empty:
        return pd.DataFrame()
    preferred_col = f"Theme_{safe_filename(selected_engine)}"
    preferred_conf = f"Confidence_{safe_filename(selected_engine)}"
    preferred_top3 = f"Top3_{safe_filename(selected_engine)}"
    preferred_evidence = f"Evidence_{safe_filename(selected_engine)}"
    if preferred_col not in comparison.columns:
        preferred_col = "Consensus_Theme"
        preferred_conf = "Max_Confidence"
        preferred_top3 = ""
        preferred_evidence = ""
    rows = []
    for _, row in comparison.iterrows():
        top3_value = row.get(preferred_top3, "") if preferred_top3 else ""
        evidence_value = row.get(preferred_evidence, "") if preferred_evidence else "Consensus entre moteurs"
        rows.append({
            "row_id": row["row_id"],
            "Respondent_Type": row.get("Respondent_Type", ""),
            "groupe": row.get("groupe", ""),
            "participants": row.get("participants", ""),
            "swot_element": row.get("swot_element", ""),
            "question": row.get("question", ""),
            "Final_Answer": row.get("Final_Answer", ""),
            "AI_Suggested_Theme": row.get(preferred_col, ""),
            "Confidence": row.get(preferred_conf, ""),
            "Consensus_Theme": row.get("Consensus_Theme", ""),
            "Agreement_Rate": row.get("Agreement_Rate", ""),
            "Consensus_Status": row.get("Consensus_Status", ""),
            "Top_3_Themes": top3_value,
            "Evidence": evidence_value,
            "Validated_Theme": row.get(preferred_col, ""),
            "New_Theme": "",
        })
    return pd.DataFrame(rows)


def render_mapping_validation(classified_df):
    if classified_df.empty:
        st.warning("Aucune réponse SWOT à classifier.")
        return classified_df

    validated_rows = []
    st.markdown("### Validation experte des classifications")
    st.caption("Les cas avec faible accord entre moteurs ou faible confiance sont prioritaires à vérifier.")

    sort_cols = [c for c in ["Consensus_Status", "Confidence"] if c in classified_df.columns]
    if sort_cols:
        classified_df = classified_df.sort_values(sort_cols, ascending=[False, True][:len(sort_cols)])

    for _, row in classified_df.iterrows():
        swot_element = row["swot_element"]
        row_key = f"{row['row_id']}_{swot_element}"
        themes = official_themes_for(row["Respondent_Type"], swot_element, include_autre=False)
        suggested = row.get("AI_Suggested_Theme", "")
        default_index = themes.index(suggested) if suggested in themes else 0
        title = f"{row['Respondent_Type']} | {row['groupe']} | {swot_element} | {row.get('Consensus_Status', '')} | {row.get('Confidence', '')}%"
        with st.expander(title, expanded=False):
            st.markdown("**Réponse corrigée**")
            st.write(row["Final_Answer"])
            st.markdown(f"**Thème proposé :** {suggested}")
            if clean(row.get("Top_3_Themes", "")):
                st.caption("Top 3 : " + clean(row.get("Top_3_Themes", "")))
            if clean(row.get("Evidence", "")):
                st.caption("Évidence : " + clean(row.get("Evidence", "")))
            selected_theme = st.selectbox(
                "Thème validé par l'expert",
                options=themes,
                index=default_index,
                key=f"theme_{row_key}"
            )
            new_theme = st.text_input(
                "Ajouter un nouveau thème si aucun thème officiel ne convient",
                value="",
                key=f"new_theme_{row_key}"
            )
            final_theme = clean(new_theme) if clean(new_theme) else selected_theme
            out = row.to_dict()
            out["Validated_Theme"] = final_theme
            out["New_Theme"] = clean(new_theme)
            validated_rows.append(out)
    return pd.DataFrame(validated_rows)


def professional_ai_header():
    html_block(f"""
    <style>
    .ai-hero {{
        background: linear-gradient(135deg, #001F5B 0%, #123C7C 55%, #8B1538 100%);
        color: white;
        border-radius: 22px;
        padding: 30px 34px;
        margin: 8px 0 24px 0;
        box-shadow: 0 16px 38px rgba(0,31,91,0.22);
    }}
    .ai-hero h2 {{ color: white !important; margin: 0 0 8px 0; font-size: 32px; font-weight: 900; }}
    .ai-hero p {{ color: #EAF2F8; margin: 0; font-size: 16px; line-height: 1.5; }}
    .ai-info-box {{
        background: #EAF2F8;
        color: #003B7A;
        border-left: 6px solid {USJ_BLUE};
        border-radius: 12px;
        padding: 14px 18px;
        margin: 14px 0;
        font-size: 15px;
    }}
    .ai-warning {{
        background: #FFF7E6;
        border-left: 6px solid #F59F00;
        border-radius: 12px;
        padding: 14px 18px;
        color: #7A4B00;
        font-weight: 700;
        margin: 12px 0;
    }}
    .ai-kpi {{
        background: white;
        border: 1px solid #D0D6E0;
        border-radius: 16px;
        padding: 16px 18px;
        box-shadow: 0 8px 20px rgba(0,0,0,0.045);
        min-height: 92px;
    }}
    .ai-kpi-number {{ color: {USJ_RED}; font-size: 28px; font-weight: 900; }}
    .ai-kpi-label {{ color: #64748B; font-size: 13px; font-weight: 700; margin-top: 3px; }}
    </style>
    <div class="ai-hero">
        <h2>Plateforme IA supervisée - Analyse thématique stratégique</h2>
        <p>Classification sur toutes les réponses importées, comparaison réelle de moteurs distincts, consensus, validation experte et synthèse institutionnelle.</p>
    </div>
    """)


def ai_platform_page(show_header=True, df_preloaded=None):
    html_block(APP_CSS)
    professional_ai_header()

    if df_preloaded is None:
        uploaded_file = st.file_uploader("Uploader le fichier Excel corrigé pour l'analyse IA", type=["xlsx", "xls"], key="ai_excel_upload")
        if uploaded_file is None:
            st.stop()
        df = read_corrected_excel(uploaded_file)
    else:
        df = df_preloaded.copy()

    df_scope = df.copy()
    swot_df = df_scope[df_scope["swot_element"].isin(["Forces", "Faiblesses", "Opportunités", "Menaces"])]

    c1, c2, c3, c4 = st.columns(4)
    with c1:
        st.markdown(f'<div class="ai-kpi"><div class="ai-kpi-number">{len(df_scope)}</div><div class="ai-kpi-label">Réponses importées</div></div>', unsafe_allow_html=True)
    with c2:
        st.markdown(f'<div class="ai-kpi"><div class="ai-kpi-number">{len(swot_df)}</div><div class="ai-kpi-label">Réponses SWOT analysées</div></div>', unsafe_allow_html=True)
    with c3:
        st.markdown(f'<div class="ai-kpi"><div class="ai-kpi-number">{df_scope["Respondent_Type"].nunique()}</div><div class="ai-kpi-label">Types de participants</div></div>', unsafe_allow_html=True)
    with c4:
        st.markdown(f'<div class="ai-kpi"><div class="ai-kpi-number">{len(available_ai_engines())}</div><div class="ai-kpi-label">Moteurs comparables</div></div>', unsafe_allow_html=True)

    st.markdown("""
    <div class="ai-info-box">
        <b>Logique professionnelle :</b> l'IA est appliquée à toute la base. Les moteurs sont réellement différents : règles expertes, TF-IDF lexical/sémantique, modèle hybride par consensus et, si installé, Sentence Transformer multilingue. Les classifications faibles ou contradictoires sont marquées pour validation humaine.
    </div>
    """, unsafe_allow_html=True)

    tab_benchmark, tab_mapping, tab_duplicates, tab_synthesis = st.tabs([
        "1. Benchmark IA global", "2. Mapping SWOT supervisé", "3. Doublons", "4. Synthèse"
    ])

    with tab_benchmark:
        st.subheader("1. Benchmark IA global sur toutes les réponses")
        st.caption("Cette étape compare les moteurs sur la même base. Sans fichier validé, on mesure le consensus et les conflits; avec fichier validé, on mesure l'accuracy réelle.")

        if st.button("Lancer le benchmark IA", key="run_professional_benchmark"):
            comparison_df, summary_df, engines = compare_ai_engines(df_scope)
            st.session_state["engine_comparison_df"] = comparison_df
            st.session_state["engine_summary_df"] = summary_df
            st.session_state["available_engines"] = engines
            st.session_state["active_ai_engine"] = ENGINE_HYBRID

        summary_df = st.session_state.get("engine_summary_df", pd.DataFrame())
        comparison_df = st.session_state.get("engine_comparison_df", pd.DataFrame())

        if not summary_df.empty:
            st.markdown("#### Synthèse des moteurs")
            st.dataframe(summary_df, use_container_width=True)
            conflict_count = int((comparison_df.get("Agreement_Rate", pd.Series(dtype=float)) < 67).sum()) if not comparison_df.empty else 0
            st.info(f"Moteur recommandé par défaut : {ENGINE_HYBRID}. Cas en désaccord à vérifier en priorité : {conflict_count}.")

            selected_engine = st.selectbox(
                "Choisir le moteur actif pour le mapping",
                options=available_ai_engines(),
                index=available_ai_engines().index(ENGINE_HYBRID) if ENGINE_HYBRID in available_ai_engines() else 0,
                key="selected_professional_engine"
            )
            if st.button("Activer le moteur sélectionné", key="activate_professional_engine"):
                st.session_state["active_ai_engine"] = selected_engine
                st.success(f"Moteur actif : {selected_engine}")

            display_cols = [
                "Respondent_Type", "groupe", "swot_element", "Final_Answer",
                "Consensus_Theme", "Agreement_Rate", "Consensus_Status", "Max_Confidence"
            ]
            existing_cols = [c for c in display_cols if c in comparison_df.columns]
            st.markdown("#### Résultats de consensus")
            st.dataframe(comparison_df[existing_cols], use_container_width=True)

            output = BytesIO()
            with pd.ExcelWriter(output, engine="openpyxl") as writer:
                comparison_df.to_excel(writer, index=False, sheet_name="Benchmark moteurs")
                summary_df.to_excel(writer, index=False, sheet_name="Synthese")
            output.seek(0)
            st.download_button(
                "Télécharger le benchmark IA",
                data=output.getvalue(),
                file_name="benchmark_IA_global.xlsx",
                mime="application/vnd.openxmlformats-officedocument.spreadsheetml.sheet"
            )
        else:
            st.warning("Lancez d'abord le benchmark IA. Le mapping utilise ensuite le moteur actif.")

    with tab_mapping:
        st.subheader("2. Mapping SWOT supervisé")
        active_engine = st.session_state.get("active_ai_engine", "")
        if not active_engine:
            st.markdown('<div class="ai-warning">Lancez d’abord le benchmark IA dans l’onglet 1.</div>', unsafe_allow_html=True)
        else:
            st.caption(f"Moteur actif : {active_engine}")
            if st.button("Classifier toutes les réponses SWOT", key="run_consensus_mapping"):
                st.session_state["classified_df"] = build_consensus_classification(df_scope, active_engine)
            classified_df = st.session_state.get("classified_df", pd.DataFrame())
            if not classified_df.empty:
                st.markdown("#### Aperçu des classifications")
                preview_cols = ["Respondent_Type", "groupe", "swot_element", "Final_Answer", "AI_Suggested_Theme", "Confidence", "Consensus_Status"]
                st.dataframe(classified_df[[c for c in preview_cols if c in classified_df.columns]], use_container_width=True)
                validated_df = render_mapping_validation(classified_df)
                st.session_state["validated_mapping_df"] = validated_df
                if not validated_df.empty:
                    output = BytesIO()
                    with pd.ExcelWriter(output, engine="openpyxl") as writer:
                        validated_df.to_excel(writer, index=False, sheet_name="Mapping valide")
                    output.seek(0)
                    st.download_button(
                        "Télécharger le mapping validé",
                        data=output.getvalue(),
                        file_name="mapping_SWOT_valide_global.xlsx",
                        mime="application/vnd.openxmlformats-officedocument.spreadsheetml.sheet"
                    )
            else:
                st.caption("Cliquez sur « Classifier toutes les réponses SWOT » pour lancer le mapping.")

    with tab_duplicates:
        st.subheader("3. Détection stricte des doublons")
        if st.button("Détecter les doublons", key="detect_duplicates_professional"):
            st.session_state["duplicates_df"] = detect_duplicate_pairs(df_scope)
        dup_df = st.session_state.get("duplicates_df", pd.DataFrame())
        if not dup_df.empty:
            st.dataframe(dup_df, use_container_width=True)
            st.download_button(
                "Télécharger les doublons détectés",
                data=dup_df.to_csv(index=False).encode("utf-8-sig"),
                file_name="doublons_detectes_global.csv",
                mime="text/csv"
            )
        else:
            st.caption("Aucun doublon affiché pour le moment.")

    with tab_synthesis:
        st.subheader("4. Synthèse automatique par thème validé")
        validated_df = st.session_state.get("validated_mapping_df", pd.DataFrame())
        if validated_df.empty:
            st.warning("Veuillez d'abord produire ou valider le mapping SWOT.")
        else:
            synthesis_df = build_theme_synthesis(validated_df)
            st.dataframe(synthesis_df, use_container_width=True)
            output = BytesIO()
            with pd.ExcelWriter(output, engine="openpyxl") as writer:
                synthesis_df.to_excel(writer, index=False, sheet_name="Synthese")
            output.seek(0)
            st.download_button(
                "Télécharger la synthèse",
                data=output.getvalue(),
                file_name="synthese_theme_SWOT_global.xlsx",
                mime="application/vnd.openxmlformats-officedocument.spreadsheetml.sheet"
            )

def print_preview_page(show_header=True, df_preloaded=None):
    if show_header:
        html_block(APP_CSS)

        html_block(f"""
        <div class="screen-only">
            <h1 style="margin-bottom:0;">{APP_TITLE}</h1>
            <h3 style="color:{USJ_RED} !important; margin-top:4px;">
                Lecture Excel - Print Preview des réponses corrigées
            </h3>
        </div>
        """)

    if df_preloaded is None:
        uploaded_file = st.file_uploader(
            "Uploader le fichier Excel corrigé",
            type=["xlsx", "xls"]
        )

        if uploaded_file is None:
            st.info(f"Veuillez uploader le fichier Excel contenant la feuille : {SHEET_NAME}")
            st.stop()

        df = read_corrected_excel(uploaded_file)
    else:
        df = df_preloaded.copy()

    participant_types = sorted(df["Respondent_Type"].dropna().unique().tolist())

    c1, c2, c3, c4 = st.columns([1.2, 1.3, 1.2, 1])

    with c1:
        selected_type = st.selectbox("Type de participants", participant_types)

    df_type = df[df["Respondent_Type"] == selected_type].copy()

    with c2:
        report_mode = st.selectbox(
            "Mode d'affichage",
            [
                "Tous les sous-groupes du type sélectionné",
                "Un sous-groupe spécifique"
            ]
        )

    subgroups = sorted(df_type["groupe"].dropna().unique().tolist())

    with c3:
        if report_mode == "Un sous-groupe spécifique":
            selected_subgroup = st.selectbox("Sous-groupe", subgroups)
        else:
            selected_subgroup = "Tous les sous-groupes"

    with c4:
        hide_names = st.checkbox("Hide participants names", value=False)

    if report_mode == "Un sous-groupe spécifique":
        df_report = df_type[df_type["groupe"] == selected_subgroup].copy()
        title_label = selected_subgroup
    else:
        df_report = df_type.copy()
        title_label = "Tous les sous-groupes"

    html_report = build_one_report_html(
        df_group=df_report,
        participant_type=selected_type,
        title_label=title_label,
        hide_names=hide_names
    )

    names_suffix = "no names" if hide_names else "with names"

    file_base = (
        f"{safe_filename(selected_type)}-"
        f"{safe_filename(title_label)}-"
        f"{safe_filename(names_suffix)}"
    )

    docx_bytes = build_word_docx(
        df_group=df_report,
        participant_type=selected_type,
        title_label=title_label,
        hide_names=hide_names
    )
    docx_base64 = base64.b64encode(docx_bytes).decode("utf-8")

    full_html = build_full_html(
        html_report=html_report,
        selected_type=selected_type,
        selected_label=title_label,
        docx_base64=docx_base64,
        docx_filename=f"{file_base}.docx"
    )

    components.html(
        full_html,
        height=1300,
        scrolling=True
    )



def main():
    st.set_page_config(
        page_title=APP_TITLE,
        layout="wide",
        initial_sidebar_state="collapsed"
    )

    html_block(APP_CSS)

    html_block(f"""
    <style>
    .usj-dashboard-header {{
        background: linear-gradient(135deg, #001F5B 0%, #1F3C88 70%, #8B1538 100%);
        border-radius: 22px;
        padding: 28px 34px;
        margin-bottom: 24px;
        box-shadow: 0 14px 34px rgba(0,31,91,0.20);
    }}
    .usj-dashboard-header h1 {{
        margin: 0;
        color: white !important;
        font-size: 36px;
        font-weight: 900;
        letter-spacing: 0.2px;
    }}
    .usj-dashboard-header h3 {{
        margin: 8px 0 0 0;
        color: #EAF2F8 !important;
        font-weight: 800;
        font-size: 20px;
    }}
    .usj-step-card {{
        background: #ffffff;
        border: 1px solid #D0D6E0;
        border-radius: 16px;
        padding: 18px 20px;
        margin: 14px 0 18px 0;
        box-shadow: 0 8px 22px rgba(0,0,0,0.045);
    }}
    .usj-step-title {{
        color: {USJ_BLUE};
        font-size: 19px;
        font-weight: 900;
        margin-bottom: 5px;
    }}
    .usj-step-subtitle {{
        color: #667085;
        font-size: 14px;
        margin-bottom: 0;
    }}
    .usj-status-box {{
        background: #F0F7FF;
        border: 1px solid #CFE4FF;
        border-left: 6px solid {USJ_BLUE};
        border-radius: 14px;
        padding: 14px 18px;
        margin: 10px 0 18px 0;
        color: {USJ_TEXT};
        font-weight: 800;
    }}
    div[role="radiogroup"] label {{
        background: #ffffff;
        border: 1px solid #D0D6E0;
        border-radius: 14px;
        padding: 11px 18px;
        margin-right: 10px;
        box-shadow: 0 5px 15px rgba(0,0,0,0.045);
    }}
    </style>
    <div class="usj-dashboard-header">
        <h1>{APP_TITLE}</h1>
        <h3>Plateforme Focus Groups USJ 2032</h3>
    </div>
    """)

    html_block("""
    <div class="usj-step-card">
        <div class="usj-step-title">1. Uploader le fichier Excel corrigé</div>
        <div class="usj-step-subtitle">Le même fichier sera utilisé ensuite pour le print preview / rapports Word ou pour la plateforme IA supervisée.</div>
    </div>
    """)

    uploaded_file = st.file_uploader(
        "Uploader le fichier Excel corrigé",
        type=["xlsx", "xls"],
        key="main_excel_upload"
    )

    if uploaded_file is None:
        st.info(f"Veuillez uploader le fichier Excel contenant la feuille : {SHEET_NAME}")
        st.stop()

    df = read_corrected_excel(uploaded_file)

    nb_rows = len(df)
    nb_types = df["Respondent_Type"].nunique()
    nb_groups = df["groupe"].nunique()

    html_block(f"""
    <div class="usj-status-box">
        Fichier chargé avec succès : {nb_rows} réponses, {nb_types} types de participants, {nb_groups} sous-groupes.
    </div>
    """)

    html_block("""
    <div class="usj-step-card">
        <div class="usj-step-title">2. Choisir le module à utiliser</div>
        <div class="usj-step-subtitle">Vous pouvez changer de module sans uploader à nouveau le fichier.</div>
    </div>
    """)

    page = st.radio(
        "Choisir le module",
        [
            "Print preview / rapports Word",
            "Plateforme IA supervisée"
        ],
        horizontal=True,
        label_visibility="collapsed",
        key="main_module_choice"
    )

    st.markdown("---")

    if page == "Print preview / rapports Word":
        print_preview_page(show_header=False, df_preloaded=df)
    else:
        ai_platform_page(show_header=False, df_preloaded=df)


if __name__ == "__main__":
    main()
